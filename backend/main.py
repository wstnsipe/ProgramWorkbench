import json
import os
import re
import xml.etree.ElementTree as ET
import zipfile
from datetime import datetime, timezone
from pathlib import Path

import yaml
from docx import Document as DocxDocument
from fastapi import FastAPI, Depends, HTTPException, Query, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from sqlalchemy import text, func
from sqlalchemy.orm import Session

from database import engine, get_db
from document_templates import get_template, normalize_llm_output, build_template_contract
from docx_builder import (
    build_acq_strategy, build_acq_strategy_smart, build_mosa_conformance_plan, build_rfi, build_sep,
    build_smart_rfi, build_smart_mosa_conformance_plan, build_sep_smart,
)
from models import Base, Program, ProgramBrief, ProgramFile, ProgramAnswer, Module, ProgramDocument, FileText, FileChunk, RagChunk
from schemas import (
    ProgramCreate, ProgramOut, ProgramUpdate, ProgramBriefIn, ProgramBriefOut, ProgramFileOut,
    WizardOut, WizardAnswersIn, ModuleIn, ModuleOut,
    GenerateDocRequest, DocumentOut,
    FileTextOut, FileChunkOut, KnowledgeSummaryOut, KnowledgeFileSummary, KnowledgeStats,
    ReextractFileResult, ReextractOut, KnowledgeSearchResult,
    RfiPlan, AcqStrategyPlan,
    ContextBuildOut, MosaPlan, SepPlan, KnowledgeIndexOut,
    KnowledgeTopFile, KnowledgeStatusOut,
)
from llm.context_builder import build_program_context, load_context, context_summary_and_gaps
from llm.retrieval import (
    index_reference_docs, index_program_files, retrieve_chunks,
    embed_texts, retrieve_chunks_vector,
    VECTOR_CHUNK_SIZE, VECTOR_CHUNK_OVERLAP,
)

app = FastAPI()


@app.on_event("startup")
def _on_startup() -> None:
    _startup_index_reference_docs()


_cors_origins_env = os.environ.get("FRONTEND_URL", "")
_cors_origins = [o.strip() for o in _cors_origins_env.split(",") if o.strip()] if _cors_origins_env else []
_cors_origins += ["http://localhost:5173", "http://localhost:5174", "http://localhost:5175"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins,
    allow_origin_regex=r"https://.*\.vercel\.app",
    allow_methods=["*"],
    allow_headers=["*"],
)

# Enable pgvector extension (must precede create_all so the vector column type is available)
with engine.connect() as _conn:
    try:
        _conn.execute(text("CREATE EXTENSION IF NOT EXISTS vector"))
        _conn.commit()
    except Exception:
        _conn.rollback()

# Create tables automatically (dev mode)
Base.metadata.create_all(bind=engine)

# Idempotent schema migrations (ALTER TABLE is safe to retry; errors = column exists)
with engine.connect() as _conn:
    for _stmt in [
        "ALTER TABLE program_files ADD COLUMN extracted_text TEXT",
        "ALTER TABLE file_text ADD COLUMN meta_json TEXT",
        "ALTER TABLE program_files ADD COLUMN source_type TEXT NOT NULL DEFAULT 'program_input'",
        "ALTER TABLE file_chunks ADD COLUMN source_type TEXT NOT NULL DEFAULT 'program_input'",
        "ALTER TABLE file_chunks ADD COLUMN embedding vector(1536)",
        "ALTER TABLE modules ADD COLUMN description TEXT",
        "ALTER TABLE modules ADD COLUMN future_recompete BOOLEAN NOT NULL DEFAULT FALSE",
        # programs — new columns (IF NOT EXISTS avoids error if already present)
        "ALTER TABLE programs ADD COLUMN IF NOT EXISTS service_branch TEXT",
        "ALTER TABLE programs ADD COLUMN IF NOT EXISTS army_pae TEXT",
        "ALTER TABLE programs ADD COLUMN IF NOT EXISTS army_branch TEXT",
        "ALTER TABLE programs ADD COLUMN IF NOT EXISTS mig_id TEXT",
        # program_briefs — columns added after initial deploy
        "ALTER TABLE program_briefs ADD COLUMN IF NOT EXISTS timeline_months INTEGER",
        "ALTER TABLE program_briefs ADD COLUMN IF NOT EXISTS software_involved BOOLEAN",
        "ALTER TABLE program_briefs ADD COLUMN IF NOT EXISTS similar_programs_exist BOOLEAN",
        # program_standards — applicability columns
        "ALTER TABLE program_standards ADD COLUMN IF NOT EXISTS applies_to_modules BOOLEAN NOT NULL DEFAULT FALSE",
        "ALTER TABLE program_standards ADD COLUMN IF NOT EXISTS applies_to_interfaces BOOLEAN NOT NULL DEFAULT FALSE",
        # program_documents — section output snapshot for per-section regeneration
        "ALTER TABLE program_documents ADD COLUMN IF NOT EXISTS assembled_json TEXT",
        # Embedding cache table (CREATE IF NOT EXISTS — fully idempotent)
        """
        CREATE TABLE IF NOT EXISTS embedding_cache (
            key        TEXT PRIMARY KEY,
            embedding  vector(1536) NOT NULL,
            created_at TIMESTAMPTZ NOT NULL DEFAULT now()
        )
        """,
    ]:
        try:
            _conn.execute(text(_stmt))
            _conn.commit()
        except Exception:
            _conn.rollback()

DATA_DIR = Path(__file__).parent / "data"
REFERENCE_DOCS_DIR = Path(__file__).parent / "reference_docs"

# File types accepted for direct upload
ALLOWED_EXTENSIONS = {".docx", ".doc", ".pdf", ".txt", ".md"}

# Index global reference docs into RagChunk on startup (idempotent – skips if already done)
def _startup_index_reference_docs() -> None:
    from database import SessionLocal
    db = SessionLocal()
    try:
        index_reference_docs(REFERENCE_DOCS_DIR, db, _extract_text)
    finally:
        db.close()

_questions_cache: list | None = None

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200


def _normalize_text(raw: str) -> str:
    """Collapse repeated spaces/tabs; reduce 3+ newlines to a paragraph break."""
    raw = re.sub(r"[ \t]+", " ", raw)
    raw = re.sub(r"\n{3,}", "\n\n", raw)
    return raw.strip()


def _is_low_quality(text: str) -> bool:
    """Return True when extracted text is too sparse to be useful.

    Triggers fallback when either:
    - fewer than 200 alphabetic characters, or
    - alphabetic ratio (alpha / total chars) is below 0.1.
    """
    if not text:
        return True
    alpha = sum(1 for c in text if c.isalpha())
    if alpha < 200:
        return True
    if alpha / len(text) < 0.1:
        return True
    return False


def _extract_docx_xml_fallback(file_path: Path) -> str:
    """Fallback DOCX extractor: treat the file as a ZIP and parse raw XML.

    Reads word/document.xml plus any header/footer XML files, collects all
    <w:t> text tokens, inserts newlines at <w:p> paragraph boundaries, then
    applies noise cleanup (drops underscore/dash/dot-only lines and very short
    punctuation-only lines).
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W_P = f"{{{W}}}p"
    W_T = f"{{{W}}}t"

    def _paragraphs_from_xml(xml_bytes: bytes) -> list[str]:
        try:
            root = ET.fromstring(xml_bytes)
        except ET.ParseError:
            return []
        paragraphs: list[str] = []
        for p_elem in root.iter(W_P):
            tokens = [t.text for t in p_elem.iter(W_T) if t.text]
            paragraphs.append(" ".join(tokens).strip())
        return paragraphs

    all_paragraphs: list[str] = []
    try:
        with zipfile.ZipFile(str(file_path)) as z:
            names = set(z.namelist())
            if "word/document.xml" in names:
                all_paragraphs.extend(_paragraphs_from_xml(z.read("word/document.xml")))
            for name in sorted(names):
                if (
                    (name.startswith("word/header") or name.startswith("word/footer"))
                    and name.endswith(".xml")
                ):
                    all_paragraphs.extend(_paragraphs_from_xml(z.read(name)))
    except (zipfile.BadZipFile, KeyError, OSError):
        return ""

    clean: list[str] = []
    for para in all_paragraphs:
        s = para.strip()
        if not s:
            continue
        # Drop lines that are only underscores / dashes / dots / whitespace
        if re.fullmatch(r"[_\-.\s]+", s):
            continue
        # Drop very short punctuation-only lines (≤3 chars, no alphanumeric)
        if len(s) <= 3 and not any(c.isalnum() for c in s):
            continue
        clean.append(s)

    return _normalize_text("\n".join(clean))


def _extract_text(file_path: Path, ext: str) -> str:
    """
    Extract plain text from a saved file.

    Supported:
      .txt / .md  — read UTF-8 (ignore errors)
      .docx       — python-docx; paragraphs + table cells (row cells joined with " | ")
      .pdf        — pypdf; join page text with double newlines
      other       — returns "" without raising

    Raises on I/O or library errors so the caller can record them in meta_json.
    """
    if ext in (".txt", ".md"):
        raw = file_path.read_text(encoding="utf-8", errors="ignore")
        return _normalize_text(raw)

    if ext == ".docx":
        doc = DocxDocument(str(file_path))

        # Collect body paragraphs in document order via the XML body children,
        # so tables are interleaved with paragraphs rather than appended after.
        from docx.oxml.ns import qn
        parts: list[str] = []
        for child in doc.element.body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                text = "".join(run.text for run in child.iter(qn("w:t")))
                text = text.strip()
                # Drop lines that are only underscores / dots (form-fill blanks)
                if text and not re.fullmatch(r"[_.\s]+", text):
                    parts.append(text)
            elif tag == "tbl":
                for tr in child.iter(qn("w:tr")):
                    cells: list[str] = []
                    for tc in tr.iter(qn("w:tc")):
                        cell_text = "".join(
                            t.text for t in tc.iter(qn("w:t"))
                        ).strip()
                        if cell_text and not re.fullmatch(r"[_.\s]+", cell_text):
                            cells.append(cell_text)
                    if cells:
                        parts.append(" | ".join(cells))

        primary = _normalize_text("\n".join(parts))
        if _is_low_quality(primary):
            return _extract_docx_xml_fallback(file_path)
        return primary

    if ext == ".pdf":
        from pypdf import PdfReader  # lazy import; avoids hard dep on startup
        reader = PdfReader(str(file_path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return _normalize_text("\n\n".join(pages))

    return ""  # .doc and any other allowed-but-unsupported type


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Split text into overlapping chunks of approximately chunk_size characters."""
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = end - overlap
    return chunks


def _load_questions() -> list:
    global _questions_cache
    if _questions_cache is None:
        config_path = Path(__file__).parent / "config" / "questions.yaml"
        with open(config_path) as f:
            _questions_cache = yaml.safe_load(f)["questions"]
    return _questions_cache


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/programs", response_model=ProgramOut)
def create_program(payload: ProgramCreate, db: Session = Depends(get_db)):
    program = Program(name=payload.name)
    db.add(program)
    db.commit()
    db.refresh(program)
    return program


@app.get("/programs", response_model=list[ProgramOut])
def list_programs(db: Session = Depends(get_db)):
    return db.query(Program).order_by(Program.id.desc()).all()


@app.get("/programs/{program_id}", response_model=ProgramOut)
def get_program(program_id: int, db: Session = Depends(get_db)):
    program = db.query(Program).filter(Program.id == program_id).first()
    if not program:
        raise HTTPException(status_code=404, detail="Program not found")
    return program


@app.patch("/programs/{program_id}", response_model=ProgramOut)
def update_program(program_id: int, payload: ProgramUpdate, db: Session = Depends(get_db)):
    program = db.query(Program).filter(Program.id == program_id).first()
    if not program:
        raise HTTPException(status_code=404, detail="Program not found")
    for field, value in payload.model_dump(exclude_unset=True).items():
        setattr(program, field, value)
    db.commit()
    db.refresh(program)
    return program


@app.get("/programs/{program_id}/brief", response_model=ProgramBriefOut)
def get_program_brief(program_id: int, db: Session = Depends(get_db)):
    program = db.query(Program).filter(Program.id == program_id).first()
    if not program:
        raise HTTPException(status_code=404, detail="Program not found")
    brief = db.query(ProgramBrief).filter(ProgramBrief.program_id == program_id).first()
    if not brief:
        raise HTTPException(status_code=404, detail="Brief not found")
    return brief


@app.put("/programs/{program_id}/brief", response_model=ProgramBriefOut)
def upsert_program_brief(program_id: int, payload: ProgramBriefIn, db: Session = Depends(get_db)):
    program = db.query(Program).filter(Program.id == program_id).first()
    if not program:
        raise HTTPException(status_code=404, detail="Program not found")
    brief = db.query(ProgramBrief).filter(ProgramBrief.program_id == program_id).first()
    if brief is None:
        brief = ProgramBrief(program_id=program_id, **payload.model_dump())
        db.add(brief)
    else:
        for field, value in payload.model_dump().items():
            setattr(brief, field, value)
    db.commit()
    db.refresh(brief)
    return brief


_VALID_SOURCE_TYPES = {"program_input", "exemplar"}


@app.post("/programs/{program_id}/upload", response_model=list[ProgramFileOut])
async def upload_program_files(
    program_id: int,
    file: UploadFile = File(...),
    source_type: str = Query("program_input"),
    db: Session = Depends(get_db),
):
    """
    Upload a single document file (.docx, .doc, .pdf, .txt, .md) for a program.

    Optional query param ``source_type`` classifies the file:
    - ``program_input`` (default) – standard program reference document
    - ``exemplar`` – example/reference artifact used as a model

    Manual test instructions:
      1. Upload a .docx via the Upload tab in the frontend.
      2. Confirm the file appears in the list (filename, size, timestamp).
      3. Check backend storage: ls backend/data/programs/<id>/uploads/<timestamp>/
      4. Confirm extracted text: SELECT extracted_text FROM program_files ORDER BY id DESC LIMIT 1;
    """
    program = db.query(Program).filter(Program.id == program_id).first()
    if not program:
        raise HTTPException(status_code=404, detail="Program not found")

    if source_type not in _VALID_SOURCE_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid source_type '{source_type}'. Accepted: {', '.join(sorted(_VALID_SOURCE_TYPES))}",
        )

    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type '{ext}' not allowed. Accepted: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    # Sanitize: take only the basename to prevent path traversal
    safe_name = Path(file.filename).name

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    upload_dir = DATA_DIR / "programs" / str(program_id) / "uploads" / timestamp
    upload_dir.mkdir(parents=True, exist_ok=True)

    contents = await file.read()
    dest = upload_dir / safe_name
    dest.write_bytes(contents)

    # Extract plain text from the saved file (all supported types)
    extraction_error: str | None = None
    try:
        knowledge_text = _extract_text(dest, ext)
    except Exception as exc:
        knowledge_text = ""
        extraction_error = str(exc)

    pf = ProgramFile(
        program_id=program_id,
        filename=safe_name,
        relative_path=str(dest.relative_to(DATA_DIR)),
        size_bytes=len(contents),
        extracted_text=knowledge_text or None,  # None when empty / failed
        source_type=source_type,
    )
    db.add(pf)
    db.flush()  # need pf.id before inserting dependents

    # Store extracted text in file_text; record any extraction error in meta_json
    ft = FileText(
        file_id=pf.id,
        extracted_text=knowledge_text,
        meta_json=json.dumps({"error": extraction_error}) if extraction_error else None,
    )
    db.add(ft)

    # Chunk the extracted text and store in file_chunks
    for i, chunk in enumerate(_chunk_text(knowledge_text)):
        db.add(FileChunk(
            program_id=program_id,
            file_id=pf.id,
            source_type=source_type,
            chunk_index=i,
            chunk_text=chunk,
            meta_json=json.dumps({"chunk_index": i, "char_start": i * (CHUNK_SIZE - CHUNK_OVERLAP)}),
        ))

    db.commit()
    db.refresh(pf)

    return [pf]


@app.get("/programs/{program_id}/files", response_model=list[ProgramFileOut])
def list_program_files(
    program_id: int,
    source_type: str | None = Query(None),
    db: Session = Depends(get_db),
):
    program = db.query(Program).filter(Program.id == program_id).first()
    if not program:
        raise HTTPException(status_code=404, detail="Program not found")
    if source_type is not None and source_type not in _VALID_SOURCE_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid source_type '{source_type}'. Accepted: {', '.join(sorted(_VALID_SOURCE_TYPES))}",
        )
    q = db.query(ProgramFile).filter(ProgramFile.program_id == program_id)
    if source_type is not None:
        q = q.filter(ProgramFile.source_type == source_type)
    return q.order_by(ProgramFile.id.desc()).all()


@app.get("/programs/{program_id}/knowledge/summary", response_model=KnowledgeSummaryOut)
def knowledge_summary(program_id: int, db: Session = Depends(get_db)):
    program = db.query(Program).filter(Program.id == program_id).first()
    if not program:
        raise HTTPException(status_code=404, detail="Program not found")

    brief = db.query(ProgramBrief).filter(ProgramBrief.program_id == program_id).first()

    answers_rows = db.query(ProgramAnswer).filter(ProgramAnswer.program_id == program_id).all()
    wizard_answers: dict[str, str | None] = {r.question_id: r.answer_text for r in answers_rows}

    modules = (
        db.query(Module)
        .filter(Module.program_id == program_id)
        .order_by(Module.id.asc())
        .all()
    )

    files = (
        db.query(ProgramFile)
        .filter(ProgramFile.program_id == program_id)
        .order_by(ProgramFile.id.asc())
        .all()
    )

    program_file_summaries = []
    exemplar_file_summaries = []
    text_file_count = 0
    for f in files:
        has_text = bool(f.extracted_text)
        if has_text:
            text_file_count += 1
        summary = {
            "id": f.id,
            "filename": f.filename,
            "uploaded_at": f.uploaded_at,
            "text_available": has_text,
            "text_preview_500chars": f.extracted_text[:500] if has_text else None,
            "source_type": f.source_type,
        }
        if f.source_type == "exemplar":
            exemplar_file_summaries.append(summary)
        else:
            program_file_summaries.append(summary)

    return KnowledgeSummaryOut(
        program=program,
        brief=brief,
        wizard_answers=wizard_answers,
        modules=modules,
        program_files=program_file_summaries,
        exemplar_files=exemplar_file_summaries,
        stats={"file_count": len(files), "text_file_count": text_file_count},
    )


@app.get("/programs/{program_id}/knowledge/search", response_model=list[KnowledgeSearchResult])
def search_knowledge(
    program_id: int,
    q: str = Query(..., min_length=1),
    top_k: int = Query(5, ge=1, le=50),
    db: Session = Depends(get_db),
):
    if not db.query(Program).filter(Program.id == program_id).first():
        raise HTTPException(status_code=404, detail="Program not found")

    q_stripped = q.strip()
    if not q_stripped:
        return []

    rows = (
        db.query(ProgramFile, FileText)
        .join(FileText, FileText.file_id == ProgramFile.id)
        .filter(ProgramFile.program_id == program_id)
        .all()
    )

    results: list[KnowledgeSearchResult] = []
    q_lower = q_stripped.lower()
    for pf, ft in rows:
        text = ft.extracted_text or ""
        if not text:
            continue
        text_lower = text.lower()
        match_count = text_lower.count(q_lower)
        if match_count == 0:
            continue

        idx = text_lower.find(q_lower)
        start = max(0, idx - 100)
        end = min(len(text), idx + len(q_stripped) + 350)
        snippet = text[start:end].strip()
        if start > 0:
            snippet = "\u2026" + snippet
        if end < len(text):
            snippet = snippet + "\u2026"

        results.append(KnowledgeSearchResult(
            file_id=pf.id,
            filename=pf.filename,
            match_count=match_count,
            snippet=snippet,
        ))

    results.sort(key=lambda r: (-r.match_count, r.filename))
    return results[:top_k]


@app.get("/programs/{program_id}/files/{file_id}/text", response_model=FileTextOut)
def get_file_text(program_id: int, file_id: int, db: Session = Depends(get_db)):
    pf = (
        db.query(ProgramFile)
        .filter(ProgramFile.program_id == program_id, ProgramFile.id == file_id)
        .first()
    )
    if not pf:
        raise HTTPException(status_code=404, detail="File not found")

    ft = db.query(FileText).filter(FileText.file_id == file_id).first()
    if not ft:
        raise HTTPException(status_code=404, detail="No extracted text record for this file")
    return ft


@app.get("/programs/{program_id}/files/{file_id}/chunks", response_model=list[FileChunkOut])
def get_file_chunks(program_id: int, file_id: int, db: Session = Depends(get_db)):
    pf = (
        db.query(ProgramFile)
        .filter(ProgramFile.program_id == program_id, ProgramFile.id == file_id)
        .first()
    )
    if not pf:
        raise HTTPException(status_code=404, detail="File not found")

    return (
        db.query(FileChunk)
        .filter(FileChunk.file_id == file_id)
        .order_by(FileChunk.chunk_index.asc())
        .limit(20)
        .all()
    )


@app.post("/programs/{program_id}/files/reextract", response_model=ReextractOut)
def reextract_program_files(program_id: int, db: Session = Depends(get_db)):
    """Re-run text extraction and chunking for every file in the program.

    Uses the same extraction logic as upload (including the XML fallback for
    low-quality DOCX results).  Overwrites extracted_text on program_files and
    file_text, then rebuilds file_chunks from scratch.
    """
    if not db.query(Program).filter(Program.id == program_id).first():
        raise HTTPException(status_code=404, detail="Program not found")

    files = (
        db.query(ProgramFile)
        .filter(ProgramFile.program_id == program_id)
        .all()
    )

    results: list[ReextractFileResult] = []
    for pf in files:
        abs_path = DATA_DIR / pf.relative_path
        ext = Path(pf.filename).suffix.lower()

        extraction_error: str | None = None
        try:
            knowledge_text = _extract_text(abs_path, ext)
        except Exception as exc:
            knowledge_text = ""
            extraction_error = str(exc)

        # Update program_files.extracted_text
        pf.extracted_text = knowledge_text or None

        # Update or create file_text record
        ft = db.query(FileText).filter(FileText.file_id == pf.id).first()
        meta: dict = {"reextracted": True}
        if extraction_error:
            meta["error"] = extraction_error
        if ft:
            ft.extracted_text = knowledge_text
            ft.meta_json = json.dumps(meta)
        else:
            db.add(FileText(
                file_id=pf.id,
                extracted_text=knowledge_text,
                meta_json=json.dumps(meta),
            ))

        # Rebuild chunks
        db.query(FileChunk).filter(FileChunk.file_id == pf.id).delete()
        for i, chunk in enumerate(_chunk_text(knowledge_text)):
            db.add(FileChunk(
                program_id=program_id,
                file_id=pf.id,
                source_type=pf.source_type,
                chunk_index=i,
                chunk_text=chunk,
                meta_json=json.dumps({"chunk_index": i, "char_start": i * (CHUNK_SIZE - CHUNK_OVERLAP)}),
            ))

        results.append(ReextractFileResult(
            file_id=pf.id,
            filename=pf.filename,
            chars=len(knowledge_text),
            error=extraction_error,
        ))

    db.commit()
    return ReextractOut(reextracted=len(files), files=results)


def _build_wizard_response(program_id: int, db: Session) -> WizardOut:
    questions = _load_questions()
    rows = db.query(ProgramAnswer).filter(ProgramAnswer.program_id == program_id).all()
    answers_map: dict[str, object] = {r.question_id: r.answer_text for r in rows}

    # Load structured modules for modules_builder question
    modules_db = db.query(Module).filter(Module.program_id == program_id).order_by(Module.id.asc()).all()
    if modules_db:
        answers_map["modules"] = [
            {
                "name": m.name,
                "description": m.description or "",
                "rationale": m.rationale or "",
                "interfaces": m.key_interfaces or "",
            }
            for m in modules_db
        ]

    questions_out = []
    answered_count = 0
    for q in questions:
        if q["type"] == "modules_builder":
            missing = len(modules_db) == 0
        else:
            answer = answers_map.get(q["id"])
            missing = answer is None or str(answer).strip() == ""
        if not missing:
            answered_count += 1
        questions_out.append({**q, "missing": missing})

    total = len(questions)
    return WizardOut(
        questions=questions_out,
        answers=answers_map,
        answered_count=answered_count,
        total_count=total,
        percent_complete=round(answered_count / total * 100, 1) if total > 0 else 0.0,
    )


@app.get("/programs/{program_id}/wizard", response_model=WizardOut)
def get_wizard(program_id: int, db: Session = Depends(get_db)):
    if not db.query(Program).filter(Program.id == program_id).first():
        raise HTTPException(status_code=404, detail="Program not found")
    return _build_wizard_response(program_id, db)


@app.put("/programs/{program_id}/wizard", response_model=WizardOut)
def put_wizard(program_id: int, payload: WizardAnswersIn, db: Session = Depends(get_db)):
    if not db.query(Program).filter(Program.id == program_id).first():
        raise HTTPException(status_code=404, detail="Program not found")

    # Handle structured modules submitted from the wizard modules builder
    modules_list = payload.answers.get("modules")
    if isinstance(modules_list, list) and modules_list:
        db.query(Module).filter(Module.program_id == program_id).delete()
        for item in modules_list:
            if not isinstance(item, dict) or not str(item.get("name", "")).strip():
                continue
            db.add(Module(
                program_id=program_id,
                name=str(item.get("name", ""))[:200],
                description=item.get("description") or None,
                rationale=item.get("rationale") or None,
                key_interfaces=item.get("interfaces") or None,
            ))

    valid_ids = {q["id"] for q in _load_questions()}
    for question_id, answer_text in payload.answers.items():
        if question_id == "modules" or question_id not in valid_ids:
            continue
        if not isinstance(answer_text, str):
            continue
        existing = (
            db.query(ProgramAnswer)
            .filter(ProgramAnswer.program_id == program_id, ProgramAnswer.question_id == question_id)
            .first()
        )
        if existing:
            existing.answer_text = answer_text
        else:
            db.add(ProgramAnswer(program_id=program_id, question_id=question_id, answer_text=answer_text))

    db.commit()
    return _build_wizard_response(program_id, db)


# ---------------------------------------------------------------------------
# Modules helpers
# ---------------------------------------------------------------------------

def _parse_module_names(h_answer: str | None) -> list[str]:
    """Extract candidate module names from wizard answer h_candidate_modules."""
    if not h_answer or not h_answer.strip():
        return []
    names: list[str] = []
    for line in h_answer.splitlines():
        line = line.strip()
        line = re.sub(r"^[-•*·◦▪▸>]+\s*", "", line)   # strip bullet markers
        line = re.sub(r"^\d+[.)]\s*", "", line)         # strip numbered lists
        line = line.strip()
        if not line:
            continue
        if len(line) < 100 and "," in line:
            names.extend(p.strip() for p in line.split(",") if p.strip())
        else:
            names.append(line)
    return names[:10]


def parse_modules_from_text(text: str) -> list[ModuleIn]:
    """Parse freeform module text into a list of ModuleIn objects (backward compat).

    Splits on blank lines OR lines starting with a bullet/number marker.
    If a chunk starts with "Name:" the value after the colon is the name;
    otherwise the first line is the name and subsequent lines are the description.
    """
    if not text or not text.strip():
        return []

    # Try splitting on blank lines first
    chunks = re.split(r"\n\s*\n", text.strip())

    # If only one chunk, try splitting on bullet/numbered list lines
    if len(chunks) == 1:
        bullet_lines = [
            l for l in text.strip().splitlines()
            if re.match(r"^[-•*·◦▪▸>]+\s+\S", l.strip()) or re.match(r"^\d+[.)]\s+\S", l.strip())
        ]
        if len(bullet_lines) > 1:
            chunks = bullet_lines

    modules: list[ModuleIn] = []
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        # Strip leading bullet/number
        chunk = re.sub(r"^[-•*·◦▪▸>]+\s*", "", chunk)
        chunk = re.sub(r"^\d+[.)]\s*", "", chunk).strip()
        if not chunk:
            continue

        if chunk.lower().startswith("name:"):
            lines = chunk.splitlines()
            name = lines[0][5:].strip()
            desc = " ".join(l.strip() for l in lines[1:] if l.strip())
        else:
            lines = chunk.splitlines()
            name = lines[0].strip()
            desc = " ".join(l.strip() for l in lines[1:] if l.strip())

        if name:
            modules.append(ModuleIn(name=name[:200], description=desc or None))

    return modules[:10]


def _default_module_names(brief: ProgramBrief | None) -> list[str]:
    if brief and brief.software_large_part:
        return ["Core Software", "Mission Software", "Infrastructure & Middleware", "External Interfaces"]
    if brief and brief.attritable:
        return ["Guidance System", "Payload", "Propulsion", "Navigation & Control"]
    return ["Mission System", "Communications", "Navigation & Control", "Software & Computing", "Power Management"]


def _find_in_text(name: str, text: str) -> bool:
    return name.lower() in text.lower()


def _extract_line_mentioning(name: str, text: str) -> str:
    for line in text.splitlines():
        if name.lower() in line.lower():
            return line.strip()
    return ""


# ---------------------------------------------------------------------------
# Module CRUD endpoints
# ---------------------------------------------------------------------------

@app.get("/programs/{program_id}/modules", response_model=list[ModuleOut])
def list_modules(program_id: int, db: Session = Depends(get_db)):
    if not db.query(Program).filter(Program.id == program_id).first():
        raise HTTPException(status_code=404, detail="Program not found")
    return (
        db.query(Module)
        .filter(Module.program_id == program_id)
        .order_by(Module.id.asc())
        .all()
    )


@app.post("/programs/{program_id}/modules", response_model=ModuleOut, status_code=201)
def create_module(program_id: int, payload: ModuleIn, db: Session = Depends(get_db)):
    if not db.query(Program).filter(Program.id == program_id).first():
        raise HTTPException(status_code=404, detail="Program not found")
    mod = Module(program_id=program_id, **payload.model_dump())
    db.add(mod)
    db.commit()
    db.refresh(mod)
    return mod


@app.post("/programs/{program_id}/modules/seed", response_model=list[ModuleOut])
def seed_modules(program_id: int, db: Session = Depends(get_db)):
    if not db.query(Program).filter(Program.id == program_id).first():
        raise HTTPException(status_code=404, detail="Program not found")

    # If structured modules already exist (saved from wizard builder), return them as-is
    existing = db.query(Module).filter(Module.program_id == program_id).order_by(Module.id.asc()).all()
    if existing:
        return existing

    brief = db.query(ProgramBrief).filter(ProgramBrief.program_id == program_id).first()
    answers_rows = db.query(ProgramAnswer).filter(ProgramAnswer.program_id == program_id).all()
    answers: dict[str, str | None] = {r.question_id: r.answer_text for r in answers_rows}

    h_answer = answers.get("h_candidate_modules") or ""
    f_answer = answers.get("f_tech_challenges_and_risk_areas") or ""
    i_answer = answers.get("i_known_standards_architectures_mapping") or ""
    j_answer = answers.get("j_obsolescence_candidates") or ""
    k_answer = answers.get("k_commercial_solutions_by_module") or ""
    n_answer = answers.get("n_software_standards_architectures") or ""

    # Prefer structured parse from text; fall back to name-only parse then defaults
    parsed_modules = parse_modules_from_text(h_answer)
    if parsed_modules:
        module_names = [m.name for m in parsed_modules]
    else:
        module_names = _parse_module_names(h_answer) or _default_module_names(brief)

    prog_desc = (brief.program_description if brief else "") or ""
    rationale_prefix = f"Identified for: {prog_desc[:120]}" if prog_desc else "Seeded from wizard answers"

    # Clear existing modules and recreate deterministically
    db.query(Module).filter(Module.program_id == program_id).delete()

    new_modules: list[Module] = []
    for i, name in enumerate(module_names):
        tech_risk = _find_in_text(name, f_answer) if f_answer else False
        obsolescence_risk = _find_in_text(name, j_answer) if j_answer else False
        cots_candidate = _find_in_text(name, k_answer) if k_answer else False

        key_interfaces = _extract_line_mentioning(name, i_answer) if i_answer else None
        standards = n_answer.strip()[:300] or None

        # Carry description from parsed module if available
        description = parsed_modules[i].description if parsed_modules and i < len(parsed_modules) else None

        mod = Module(
            program_id=program_id,
            name=name[:200],
            description=description,
            rationale=rationale_prefix,
            key_interfaces=key_interfaces or None,
            standards=standards,
            tech_risk=tech_risk,
            obsolescence_risk=obsolescence_risk,
            cots_candidate=cots_candidate,
        )
        db.add(mod)
        new_modules.append(mod)

    db.commit()
    for mod in new_modules:
        db.refresh(mod)
    return new_modules


@app.get("/programs/{program_id}/modules/{module_id}", response_model=ModuleOut)
def get_module(program_id: int, module_id: int, db: Session = Depends(get_db)):
    mod = (
        db.query(Module)
        .filter(Module.program_id == program_id, Module.id == module_id)
        .first()
    )
    if not mod:
        raise HTTPException(status_code=404, detail="Module not found")
    return mod


@app.put("/programs/{program_id}/modules/{module_id}", response_model=ModuleOut)
def update_module(program_id: int, module_id: int, payload: ModuleIn, db: Session = Depends(get_db)):
    mod = (
        db.query(Module)
        .filter(Module.program_id == program_id, Module.id == module_id)
        .first()
    )
    if not mod:
        raise HTTPException(status_code=404, detail="Module not found")
    for field, value in payload.model_dump().items():
        setattr(mod, field, value)
    db.commit()
    db.refresh(mod)
    return mod


@app.delete("/programs/{program_id}/modules/{module_id}", status_code=204)
def delete_module(program_id: int, module_id: int, db: Session = Depends(get_db)):
    mod = (
        db.query(Module)
        .filter(Module.program_id == program_id, Module.id == module_id)
        .first()
    )
    if not mod:
        raise HTTPException(status_code=404, detail="Module not found")
    db.delete(mod)
    db.commit()


# ---------------------------------------------------------------------------
# Smart RFI – OpenAI JSON schema description (used in system prompt)
# ---------------------------------------------------------------------------

_SMART_RFI_JSON_SCHEMA = """\
Return ONLY a JSON object with exactly these keys (no extra keys, no markdown fences):
{
  "overview": "<string: 2-4 sentence overview of the RFI document and its purpose>",
  "rfi_purpose": "<string: statutory/policy purpose of this RFI, citing relevant law or policy>",
  "program_context": "<string: detailed program context and background, 2-5 paragraphs>",
  "mosa_requirements": ["<string>", ...],
  "questions_to_industry": ["<string>", ...],
  "requested_deliverables": ["<string>", ...],
  "submission_instructions": ["<string>", ...],
  "module_table_rows": [
    {
      "module_name": "<string>",
      "rationale": "<string>",
      "key_interfaces": "<string>",
      "standards": "<string>",
      "tech_risk": "<'Yes'|'No'|brief descriptor>",
      "obsolescence_risk": "<'Yes'|'No'|brief descriptor>",
      "cots_candidate": "<'Yes'|'No'|brief descriptor>"
    }
  ],
  "sources_used": [
    {
      "file_id": <integer>,
      "filename": "<string>",
      "excerpt": "<string: 1-3 sentence excerpt from the file>"
    }
  ],
  "citations": {
    "<section_key e.g. overview>": [<source_number_integer>, ...]
  }
}
\
"""


# ---------------------------------------------------------------------------
# Document generation  (see docx_builder.py)
# ---------------------------------------------------------------------------


@app.post("/programs/{program_id}/documents/generate", response_model=list[DocumentOut])
def generate_documents(
    program_id: int,
    payload: GenerateDocRequest,
    db: Session = Depends(get_db),
):
    program = db.query(Program).filter(Program.id == program_id).first()
    if not program:
        raise HTTPException(status_code=404, detail="Program not found")

    if not payload.doc_types:
        raise HTTPException(status_code=400, detail="doc_types must not be empty")

    supported = {"RFI", "ACQ_STRATEGY", "SEP", "MOSA_CONFORMANCE_PLAN"}
    unknown = set(payload.doc_types) - supported
    if unknown:
        raise HTTPException(status_code=400, detail=f"Unsupported doc_types: {', '.join(sorted(unknown))}")

    brief = db.query(ProgramBrief).filter(ProgramBrief.program_id == program_id).first()
    answers_rows = db.query(ProgramAnswer).filter(ProgramAnswer.program_id == program_id).all()
    answers: dict[str, str | None] = {r.question_id: r.answer_text for r in answers_rows}
    files = (
        db.query(ProgramFile)
        .filter(ProgramFile.program_id == program_id)
        .order_by(ProgramFile.id.asc())
        .all()
    )
    modules = (
        db.query(Module)
        .filter(Module.program_id == program_id)
        .order_by(Module.id.asc())
        .all()
    )

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    generated_date = datetime.now(timezone.utc).strftime("%B %d, %Y")
    out_dir = DATA_DIR / "programs" / str(program_id) / "docs" / timestamp
    out_dir.mkdir(parents=True, exist_ok=True)

    new_docs: list[ProgramDocument] = []

    for doc_type in payload.doc_types:
        if doc_type == "RFI":
            docx = build_rfi(program, brief, answers, files, modules, generated_date)
            file_name = "RFI.docx"
        elif doc_type == "ACQ_STRATEGY":
            docx = build_acq_strategy(program, brief, answers, files, modules, generated_date)
            file_name = "AcqStrategy.docx"
        elif doc_type == "SEP":
            docx = build_sep(program, brief, answers, files, modules, generated_date)
            file_name = "SEP.docx"
        elif doc_type == "MOSA_CONFORMANCE_PLAN":
            docx = build_mosa_conformance_plan(program, brief, answers, files, modules, generated_date)
            file_name = "MOSAConformancePlan.docx"
        else:
            continue  # guarded by supported check above

        abs_path = out_dir / file_name
        docx.save(str(abs_path))
        rel_path = str(abs_path.relative_to(DATA_DIR))

        doc_row = ProgramDocument(
            program_id=program_id,
            doc_type=doc_type,
            file_path=rel_path,
        )
        db.add(doc_row)
        new_docs.append(doc_row)

    db.commit()
    for d in new_docs:
        db.refresh(d)

    return new_docs


@app.get("/programs/{program_id}/documents", response_model=list[DocumentOut])
def list_documents(program_id: int, db: Session = Depends(get_db)):
    if not db.query(Program).filter(Program.id == program_id).first():
        raise HTTPException(status_code=404, detail="Program not found")
    return (
        db.query(ProgramDocument)
        .filter(ProgramDocument.program_id == program_id)
        .order_by(ProgramDocument.id.desc())
        .all()
    )


# ---------------------------------------------------------------------------
# Shared vector-RAG helpers for all Smart doc generators
# ---------------------------------------------------------------------------

def _build_program_query(
    program_name: str,
    brief_description: str,
    answers: dict,
    doc_type_hint: str,
) -> str:
    """Build a concise query string for vector similarity retrieval.

    Combines program name, brief description, key wizard answers, and the
    document-type hint into a single string that captures the program's
    essential characteristics for embedding-based retrieval.
    """
    parts: list[str] = [program_name]
    if brief_description and brief_description.strip():
        desc = brief_description.strip()
        parts.append(desc[:300] if len(desc) > 300 else desc)
    for q_id in (
        "a_program_description",
        "f_tech_challenges_and_risk_areas",
        "g_mosa_scenarios",
        "h_candidate_modules",
    ):
        val = answers.get(q_id)
        if val and str(val).strip():
            snippet = str(val).strip()
            parts.append(snippet[:150] if len(snippet) > 150 else snippet)
    parts.append(doc_type_hint)
    return ". ".join(parts)


def _retrieve_rag_context(
    program_id: int,
    program_query: str,
    db,
) -> tuple[str, str]:
    """Retrieve vector RAG chunks split by source_type for Smart doc prompts.

    Calls retrieve_chunks_vector twice:
      - top 12 chunks with source_type='program_input'  → Program Evidence
      - top 6 chunks  with source_type='exemplar'       → Exemplar Patterns

    Degrades gracefully (returns placeholder strings) if the vector index
    has not been built yet for the program.

    Returns:
        (program_evidence_text, exemplar_patterns_text)
    """
    from llm.retrieval import retrieve_chunks_vector

    def _fmt(chunks: list[dict]) -> str:
        parts = []
        for c in chunks:
            parts.append(
                f"[{c['filename']} | chunk {c['chunk_index']}]\n{c['chunk_text']}"
            )
        return "\n\n---\n\n".join(parts)

    try:
        prog_chunks = retrieve_chunks_vector(
            program_id=program_id,
            query=program_query,
            db=db,
            source_type="program_input",
            top_k=12,
        )
        exmp_chunks = retrieve_chunks_vector(
            program_id=program_id,
            query=program_query,
            db=db,
            source_type="exemplar",
            top_k=6,
        )
    except Exception:
        # Vector index may not exist yet; degrade gracefully
        return (
            "No program input documents have been indexed for vector search yet. "
            "Run POST /programs/{id}/knowledge/index to enable richer evidence.",
            "No exemplar documents have been indexed for vector search yet.",
        )

    prog_text = (
        _fmt(prog_chunks)
        if prog_chunks
        else "No program input chunks found in the vector index."
    )
    exmp_text = (
        _fmt(exmp_chunks)
        if exmp_chunks
        else "No exemplar chunks found in the vector index."
    )
    return prog_text, exmp_text


# ---------------------------------------------------------------------------
# Grounded-source RAG helpers (used by all smart doc generators)
# ---------------------------------------------------------------------------

_GROUNDING_INSTRUCTION = (
    "GROUNDING RULES (highest priority):\n"
    "- Only state facts that appear in the SOURCES block or the PROGRAM FACT-PACK.\n"
    "- If a required piece of information is absent from those sources, output "
    "\"Not provided\" rather than inventing a value.\n"
    "- Do not fabricate citations, file names, acronyms, or statistics not present "
    "in the sources.\n"
    "- In the \"citations\" field, map each top-level JSON section key (e.g. "
    "\"executive_summary\") to a list of integer source numbers from the SOURCES "
    "block (e.g. [1, 3, 5]) that support that section. Use [] if no source applies.\n\n"
)

_MAX_CHUNK_CHARS = 1200  # per-chunk character cap
_MAX_CONTEXT_CHARS = int(os.environ.get("MAX_CONTEXT_CHARS", "12000"))


def _build_sources_block(chunks: list[dict]) -> str:
    """Format a numbered SOURCES block from retrieve_chunks() output.

    Each entry: [#N filename (file_id:x chunk:y type:z)] <truncated text>

    Enforces two budgets:
    - Per-chunk: _MAX_CHUNK_CHARS characters (hard truncation + ellipsis).
    - Total context: _MAX_CONTEXT_CHARS (MAX_CONTEXT_CHARS env var, default 12 000).
      Chunks are added in score order until the budget is exhausted; remaining
      chunks are dropped and a note is appended.

    Returns the full block string (empty string if no chunks).
    """
    if not chunks:
        return ""

    lines: list[str] = []
    total_chars = 0
    included = 0

    for i, c in enumerate(chunks, 1):
        text = c["chunk_text"]
        if len(text) > _MAX_CHUNK_CHARS:
            text = text[:_MAX_CHUNK_CHARS] + "..."
        header = (
            f"[#{i} {c['filename']} "
            f"(file_id:{c['file_id']} chunk:{c['chunk_index']} type:{c['source_type']})]"
        )
        entry = f"{header}\n{text}"
        entry_chars = len(entry)

        if total_chars + entry_chars > _MAX_CONTEXT_CHARS:
            break  # budget exhausted

        lines.append(entry)
        total_chars += entry_chars
        included += 1

    skipped = len(chunks) - included
    block = "\n\n".join(lines)
    if skipped:
        block += f"\n\n[{skipped} additional source(s) omitted — context budget reached]"

    print(
        f"[rag] sources_block  included={included}/{len(chunks)} "
        f"total_chars={total_chars} budget={_MAX_CONTEXT_CHARS}"
    )
    return block


def _retrieve_grounded_sources(
    program_id: int,
    query: str,
    db,
    k: int = 12,
) -> tuple[str, list]:
    """Retrieve pgvector chunks and return a numbered SOURCES block.

    Calls rag.retrieve_chunks; degrades gracefully when no vector index exists.
    Returns (sources_section_text, raw_chunks_list).
    The sources_section_text is ready to embed in the user message under a
    '## SOURCES' heading.
    """
    from rag import retrieve_chunks as _rag_retrieve

    try:
        chunks = _rag_retrieve(db, program_id, query, k)
    except ValueError:
        # No vector index — return a clear placeholder so the LLM knows
        return (
            "No vector index found for this program. "
            "Run POST /programs/{id}/knowledge/index to enable grounded retrieval.",
            [],
        )

    block = _build_sources_block(chunks)
    if not block:
        return "No matching chunks found in the vector index.", []
    return block, chunks


@app.post("/programs/{program_id}/docs/rfi/smart")
def generate_smart_rfi(program_id: int, db: Session = Depends(get_db)):
    """Generate an AI-assisted RFI using structured LLM output (RfiPlan schema).

    Requires OPENAI_API_KEY to be set in the environment.
    Returns the generated DOCX directly as a file download.
    """
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=503,
            detail=(
                "Smart RFI generation requires an LLM API key. "
                "Please set the OPENAI_API_KEY environment variable and restart the server."
            ),
        )

    program = db.query(Program).filter(Program.id == program_id).first()
    if not program:
        raise HTTPException(status_code=404, detail="Program not found")

    brief = db.query(ProgramBrief).filter(ProgramBrief.program_id == program_id).first()
    answers_rows = db.query(ProgramAnswer).filter(ProgramAnswer.program_id == program_id).all()
    answers: dict[str, str | None] = {r.question_id: r.answer_text for r in answers_rows}
    files = (
        db.query(ProgramFile)
        .filter(ProgramFile.program_id == program_id)
        .order_by(ProgramFile.id.asc())
        .all()
    )
    modules = (
        db.query(Module)
        .filter(Module.program_id == program_id)
        .order_by(Module.id.asc())
        .all()
    )

    # ---- Build fact-pack ----
    fact_lines: list[str] = [f"Program Name: {program.name}"]
    if brief:
        if brief.program_description:
            fact_lines.append(f"Program Description: {brief.program_description}")
        if brief.dev_cost_estimate is not None:
            fact_lines.append(f"Estimated Development Cost: ${brief.dev_cost_estimate:,.0f}M")
        if brief.production_unit_cost is not None:
            fact_lines.append(f"Production Unit Cost: ${brief.production_unit_cost:,.0f}M")
        flags = []
        if brief.attritable:
            flags.append("Attritable")
        if brief.sustainment_tail:
            flags.append("Sustainment Tail")
        if brief.software_large_part:
            flags.append("Software-Intensive")
        if brief.mission_critical:
            flags.append("Mission-Critical")
        if brief.safety_critical:
            flags.append("Safety-Critical")
        if flags:
            fact_lines.append(f"Program Flags: {', '.join(flags)}")

    wizard_fields = {
        "a_program_description": "Program Description (Wizard)",
        "e_similar_previous_programs": "Heritage / Analogous Programs",
        "f_tech_challenges_and_risk_areas": "Technical Challenges and Risk Areas",
        "g_mosa_scenarios": "MOSA Scenarios",
        "h_candidate_modules": "Candidate Modules",
        "i_known_standards_architectures_mapping": "Known Standards / Architectures by Module",
        "j_obsolescence_candidates": "Obsolescence Candidates",
        "k_commercial_solutions_by_module": "Commercial Solutions by Module",
        "n_software_standards_architectures": "Software Standards / Architectures",
        "o_mosa_repo_searched": "MOSA Repository Search Performed",
    }
    for q_id, label in wizard_fields.items():
        val = answers.get(q_id)
        if val and val.strip():
            fact_lines.append(f"\n{label}:\n{val.strip()}")

    if modules:
        fact_lines.append("\nDefined Modules:")
        for m in modules:
            row = f"  - [Module] {m.name}"
            if m.rationale:
                row += f": {m.rationale}"
            if m.key_interfaces:
                row += f" | Interfaces: {m.key_interfaces}"
            if m.standards:
                row += f" | Standards: {m.standards}"
            row += f" | Tech Risk: {'Yes' if m.tech_risk else 'No'}"
            row += f" | Obsolescence Risk: {'Yes' if m.obsolescence_risk else 'No'}"
            row += f" | COTS Candidate: {'Yes' if m.cots_candidate else 'No'}"
            fact_lines.append(row)

    if files:
        fact_lines.append("\nUploaded Reference Files (file_id | filename):")
        for f in files:
            fact_lines.append(f"  - [{f.id}] {f.filename}")

    fact_pack = "\n".join(fact_lines)

    # ---- Grounded source retrieval ----
    _rfi_brief_desc = brief.program_description if brief else ""
    _rfi_query = _build_program_query(
        program.name, _rfi_brief_desc, answers,
        "RFI request for information MOSA modular open systems industry questions",
    )
    _rfi_sources, _rfi_chunks = _retrieve_grounded_sources(program_id, _rfi_query, db)

    # ---- Build LLM prompt ----
    _rfi_tmpl = get_template("rfi")
    system_prompt = (
        "You are a defense acquisition expert specializing in Modular Open Systems Approach (MOSA) "
        "and Request for Information (RFI) documents.\n\n"
        "Generate a comprehensive, program-specific RFI plan.\n\n"
        "Requirements:\n"
        "1. All content must be specific to this program — not generic boilerplate.\n"
        "2. module_table_rows must contain one row per module listed in 'Defined Modules' below "
        "(if no modules are listed, propose sensible defaults based on the program context).\n"
        "3. sources_used must only reference file_ids that appear in the SOURCES block; "
        "if no sources are available, set sources_used to an empty list.\n"
        "4. MOSA requirements and questions to industry must be tailored to the program's specific characteristics.\n"
        "5. tech_risk, obsolescence_risk, and cots_candidate in module rows should be 'Yes', 'No', or a brief descriptor.\n\n"
        + build_template_contract("rfi")
        + _GROUNDING_INSTRUCTION
        + _SMART_RFI_JSON_SCHEMA
    )
    user_message = (
        "## PROGRAM FACT-PACK\n"
        f"{fact_pack}\n\n"
        "## SOURCES\n"
        "Use only these numbered sources for facts. Cite them in the citations field "
        "by their source number.\n\n"
        f"{_rfi_sources}"
    )

    # ---- Call OpenAI API ----
    try:
        import openai as _openai  # lazy import; only needed for this endpoint
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="The 'openai' Python package is not installed. Run: pip install openai",
        )

    model = os.environ.get("OPENAI_MODEL", "gpt-5-mini")
    try:
        client = _openai.OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model=model,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
        )
    except _openai.AuthenticationError:
        raise HTTPException(
            status_code=503,
            detail="Invalid OPENAI_API_KEY. Please verify your API key and try again.",
        )
    except Exception as exc:
        raise HTTPException(status_code=503, detail=f"LLM service error: {exc}")

    # ---- Parse JSON response ----
    raw_content = response.choices[0].message.content or ""
    try:
        raw_data = json.loads(raw_content)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=502, detail=f"LLM did not return valid JSON: {exc}")

    # ---- Enforce template conformance, fill "Not provided" defaults ----
    try:
        normalize_llm_output("rfi", raw_data)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

    # ---- Validate with Pydantic ----
    try:
        rfi_plan = RfiPlan.model_validate(raw_data)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"LLM produced an invalid RFI plan structure: {exc}")

    # ---- Render DOCX ----
    generated_date = datetime.now(timezone.utc).strftime("%B %d, %Y")
    docx = build_smart_rfi(program, generated_date, rfi_plan)

    # ---- Persist to disk + DB ----
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    out_dir = DATA_DIR / "programs" / str(program_id) / "docs" / timestamp
    out_dir.mkdir(parents=True, exist_ok=True)
    file_name = "RFI_Smart.docx"
    abs_path = out_dir / file_name
    docx.save(str(abs_path))

    # Traceability: save plan JSON + source file IDs + citations
    traceability = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "model": model,
        "source_file_ids": [c["file_id"] for c in _rfi_chunks],
        "citations": raw_data.get("citations", {}),
        "plan": raw_data,
    }
    (out_dir / "rfi_plan_trace.json").write_text(
        json.dumps(traceability, indent=2, default=str)
    )

    rel_path = str(abs_path.relative_to(DATA_DIR))
    doc_row = ProgramDocument(program_id=program_id, doc_type="RFI_SMART", file_path=rel_path)
    db.add(doc_row)
    db.commit()

    return FileResponse(
        path=str(abs_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=file_name,
    )


# ---------------------------------------------------------------------------
# Smart Acquisition Strategy – OpenAI JSON schema description
# ---------------------------------------------------------------------------

_SMART_ACQ_STRATEGY_JSON_SCHEMA = """\
Return ONLY a JSON object with exactly these keys (no extra keys, no markdown fences):
{
  "title_block": {
    "program_name": "<string: full program name>",
    "date": "<string: e.g. 'February 2026'>",
    "organization": "<string: e.g. 'Program Executive Office Aviation'>"
  },
  "executive_summary": "<string: 2-4 paragraph executive summary of the acquisition strategy>",
  "acquisition_approach": "<string: 2-3 paragraph narrative describing overall acquisition approach and rationale>",
  "schedule_milestones": [
    {
      "name": "<string: milestone name, e.g. 'Milestone A', 'RFP Release', 'Contract Award'>",
      "date": "<string: e.g. 'Q2 FY26'>",
      "description": "<string: 1-2 sentence description of the milestone>"
    }
  ],
  "cost_estimates": {
    "development": "<string: formatted estimate, e.g. '$15.5M'>",
    "production_unit": "<string: formatted estimate per unit, e.g. '$2.3M per unit'>",
    "sustainment_annual": "<string: formatted annual estimate, e.g. '$1.2M per year'>"
  },
  "risk_register": [
    {
      "risk_id": "<string: e.g. 'R-01'>",
      "description": "<string: clear risk description>",
      "probability": "<'Low'|'Medium'|'High'>",
      "impact": "<'Low'|'Medium'|'High'>",
      "mitigation": "<string: specific mitigation strategy>",
      "owner": "<string: responsible party or role>"
    }
  ],
  "standards_references": [
    {
      "name": "<string: standard or regulation name, e.g. 'MIL-STD-882E'>",
      "description": "<string: one-line description of relevance>"
    }
  ],
  "mosa_approach": "<string: 2-3 paragraph narrative on Modular Open Systems Approach implementation>",
  "mosa_bullets": ["<string: specific MOSA commitment or requirement>", ...],
  "data_rights_approach": "<string: 1-2 paragraph data rights and technical data strategy>",
  "test_verification_approach": "<string: 1-2 paragraph test, evaluation, and verification strategy>",
  "contracting_strategy": "<string: 1-2 paragraph contracting vehicle and competition strategy>",
  "module_table_rows": [
    {
      "module_name": "<string>",
      "rationale": "<string>",
      "key_interfaces": "<string>",
      "standards": "<string>",
      "tech_risk": "<'Yes'|'No'|brief descriptor>",
      "obsolescence_risk": "<'Yes'|'No'|brief descriptor>",
      "cots_candidate": "<'Yes'|'No'|brief descriptor>"
    }
  ],
  "sources_used": [
    {
      "file_id": <integer>,
      "filename": "<string>",
      "excerpt": "<string: 1-3 sentence excerpt from the file>"
    }
  ],
  "citations": {
    "<section_key e.g. executive_summary>": [<source_number_integer>, ...]
  }
}
"""

_ACQ_STRATEGY_SYSTEM_PROMPT_PREFIX = (
    "You are a defense acquisition expert specializing in Acquisition Strategy documents "
    "per DoDI 5000.02, 10 U.S.C. § 4401, and related acquisition policy.\n\n"
    "Generate a comprehensive, program-specific Acquisition Strategy.\n\n"
    "Requirements:\n"
    "1. All content must be specific to this program — not generic boilerplate.\n"
    "2. module_table_rows must contain one row per module listed in 'Defined Modules' "
    "(or propose sensible defaults based on program context if no modules are defined).\n"
    "3. risk_register must include at least 4-6 acquisition-relevant risks.\n"
    "4. schedule_milestones must include Milestone A/B/C or equivalent program decision points.\n"
    "5. sources_used must only reference file_ids that appear in the Retrieved Evidence section; "
    "set to an empty list if no evidence files are available.\n"
    "6. cost_estimates should use data from the program fact-pack if available; otherwise estimate.\n\n"
)


@app.post("/programs/{program_id}/docs/acq_strategy/smart")
def generate_smart_acq_strategy(program_id: int, db: Session = Depends(get_db)):
    """Generate an AI-assisted Acquisition Strategy using structured LLM output (AcqStrategyPlan schema).

    Requires OPENAI_API_KEY to be set in the environment.
    Returns the generated DOCX directly as a file download.
    """
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=503,
            detail=(
                "Smart Acquisition Strategy generation requires an LLM API key. "
                "Please set the OPENAI_API_KEY environment variable and restart the server."
            ),
        )

    program = db.query(Program).filter(Program.id == program_id).first()
    if not program:
        raise HTTPException(status_code=404, detail="Program not found")

    brief = db.query(ProgramBrief).filter(ProgramBrief.program_id == program_id).first()
    answers_rows = db.query(ProgramAnswer).filter(ProgramAnswer.program_id == program_id).all()
    answers: dict[str, str | None] = {r.question_id: r.answer_text for r in answers_rows}
    files = (
        db.query(ProgramFile)
        .filter(ProgramFile.program_id == program_id)
        .order_by(ProgramFile.id.asc())
        .all()
    )
    modules = (
        db.query(Module)
        .filter(Module.program_id == program_id)
        .order_by(Module.id.asc())
        .all()
    )

    # ---- Build fact-pack ----
    fact_lines: list[str] = [f"Program Name: {program.name}"]
    if brief:
        if brief.program_description:
            fact_lines.append(f"Program Description: {brief.program_description}")
        if brief.dev_cost_estimate is not None:
            fact_lines.append(f"Estimated Development Cost: ${brief.dev_cost_estimate:,.0f}M")
        if brief.production_unit_cost is not None:
            fact_lines.append(f"Production Unit Cost: ${brief.production_unit_cost:,.0f}M")
        flags = []
        if brief.attritable:
            flags.append("Attritable")
        if brief.sustainment_tail:
            flags.append("Sustainment Tail")
        if brief.software_large_part:
            flags.append("Software-Intensive")
        if brief.mission_critical:
            flags.append("Mission-Critical")
        if brief.safety_critical:
            flags.append("Safety-Critical")
        if flags:
            fact_lines.append(f"Program Flags: {', '.join(flags)}")

    wizard_fields = {
        "a_program_description": "Program Description (Wizard)",
        "e_similar_previous_programs": "Heritage / Analogous Programs",
        "f_tech_challenges_and_risk_areas": "Technical Challenges and Risk Areas",
        "g_mosa_scenarios": "MOSA Scenarios",
        "h_candidate_modules": "Candidate Modules",
        "i_known_standards_architectures_mapping": "Known Standards / Architectures by Module",
        "j_obsolescence_candidates": "Obsolescence Candidates",
        "k_commercial_solutions_by_module": "Commercial Solutions by Module",
        "n_software_standards_architectures": "Software Standards / Architectures",
        "o_mosa_repo_searched": "MOSA Repository Search Performed",
    }
    for q_id, label in wizard_fields.items():
        val = answers.get(q_id)
        if val and val.strip():
            fact_lines.append(f"\n{label}:\n{val.strip()}")

    if modules:
        fact_lines.append("\nDefined Modules:")
        for m in modules:
            row = f"  - [Module] {m.name}"
            if m.rationale:
                row += f": {m.rationale}"
            if m.key_interfaces:
                row += f" | Interfaces: {m.key_interfaces}"
            if m.standards:
                row += f" | Standards: {m.standards}"
            row += f" | Tech Risk: {'Yes' if m.tech_risk else 'No'}"
            row += f" | Obsolescence Risk: {'Yes' if m.obsolescence_risk else 'No'}"
            row += f" | COTS Candidate: {'Yes' if m.cots_candidate else 'No'}"
            fact_lines.append(row)

    if files:
        fact_lines.append("\nUploaded Reference Files (file_id | filename):")
        for f in files:
            fact_lines.append(f"  - [{f.id}] {f.filename}")

    fact_pack = "\n".join(fact_lines)

    # ---- Grounded source retrieval ----
    _acq_brief_desc = brief.program_description if brief else ""
    _acq_query = _build_program_query(
        program.name, _acq_brief_desc, answers,
        "Acquisition Strategy contracting schedule milestones cost MOSA data rights competition",
    )
    _acq_sources, _acq_chunks = _retrieve_grounded_sources(program_id, _acq_query, db)

    # ---- Build LLM prompt ----
    system_prompt = (
        _ACQ_STRATEGY_SYSTEM_PROMPT_PREFIX
        + build_template_contract("acq_strategy")
        + _GROUNDING_INSTRUCTION
        + _SMART_ACQ_STRATEGY_JSON_SCHEMA
    )
    user_message = (
        "## PROGRAM FACT-PACK\n"
        f"{fact_pack}\n\n"
        "## SOURCES\n"
        "Use only these numbered sources for facts. Cite them in the citations field "
        "by their source number.\n\n"
        f"{_acq_sources}"
    )

    # ---- Call OpenAI API ----
    try:
        import openai as _openai
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="The 'openai' Python package is not installed. Run: pip install openai",
        )

    model = os.environ.get("OPENAI_MODEL", "gpt-5-mini")
    try:
        client = _openai.OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model=model,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
        )
    except _openai.AuthenticationError:
        raise HTTPException(
            status_code=503,
            detail="Invalid OPENAI_API_KEY. Please verify your API key and try again.",
        )
    except Exception as exc:
        raise HTTPException(status_code=503, detail=f"LLM service error: {exc}")

    # ---- Parse and validate ----
    raw_content = response.choices[0].message.content or ""
    try:
        raw_data = json.loads(raw_content)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=502, detail=f"LLM did not return valid JSON: {exc}")

    # ---- Enforce template conformance, fill "Not provided" defaults ----
    try:
        normalize_llm_output("acq_strategy", raw_data)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

    try:
        acq_plan = AcqStrategyPlan.model_validate(raw_data)
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"LLM produced an invalid Acquisition Strategy plan structure: {exc}",
        )

    # ---- Render DOCX ----
    generated_date = datetime.now(timezone.utc).strftime("%B %d, %Y")
    docx = build_acq_strategy_smart(program, brief, answers, files, modules, generated_date, acq_plan)

    # ---- Persist to disk + DB ----
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    out_dir = DATA_DIR / "programs" / str(program_id) / "docs" / timestamp
    out_dir.mkdir(parents=True, exist_ok=True)
    file_name = "AcqStrategy_Smart.docx"
    abs_path = out_dir / file_name
    docx.save(str(abs_path))

    # Traceability: save full plan JSON + source file IDs + citations
    traceability = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "model": model,
        "source_file_ids": [c["file_id"] for c in _acq_chunks],
        "citations": raw_data.get("citations", {}),
        "plan": raw_data,
    }
    (out_dir / "acq_strategy_plan_trace.json").write_text(
        json.dumps(traceability, indent=2, default=str)
    )

    rel_path = str(abs_path.relative_to(DATA_DIR))
    doc_row = ProgramDocument(program_id=program_id, doc_type="ACQ_STRATEGY_SMART", file_path=rel_path)
    db.add(doc_row)
    db.commit()

    return FileResponse(
        path=str(abs_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=file_name,
    )


@app.get("/documents/{doc_id}/download")
def download_document(doc_id: int, db: Session = Depends(get_db)):
    doc_row = db.query(ProgramDocument).filter(ProgramDocument.id == doc_id).first()
    if not doc_row:
        raise HTTPException(status_code=404, detail="Document not found")

    abs_path = DATA_DIR / doc_row.file_path
    if not abs_path.exists():
        raise HTTPException(status_code=404, detail="Document file not found on disk")

    return FileResponse(
        path=str(abs_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=abs_path.name,
    )


# ---------------------------------------------------------------------------
# A) Program context builder
# ---------------------------------------------------------------------------


@app.post("/programs/{program_id}/knowledge/build_context", response_model=ContextBuildOut)
def build_context(program_id: int, db: Session = Depends(get_db)):
    """Build and save a program knowledge context snapshot (context.json).

    Gathers: program record, brief, wizard answers, modules, uploaded file metadata.
    Saves to: backend/data/programs/{id}/context/context.json
    Returns: short summary + list of missing-info questions.
    """
    program = db.query(Program).filter(Program.id == program_id).first()
    if not program:
        raise HTTPException(status_code=404, detail="Program not found")

    brief = db.query(ProgramBrief).filter(ProgramBrief.program_id == program_id).first()
    answers_rows = db.query(ProgramAnswer).filter(ProgramAnswer.program_id == program_id).all()
    answers: dict[str, str | None] = {r.question_id: r.answer_text for r in answers_rows}
    modules = (
        db.query(Module)
        .filter(Module.program_id == program_id)
        .order_by(Module.id.asc())
        .all()
    )
    files = (
        db.query(ProgramFile)
        .filter(ProgramFile.program_id == program_id)
        .order_by(ProgramFile.id.asc())
        .all()
    )

    # (Re-)index program files into RagChunk for up-to-date retrieval
    index_program_files(program_id, files, db)

    ctx = build_program_context(program, brief, answers, modules, files, DATA_DIR)
    summary, missing = context_summary_and_gaps(ctx)

    ctx_path = str(
        (DATA_DIR / "programs" / str(program_id) / "context" / "context.json")
        .relative_to(DATA_DIR)
    )
    return ContextBuildOut(
        program_id=program_id,
        summary=summary,
        missing_info_questions=missing,
        context_path=ctx_path,
    )


def _compute_knowledge_status(program_id: int, db: Session) -> KnowledgeStatusOut:
    """Compute the knowledge health status for a program."""
    files = (
        db.query(ProgramFile)
        .filter(ProgramFile.program_id == program_id)
        .all()
    )
    file_count = len(files)
    text_file_count = sum(1 for f in files if f.extracted_text)

    chunk_count = db.query(FileChunk).filter(FileChunk.program_id == program_id).count()

    last_indexed_at = (
        db.query(func.max(FileChunk.created_at))
        .filter(FileChunk.program_id == program_id)
        .scalar()
    )

    # Per-file chunk counts joined with filenames
    per_file_rows = (
        db.query(FileChunk.file_id, func.count(FileChunk.id).label("cnt"))
        .filter(FileChunk.program_id == program_id)
        .group_by(FileChunk.file_id)
        .order_by(func.count(FileChunk.id).desc())
        .all()
    )
    file_id_to_name = {f.id: f.filename for f in files}
    top_files = [
        KnowledgeTopFile(file_id=row.file_id, filename=file_id_to_name.get(row.file_id, ""), chunk_count=row.cnt)
        for row in per_file_rows
    ]

    return KnowledgeStatusOut(
        program_id=program_id,
        file_count=file_count,
        text_file_count=text_file_count,
        chunk_count=chunk_count,
        last_indexed_at=last_indexed_at,
        top_files=top_files,
    )


@app.get("/programs/{program_id}/knowledge/status", response_model=KnowledgeStatusOut)
def knowledge_status(program_id: int, db: Session = Depends(get_db)):
    """Return knowledge health stats for the program's vector index."""
    if not db.query(Program).filter(Program.id == program_id).first():
        raise HTTPException(status_code=404, detail="Program not found")
    return _compute_knowledge_status(program_id, db)


@app.post("/programs/{program_id}/knowledge/index", response_model=KnowledgeStatusOut)
def index_knowledge(program_id: int, db: Session = Depends(get_db)):
    """Wipe and rebuild the vector embedding index for all files in the program.

    Loads the extracted text already stored on each ProgramFile, re-chunks it
    using the vector-optimised parameters (1 000-char chunks / 100-char overlap),
    computes OpenAI text-embedding-3-small embeddings in batches, and persists
    the resulting rows to file_chunks (replacing any prior rows for this program).

    Requires OPENAI_API_KEY to be set in the environment.
    """
    if not db.query(Program).filter(Program.id == program_id).first():
        raise HTTPException(status_code=404, detail="Program not found")

    files = (
        db.query(ProgramFile)
        .filter(ProgramFile.program_id == program_id)
        .all()
    )

    # Wipe all existing chunks for this program
    db.query(FileChunk).filter(FileChunk.program_id == program_id).delete()
    db.flush()

    # Build chunk records from each file's extracted text
    chunk_records: list[dict] = []
    for pf in files:
        raw_text = pf.extracted_text or ""
        if not raw_text.strip():
            continue
        for i, chunk in enumerate(_chunk_text(raw_text, VECTOR_CHUNK_SIZE, VECTOR_CHUNK_OVERLAP)):
            chunk_records.append({
                "program_id": program_id,
                "file_id": pf.id,
                "source_type": pf.source_type,
                "chunk_index": i,
                "chunk_text": chunk,
            })

    if not chunk_records:
        db.commit()
        return _compute_knowledge_status(program_id, db)

    # Embed all chunks in one batched call (≤100 texts per API request)
    texts = [r["chunk_text"] for r in chunk_records]
    try:
        embeddings = embed_texts(texts)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Embedding API error: {exc}") from exc

    for record, embedding in zip(chunk_records, embeddings):
        db.add(FileChunk(
            program_id=record["program_id"],
            file_id=record["file_id"],
            source_type=record["source_type"],
            chunk_index=record["chunk_index"],
            chunk_text=record["chunk_text"],
            embedding=embedding,
        ))

    db.commit()

    return _compute_knowledge_status(program_id, db)


@app.get("/programs/{program_id}/knowledge/retrieve")
def retrieve_knowledge(
    program_id: int,
    q: str = Query(..., min_length=1, description="Query string"),
    k: int = Query(12, ge=1, le=100, description="Number of chunks to return"),
    db: Session = Depends(get_db),
):
    """Vector similarity search over embedded chunks for a program.

    Returns the top-k most similar chunks ordered by descending cosine
    similarity score (1 = identical, 0 = orthogonal).

    Requires chunks to have been indexed via POST /programs/{id}/knowledge/index.
    """
    if not db.query(Program).filter(Program.id == program_id).first():
        raise HTTPException(status_code=404, detail="Program not found")

    from rag import retrieve_chunks as rag_retrieve

    try:
        results = rag_retrieve(db, program_id, q, k)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    return {"query": q, "k": k, "results": results}


# ---------------------------------------------------------------------------
# C) Smart MOSA Conformance Plan generator
# ---------------------------------------------------------------------------

_MOSA_SYSTEM_PROMPT_PREFIX = (
    "You are a defense acquisition expert specializing in Modular Open Systems Approach (MOSA) "
    "and MOSA Conformance Plans per PEO AVN MIG v2.2, 10 U.S.C. \u00a7 4401, and DoD Directive 5000.01.\n\n"
    "Generate a comprehensive, program-specific MOSA Conformance Plan.\n\n"
    "Requirements:\n"
    "1. All content must be specific to this program — not generic boilerplate.\n"
    "2. module_inventory must contain one row per module listed in the program fact-pack "
    "(or propose sensible defaults if no modules are defined).\n"
    "3. assessment_matrix must include at minimum the four MOSA minimum criteria: "
    "Open Standards Adoption, Competitive Upgrade Path, Government Technical Data Rights, "
    "and Well-Defined Modular Boundaries.\n"
    "4. sources_used must only reference chunk_ids / file_ids that appear in the SOURCES block; "
    "set to an empty list if no sources are available.\n"
    "5. verification_milestones should cover PDR, CDR, and IOT&E at minimum.\n"
    "6. risk_register should include at least 3-5 MOSA-relevant risks.\n\n"
)


@app.post("/programs/{program_id}/docs/mosa_conformance/smart")
def generate_smart_mosa_conformance_plan(program_id: int, db: Session = Depends(get_db)):
    """Generate an AI-assisted MOSA Conformance Plan using structured LLM output.

    Pipeline:
      1) Load context.json (build it if missing).
      2) (Re-)index program files into RagChunk.
      3) Retrieve relevant chunks for MOSA conformance plan content.
      4) Call OpenAI with MosaPlan JSON schema.
      5) Validate response with Pydantic MosaPlan.
      6) Render DOCX via build_smart_mosa_conformance_plan.
      7) Persist DOCX + plan JSON + chunk IDs for traceability.
      8) Return DOCX as file download.
    """
    from llm.client import get_client, get_model
    from llm.mosa_schema import MOSA_PLAN_JSON_SCHEMA

    # Verify API key before doing any DB work
    try:
        client = get_client()
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc))

    program = db.query(Program).filter(Program.id == program_id).first()
    if not program:
        raise HTTPException(status_code=404, detail="Program not found")

    brief = db.query(ProgramBrief).filter(ProgramBrief.program_id == program_id).first()
    answers_rows = db.query(ProgramAnswer).filter(ProgramAnswer.program_id == program_id).all()
    answers: dict[str, str | None] = {r.question_id: r.answer_text for r in answers_rows}
    modules = (
        db.query(Module)
        .filter(Module.program_id == program_id)
        .order_by(Module.id.asc())
        .all()
    )
    files = (
        db.query(ProgramFile)
        .filter(ProgramFile.program_id == program_id)
        .order_by(ProgramFile.id.asc())
        .all()
    )

    # Step 1: Load or build context
    ctx = load_context(program_id, DATA_DIR)
    if ctx is None:
        index_program_files(program_id, files, db)
        ctx = build_program_context(program, brief, answers, modules, files, DATA_DIR)
    else:
        # Re-index in case files changed since last context build
        index_program_files(program_id, files, db)

    # Step 2: Ensure reference docs are indexed
    index_reference_docs(REFERENCE_DOCS_DIR, db, _extract_text)

    # Step 3: Build fact-pack from context
    fact_lines: list[str] = [f"Program Name: {program.name}"]
    brief_data = ctx.get("brief") or {}
    if brief_data.get("program_description"):
        fact_lines.append(f"Program Description: {brief_data['program_description']}")
    if brief_data.get("dev_cost_estimate") is not None:
        fact_lines.append(f"Estimated Development Cost: ${brief_data['dev_cost_estimate']:,.0f}")
    if brief_data.get("production_unit_cost") is not None:
        fact_lines.append(f"Production Unit Cost: ${brief_data['production_unit_cost']:,.0f}")

    flags = [
        k.replace("_", " ").title()
        for k in ("attritable", "sustainment_tail", "software_large_part", "mission_critical", "safety_critical")
        if brief_data.get(k)
    ]
    if flags:
        fact_lines.append(f"Program Flags: {', '.join(flags)}")

    wizard_fields = {
        "a_program_description": "Program Description (Wizard)",
        "f_tech_challenges_and_risk_areas": "Technical Challenges and Risk Areas",
        "g_mosa_scenarios": "MOSA Scenarios",
        "h_candidate_modules": "Candidate Modules",
        "i_known_standards_architectures_mapping": "Known Standards / Architectures by Module",
        "j_obsolescence_candidates": "Obsolescence Candidates",
        "k_commercial_solutions_by_module": "Commercial Solutions by Module",
        "n_software_standards_architectures": "Software Standards / Architectures",
        "o_mosa_repo_searched": "MOSA Repository Search Performed",
    }
    wizard_answers = ctx.get("wizard_answers") or {}
    for q_id, label in wizard_fields.items():
        val = wizard_answers.get(q_id)
        if val and str(val).strip():
            fact_lines.append(f"\n{label}:\n{val.strip()}")

    ctx_modules = ctx.get("modules") or []
    if ctx_modules:
        fact_lines.append("\nDefined Modules:")
        for m in ctx_modules:
            row = f"  - [Module] {m['name']}"
            if m.get("rationale"):
                row += f": {m['rationale']}"
            if m.get("key_interfaces"):
                row += f" | Interfaces: {m['key_interfaces']}"
            if m.get("standards"):
                row += f" | Standards: {m['standards']}"
            row += f" | Tech Risk: {'Yes' if m.get('tech_risk') else 'No'}"
            row += f" | Obsolescence Risk: {'Yes' if m.get('obsolescence_risk') else 'No'}"
            row += f" | COTS Candidate: {'Yes' if m.get('cots_candidate') else 'No'}"
            fact_lines.append(row)

    fact_pack = "\n".join(fact_lines)

    # ---- Grounded source retrieval ----
    _mosa_brief_desc = brief_data.get("program_description", "")
    _mosa_query = _build_program_query(
        program.name, _mosa_brief_desc, answers,
        "MOSA conformance plan modular open systems interface standards open architecture",
    )
    _mosa_sources, _mosa_source_chunks = _retrieve_grounded_sources(program_id, _mosa_query, db)

    # Step 5: Call OpenAI
    system_prompt = (
        _MOSA_SYSTEM_PROMPT_PREFIX
        + build_template_contract("mcp")
        + _GROUNDING_INSTRUCTION
        + MOSA_PLAN_JSON_SCHEMA
    )
    user_message = (
        "## PROGRAM FACT-PACK\n"
        f"{fact_pack}\n\n"
        "## SOURCES\n"
        "Use only these numbered sources for facts. Cite them in the citations field "
        "by their source number.\n\n"
        f"{_mosa_sources}"
    )

    model = get_model()
    try:
        response = client.chat.completions.create(
            model=model,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
        )
    except Exception as exc:
        import openai as _openai
        if isinstance(exc, _openai.AuthenticationError):
            raise HTTPException(status_code=503, detail="Invalid OPENAI_API_KEY.")
        raise HTTPException(status_code=503, detail=f"LLM service error: {exc}")

    raw_content = response.choices[0].message.content or ""
    try:
        raw_data = json.loads(raw_content)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=502, detail=f"LLM did not return valid JSON: {exc}")

    # ---- Enforce template conformance, fill "Not provided" defaults ----
    try:
        normalize_llm_output("mcp", raw_data)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

    try:
        mosa_plan = MosaPlan.model_validate(raw_data)
    except Exception as exc:
        raise HTTPException(
            status_code=400, detail=f"LLM produced an invalid MosaPlan structure: {exc}"
        )

    # Step 6: Render DOCX
    generated_date = datetime.now(timezone.utc).strftime("%B %d, %Y")
    docx = build_smart_mosa_conformance_plan(program, generated_date, mosa_plan)

    # Step 7: Persist DOCX + plan JSON + chunk IDs
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    out_dir = DATA_DIR / "programs" / str(program_id) / "docs" / timestamp
    out_dir.mkdir(parents=True, exist_ok=True)

    file_name = "MOSAConformancePlan_Smart.docx"
    abs_path = out_dir / file_name
    docx.save(str(abs_path))

    # Traceability: save the full plan JSON + source file IDs used
    traceability = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "model": model,
        "source_file_ids": [c["file_id"] for c in _mosa_source_chunks],
        "citations": raw_data.get("citations", {}),
        "plan": raw_data,
    }
    (out_dir / "mosa_plan_trace.json").write_text(
        json.dumps(traceability, indent=2, default=str)
    )

    rel_path = str(abs_path.relative_to(DATA_DIR))
    doc_row = ProgramDocument(
        program_id=program_id,
        doc_type="MOSA_CONFORMANCE_PLAN_SMART",
        file_path=rel_path,
    )
    db.add(doc_row)
    db.commit()

    return FileResponse(
        path=str(abs_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=file_name,
    )


# ---------------------------------------------------------------------------
# D) Smart SEP (Systems Engineering Plan) – shared LLM helper + two endpoints
# ---------------------------------------------------------------------------

_SEP_SYSTEM_PROMPT_PREFIX = (
    "You are a defense acquisition systems engineering expert specializing in Systems Engineering Plans (SEPs) "
    "per the OSD SEP Guide v4.1, DoDI 5000.02, and DoD Instruction 5000.88.\n\n"
    "Generate a comprehensive, program-specific Systems Engineering Plan.\n\n"
    "Requirements:\n"
    "1. All content must be specific to this program — not generic boilerplate.\n"
    "2. tech_reviews must include at minimum: SRR, SFR, PDR, CDR, and TRR.\n"
    "3. risk_register must contain at least 4 program-relevant SE risks with owners.\n"
    "4. specialty_eng must address all six fields based on the program's characteristics; "
    "write 'Not applicable – [reason]' only when genuinely inapplicable.\n"
    "5. sources_used must only reference file_ids that appear in the SOURCES block; "
    "set to an empty list if no sources are available.\n"
    "6. appendices.references must include relevant DoD/service regulations and standards.\n"
    "7. config_mgmt.baselines must cover at minimum: Functional, Allocated, and Product baselines.\n\n"
)

_SEP_QUERIES = [
    "systems engineering plan SEP approach methodology",
    "technical review PDR CDR SRR SFR TRR milestone",
    "requirements management traceability RTM DOORS",
    "architecture design interface control MOSA modular open systems",
    "risk management technical risk mitigation",
    "configuration management baseline ECP change control",
    "verification validation test evaluation acceptance",
    "data management technical data rights CDRL data items",
    "cybersecurity RMF cyber security",
    "safety system safety airworthiness certification",
    "human factors HSI human systems integration",
    "reliability maintainability RAM sustainment",
]

_SEP_PREFERRED_KEYWORDS = ("sep", "se plan", "systems eng", "semp", "mosa", "rfi", "strategy")


def _sep_plan_from_llm(
    program_id: int,
    db: Session,
) -> tuple:
    """Shared helper: build context, call OpenAI, return (program, sep_plan, raw_data, source_chunks, model).

    Raises HTTPException on any error so callers can propagate cleanly.
    """
    from llm.client import get_client, get_model
    from llm.sep_schema import SEP_PLAN_JSON_SCHEMA

    try:
        client = get_client()
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc))

    program = db.query(Program).filter(Program.id == program_id).first()
    if not program:
        raise HTTPException(status_code=404, detail="Program not found")

    brief = db.query(ProgramBrief).filter(ProgramBrief.program_id == program_id).first()
    answers_rows = db.query(ProgramAnswer).filter(ProgramAnswer.program_id == program_id).all()
    answers: dict[str, str | None] = {r.question_id: r.answer_text for r in answers_rows}
    modules = (
        db.query(Module)
        .filter(Module.program_id == program_id)
        .order_by(Module.id.asc())
        .all()
    )
    files = (
        db.query(ProgramFile)
        .filter(ProgramFile.program_id == program_id)
        .order_by(ProgramFile.id.asc())
        .all()
    )

    # Load or build context; re-index program files
    ctx = load_context(program_id, DATA_DIR)
    if ctx is None:
        index_program_files(program_id, files, db)
        ctx = build_program_context(program, brief, answers, modules, files, DATA_DIR)
    else:
        index_program_files(program_id, files, db)

    index_reference_docs(REFERENCE_DOCS_DIR, db, _extract_text)

    # Build fact-pack
    fact_lines: list[str] = [f"Program Name: {program.name}"]
    if brief:
        if brief.program_description:
            fact_lines.append(f"Program Description: {brief.program_description}")
        if brief.dev_cost_estimate is not None:
            fact_lines.append(f"Estimated Development Cost: ${brief.dev_cost_estimate:,.0f}M")
        if brief.production_unit_cost is not None:
            fact_lines.append(f"Production Unit Cost: ${brief.production_unit_cost:,.0f}M")
        flags = []
        if brief.attritable:
            flags.append("Attritable")
        if brief.sustainment_tail:
            flags.append("Sustainment Tail")
        if brief.software_large_part:
            flags.append("Software-Intensive")
        if brief.mission_critical:
            flags.append("Mission-Critical")
        if brief.safety_critical:
            flags.append("Safety-Critical")
        if flags:
            fact_lines.append(f"Program Flags: {', '.join(flags)}")

    wizard_fields = {
        "a_program_description": "Program Description (Wizard)",
        "b_acquisition_phase": "Acquisition Phase",
        "c_program_type": "Program Type",
        "d_contracting_approach": "Contracting Approach",
        "e_similar_previous_programs": "Heritage / Analogous Programs",
        "f_tech_challenges_and_risk_areas": "Technical Challenges and Risk Areas",
        "g_mosa_scenarios": "MOSA Scenarios",
        "h_candidate_modules": "Candidate Modules",
        "i_known_standards_architectures_mapping": "Known Standards / Architectures by Module",
        "j_obsolescence_candidates": "Obsolescence Candidates",
        "k_commercial_solutions_by_module": "Commercial Solutions by Module",
        "n_software_standards_architectures": "Software Standards / Architectures",
        "o_mosa_repo_searched": "MOSA Repository Search Performed",
    }
    wizard_answers = ctx.get("wizard_answers") or {}
    for q_id, label in wizard_fields.items():
        val = wizard_answers.get(q_id) or answers.get(q_id)
        if val and str(val).strip():
            fact_lines.append(f"\n{label}:\n{str(val).strip()}")

    ctx_modules = ctx.get("modules") or []
    if ctx_modules:
        fact_lines.append("\nDefined Modules:")
        for m in ctx_modules:
            row = f"  - [Module] {m['name']}"
            if m.get("rationale"):
                row += f": {m['rationale']}"
            if m.get("key_interfaces"):
                row += f" | Interfaces: {m['key_interfaces']}"
            if m.get("standards"):
                row += f" | Standards: {m['standards']}"
            row += f" | Tech Risk: {'Yes' if m.get('tech_risk') else 'No'}"
            row += f" | Obsolescence Risk: {'Yes' if m.get('obsolescence_risk') else 'No'}"
            row += f" | COTS Candidate: {'Yes' if m.get('cots_candidate') else 'No'}"
            fact_lines.append(row)

    if files:
        fact_lines.append("\nUploaded Reference Files (file_id | filename):")
        for f in files:
            fact_lines.append(f"  - [{f.id}] {f.filename}")

    fact_pack = "\n".join(fact_lines)
    knowledge_summary = ctx.get("knowledge_summary") or ""

    # ---- Grounded source retrieval ----
    _sep_brief_desc = brief.program_description if brief else ""
    _sep_query = _build_program_query(
        program.name, _sep_brief_desc, answers,
        "Systems Engineering Plan technical reviews requirements architecture MOSA risk verification",
    )
    _sep_sources, _sep_source_chunks = _retrieve_grounded_sources(program_id, _sep_query, db)

    # Call OpenAI
    system_prompt = (
        _SEP_SYSTEM_PROMPT_PREFIX
        + build_template_contract("sep")
        + _GROUNDING_INSTRUCTION
        + SEP_PLAN_JSON_SCHEMA
    )
    user_message = f"## PROGRAM FACT-PACK\n{fact_pack}\n\n"
    if knowledge_summary:
        user_message += f"## KNOWLEDGE SUMMARY\n{knowledge_summary}\n\n"
    user_message += (
        "## SOURCES\n"
        "Use only these numbered sources for facts. Cite them in the citations field "
        "by their source number.\n\n"
        f"{_sep_sources}"
    )

    model = get_model()
    try:
        response = client.chat.completions.create(
            model=model,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
        )
    except Exception as exc:
        import openai as _openai
        if isinstance(exc, _openai.AuthenticationError):
            raise HTTPException(status_code=503, detail="Invalid OPENAI_API_KEY.")
        raise HTTPException(status_code=503, detail=f"LLM service error: {exc}")

    raw_content = response.choices[0].message.content or ""
    try:
        raw_data = json.loads(raw_content)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=502, detail=f"LLM did not return valid JSON: {exc}")

    # ---- Enforce template conformance, fill "Not provided" defaults ----
    try:
        normalize_llm_output("sep", raw_data)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

    try:
        sep_plan = SepPlan.model_validate(raw_data)
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"LLM produced an invalid SepPlan structure: {exc}",
        )

    return program, sep_plan, raw_data, _sep_source_chunks, model


@app.post("/programs/{program_id}/docs/sep/smart/plan")
def generate_smart_sep_plan(program_id: int, db: Session = Depends(get_db)):
    """Generate an AI-assisted SEP as a validated JSON plan (no DOCX rendering).

    Calls _sep_plan_from_llm for all LLM work, then persists a traceability JSON
    and returns the validated SepPlan as JSON.
    """
    _program, sep_plan, raw_data, source_chunks, model = _sep_plan_from_llm(program_id, db)

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    out_dir = DATA_DIR / "programs" / str(program_id) / "docs" / timestamp
    out_dir.mkdir(parents=True, exist_ok=True)

    traceability = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "model": model,
        "source_file_ids": [c["file_id"] for c in source_chunks],
        "citations": raw_data.get("citations", {}),
        "plan": raw_data,
    }
    trace_path = out_dir / "sep_plan_trace.json"
    trace_path.write_text(json.dumps(traceability, indent=2, default=str))

    doc_row = ProgramDocument(
        program_id=program_id,
        doc_type="SEP_SMART_PLAN",
        file_path=str(trace_path.relative_to(DATA_DIR)),
    )
    db.add(doc_row)
    db.commit()

    return sep_plan.model_dump()


@app.post("/programs/{program_id}/docs/sep/smart")
def generate_smart_sep_docx(program_id: int, db: Session = Depends(get_db)):
    """Generate an AI-assisted SEP and return it as a DOCX file attachment.

    Calls _sep_plan_from_llm for all LLM work, renders the DOCX via
    build_sep_smart, persists the DOCX + traceability JSON, and streams the
    file back as 'SEP_Smart.docx'.
    """
    program, sep_plan, raw_data, source_chunks, model = _sep_plan_from_llm(program_id, db)

    generated_date = datetime.now(timezone.utc).strftime("%B %d, %Y")
    docx = build_sep_smart(sep_plan, program, generated_date)

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    out_dir = DATA_DIR / "programs" / str(program_id) / "docs" / timestamp
    out_dir.mkdir(parents=True, exist_ok=True)

    file_name = "SEP_Smart.docx"
    abs_path = out_dir / file_name
    docx.save(str(abs_path))

    traceability = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "model": model,
        "source_file_ids": [c["file_id"] for c in source_chunks],
        "citations": raw_data.get("citations", {}),
        "plan": raw_data,
    }
    (out_dir / "sep_plan_trace.json").write_text(
        json.dumps(traceability, indent=2, default=str)
    )

    doc_row = ProgramDocument(
        program_id=program_id,
        doc_type="SEP_SMART",
        file_path=str(abs_path.relative_to(DATA_DIR)),
    )
    db.add(doc_row)
    db.commit()

    return FileResponse(
        path=str(abs_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=file_name,
    )
