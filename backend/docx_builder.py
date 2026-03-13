"""docx_builder.py – Word document generation helpers and section builders.

Public API:
    build_acq_strategy(...) -> DocxDocument
    build_rfi(...)          -> DocxDocument
"""

from __future__ import annotations

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import Module, Program, ProgramBrief, ProgramFile

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

MOSA_MINIMUM_CRITERIA: list[tuple[str, str]] = [
    (
        "Adoption of Open Standards",
        "Design the system using open, consensus-based interface standards to prevent proprietary lock-in "
        "and enable competitive sourcing of components throughout the program lifecycle.",
    ),
    (
        "Competitive Upgrade Path",
        "Architect the system so that individual modules or subsystems can be upgraded or replaced by "
        "competing suppliers, reducing total ownership cost and stimulating innovation.",
    ),
    (
        "Government Technical Data Rights",
        "Acquire sufficient technical data rights — including interface control documents and design data — "
        "to allow the Government to re-compete, support, or modify components independently of the original developer.",
    ),
    (
        "Well-Defined Modular Boundaries",
        "Implement clearly defined functional boundaries and standardized interfaces between modules so that "
        "changes to one module do not cascade to others, enabling incremental modernization.",
    ),
]

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def format_currency(value: float) -> str:
    """Return a dollar-formatted string without an 'M' suffix: $15,000,000"""
    return f"${value:,.0f}"


def safe_text(value) -> str:
    """Return value as a stripped string, or 'Not provided' if None/empty."""
    if value is None:
        return "Not provided"
    text = str(value).strip()
    return text if text else "Not provided"


def _add_page_number_field(paragraph) -> None:
    """Append a PAGE field code run to a paragraph."""
    run = paragraph.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _configure_footer(doc: DocxDocument, program_name: str) -> None:
    """Set footer on every section: program name left, page number centered."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        name_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        name_para.clear()
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = name_para.add_run(program_name)
        name_run.font.size = Pt(9)
        name_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        page_para = footer.add_paragraph()
        page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number_field(page_para)


def _set_cell_text(
    cell, text: str, bold: bool = False, center: bool = False, font_size_pt: int = 10
) -> None:
    """Write text into a table cell with optional bold/center formatting."""
    para = cell.paragraphs[0]
    # Remove all existing run elements so runs[0] is the one we control
    for r_elem in para._p.findall(qn("w:r")):
        para._p.remove(r_elem)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size_pt)


def _shade_cell(cell, hex_color: str = "D9E1F2") -> None:
    """Apply a solid background fill to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _make_table(doc: DocxDocument, headers: list[str], col_widths: list) -> object:
    """Create a Table Grid table with a bold, shaded header row."""
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, header_text in enumerate(headers):
        _set_cell_text(hdr_cells[i], header_text, bold=True)
        _shade_cell(hdr_cells[i])
        hdr_cells[i].width = col_widths[i]
    return tbl


# ---------------------------------------------------------------------------
# Shared page / doc setup
# ---------------------------------------------------------------------------


def _setup_margins(doc: DocxDocument) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)


def _add_title_page(
    doc: DocxDocument, program: Program, generated_date: str, subtitle: str
) -> None:
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    run = title_para.add_run(program.name)
    run.bold = True
    run.font.size = Pt(24)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_before = Pt(12)
    sub_run = sub_para.add_run(subtitle)
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(8)
    date_para.add_run(generated_date).font.size = Pt(11)


# ---------------------------------------------------------------------------
# Heading helpers
# ---------------------------------------------------------------------------


def _h1(doc: DocxDocument, text: str) -> None:
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(8)


def _h2(doc: DocxDocument, text: str) -> None:
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)


def _h3(doc: DocxDocument, text: str) -> None:
    h = doc.add_heading(text, level=3)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# ACQ STRATEGY section builders
# ---------------------------------------------------------------------------


def _build_executive_summary(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
) -> None:
    _h1(doc, "1. Executive Summary")
    doc.add_paragraph(
        f"This Acquisition Strategy documents the Government's approach to acquiring the {program.name} program "
        "in accordance with a Modular Open Systems Approach (MOSA) as required by 10 U.S.C. \u00a7\u00a04401. "
        "This strategy outlines program objectives, modular contracting approach, data rights posture, "
        "standards and architectures, industry engagement activities, and risk mitigation through modularity."
    )
    bullets: list[str] = [f"Program: {program.name}"]
    if brief:
        if brief.dev_cost_estimate is not None:
            bullets.append(f"Estimated development cost: {format_currency(brief.dev_cost_estimate)}")
        if brief.production_unit_cost is not None:
            bullets.append(f"Estimated production unit cost: {format_currency(brief.production_unit_cost)}")
        if brief.mission_critical:
            bullets.append("Mission-critical system")
        if brief.safety_critical:
            bullets.append("Safety-critical system")
        if brief.attritable:
            bullets.append("Attritable system (low-cost, expendable)")
        if brief.sustainment_tail:
            bullets.append("Long sustainment tail anticipated")
        if brief.software_large_part:
            bullets.append("Software constitutes a major portion of the system")
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")


def _build_program_overview(
    doc: DocxDocument,
    program_desc: str,
    answers: dict[str, str | None],
) -> None:
    _h1(doc, "2. Program Overview")
    doc.add_paragraph(program_desc)
    similar = (answers.get("e_similar_previous_programs") or "").strip()
    if similar:
        _h2(doc, "2.1 Similar Previous Programs")
        doc.add_paragraph(similar)


def _build_constraints(
    doc: DocxDocument,
    brief: ProgramBrief | None,
    answers: dict[str, str | None],
    files: list[ProgramFile],
    tech: str,
) -> None:
    _h1(doc, "3. Constraints")

    _h2(doc, "3.1 Technical Constraints")
    if tech:
        for line in tech.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("No specific technical constraints identified at this time.")

    _h2(doc, "3.2 Cost & Schedule Constraints")
    cost_bullets: list[str] = []
    if brief:
        if brief.dev_cost_estimate is not None:
            cost_bullets.append(f"Development cost estimate: {format_currency(brief.dev_cost_estimate)}")
        if brief.production_unit_cost is not None:
            cost_bullets.append(f"Production unit cost estimate: {format_currency(brief.production_unit_cost)}")
    if cost_bullets:
        for b in cost_bullets:
            doc.add_paragraph(b, style="List Bullet")
    else:
        doc.add_paragraph("Cost and schedule constraints are to be determined.")

    _h2(doc, "3.3 Reference Documents")
    if files:
        doc.add_paragraph("The following reference documents have been uploaded to support this acquisition:")
        for f in files:
            doc.add_paragraph(f.filename, style="List Bullet")
    else:
        doc.add_paragraph("No reference documents have been attached.")


def _build_modular_contracting_strategy(
    doc: DocxDocument,
    modules: list[Module],
    mosa_scenarios: str,
) -> None:
    _h1(doc, "4. Modular Contracting Strategy")
    doc.add_paragraph(
        "The program will employ a modular contracting strategy consistent with MOSA principles. "
        "Each module will be competed independently where market conditions allow, using open interface "
        "standards to preserve the Government's ability to re-compete, upgrade, and sustain components "
        "across the program lifecycle."
    )

    _h2(doc, "4.1 Module Competition Plan")
    col_widths = [Inches(1.4), Inches(1.4), Inches(1.6), Inches(1.6)]
    tbl = doc.add_table(rows=len(modules) + 1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Module", "Competition Strategy", "Functional Boundary", "Key Interfaces"]):
        _set_cell_text(hdr[i], text, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, mod in enumerate(modules, start=1):
        cells = tbl.rows[r].cells
        strategy = "Full & Open – COTS" if mod.cots_candidate else "Competitive – MOSA Compliant"
        _set_cell_text(cells[0], safe_text(mod.name))
        _set_cell_text(cells[1], strategy)
        _set_cell_text(cells[2], safe_text(mod.rationale))
        _set_cell_text(cells[3], safe_text(mod.key_interfaces))
        for i, cell in enumerate(cells):
            cell.width = col_widths[i]

    if mosa_scenarios:
        _h2(doc, "4.2 MOSA Contracting Scenarios")
        for line in mosa_scenarios.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")


def _build_data_rights(
    doc: DocxDocument,
    i_answer: str,
) -> None:
    _h1(doc, "5. Data Rights / Technical Data Strategy")
    doc.add_paragraph(
        "The Government will acquire sufficient technical data rights to enable competition, independent "
        "support, and future modernization. The following posture governs data rights across all contract "
        "actions under this program:"
    )
    data_rights_items = [
        (
            "Interface Control Documents (ICDs)",
            "Government will obtain unlimited rights to all ICDs to preserve competitive re-sourcing of modules.",
        ),
        (
            "Design Data",
            "Government will negotiate for at minimum Government Purpose Rights (GPR) on all major design "
            "artifacts; unlimited rights will be sought for software developed exclusively with Government funding.",
        ),
        (
            "Firmware & Embedded Software",
            "Rights will be negotiated to allow the Government to modify and redistribute firmware for "
            "sustainment and security patching.",
        ),
        (
            "Licensing Posture",
            "Preference for royalty-free, perpetual licenses on all deliverable technical data. Proprietary "
            "licenses will be documented and approved by the Program Manager.",
        ),
        (
            "Standards & Specifications",
            "All applicable open standards incorporated by reference will be identified in the contract, "
            "with Government retention of copies.",
        ),
    ]
    for title, desc in data_rights_items:
        p = doc.add_paragraph(style="List Bullet")
        run_title = p.add_run(f"{title}: ")
        run_title.bold = True
        p.add_run(desc)

    if i_answer:
        doc.add_paragraph()
        doc.add_paragraph("Known ICDs and interface standards by module:")
        doc.add_paragraph(i_answer)


def _build_standards_architectures(
    doc: DocxDocument,
    modules: list[Module],
    answers: dict[str, str | None],
    n_answer: str,
) -> None:
    _h1(doc, "6. Standards & Architectures")
    doc.add_paragraph(
        "The following standards and reference architectures have been identified for this program. "
        "Contractors shall demonstrate compliance with applicable standards at each design review."
    )

    if n_answer:
        _h2(doc, "6.1 Software Standards & Architectures")
        doc.add_paragraph(n_answer)

    _h2(doc, "6.2 Module-Level Standards")
    col_widths = [Inches(2.0), Inches(4.0)]
    tbl = doc.add_table(rows=len(modules) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Module", "Applicable Standards"]):
        _set_cell_text(hdr[i], text, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, mod in enumerate(modules, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], safe_text(mod.name))
        _set_cell_text(cells[1], safe_text(mod.standards or mod.key_interfaces))
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]

    o_answer = (answers.get("o_mosa_repo_searched") or "").strip()
    if o_answer:
        _h2(doc, "6.3 MOSA Repository Search")
        status = "has been completed" if o_answer.lower() == "yes" else "has not yet been completed"
        doc.add_paragraph(
            f"A search of the MOSA repository {status} for this program. "
            "Results will be incorporated into the standards selection process."
        )


def _build_industry_engagement(doc: DocxDocument) -> None:
    _h1(doc, "7. Industry Engagement Plan")
    doc.add_paragraph(
        "The program office will engage with industry early and often to inform the acquisition strategy, "
        "validate technical assumptions, and promote competition. The following activities are planned:"
    )
    engagement_items = [
        (
            "Request for Information (RFI)",
            "Issue an RFI to solicit industry capability statements, proposed architectures, and open "
            "standards recommendations prior to Draft RFP release.",
        ),
        (
            "Industry Day",
            "Host an Industry Day to present program objectives, MOSA requirements, and modular "
            "decomposition concepts; solicit vendor feedback on proposed interfaces.",
        ),
        (
            "One-on-One Sessions",
            "Conduct voluntary one-on-one sessions with interested vendors to discuss proprietary "
            "capabilities without disclosure to competitors.",
        ),
        (
            "Draft RFP Review",
            "Release Draft RFP for public comment with a minimum 30-day review period; publish "
            "responses to industry questions.",
        ),
        (
            "Pre-Solicitation Conference",
            "Hold a pre-solicitation conference to clarify requirements and answer questions from "
            "prospective offerors.",
        ),
    ]
    for title, desc in engagement_items:
        p = doc.add_paragraph(style="List Bullet")
        run_title = p.add_run(f"{title}: ")
        run_title.bold = True
        p.add_run(desc)


def _build_risk_register(
    doc: DocxDocument,
    modules: list[Module],
    tech: str,
) -> None:
    _h1(doc, "8. Risk Register")
    risks: list[tuple[str, str, str, str, str]] = [
        (
            "Vendor Lock-in",
            "Single-source dependencies on proprietary interfaces limiting re-competition",
            "Medium", "High",
            "Enforce open standards; acquire ICDs; modular boundaries in contract",
        ),
        (
            "Technology Obsolescence",
            "Rapid COTS obsolescence requiring costly mid-cycle replacements",
            "Medium", "Medium",
            "COTS insertion planning; modular design enables component swap without system redesign",
        ),
        (
            "Integration Complexity",
            "Increased integration risk from multiple competing vendors across modules",
            "Medium", "Medium",
            "Rigorous ICD management; system integrator role; integration testing milestones",
        ),
        (
            "Schedule Risk",
            "Delays in one module cascading to system-level schedule",
            "Low", "High",
            "Independent module schedules; parallel development; well-defined acceptance criteria",
        ),
    ]

    for mod in modules:
        if mod.tech_risk:
            risks.append((
                f"Tech Risk: {mod.name[:30]}",
                f"{mod.name} identified as a high-risk technical area",
                "High", "High",
                f"Early prototype; incremental demo milestones; modular boundary isolates risk to {mod.name[:30]}",
            ))
        if mod.obsolescence_risk:
            risks.append((
                f"Obsolescence: {mod.name[:25]}",
                f"{mod.name} identified as an obsolescence risk area",
                "Medium", "Medium",
                "Modular replacement path; COTS insertion plan; technology refresh in LCMP",
            ))

    if tech:
        first_challenge = next(
            (ln.strip() for ln in tech.splitlines() if ln.strip()), None
        )
        if first_challenge:
            risks.append((
                "Program Tech Risk",
                first_challenge[:120],
                "TBD", "TBD",
                "Modular architecture limits blast radius; MOSA enables competitive re-sourcing",
            ))

    col_widths = [Inches(1.1), Inches(1.7), Inches(0.8), Inches(0.7), Inches(1.7)]
    tbl = doc.add_table(rows=len(risks) + 1, cols=5)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Risk Area", "Description", "Likelihood", "Impact", "Mitigation"]):
        _set_cell_text(hdr[i], text, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (risk_area, desc, likelihood, impact, mitigation) in enumerate(risks, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], risk_area)
        _set_cell_text(cells[1], desc)
        _set_cell_text(cells[2], likelihood, center=True)
        _set_cell_text(cells[3], impact, center=True)
        _set_cell_text(cells[4], mitigation)
        for i, cell in enumerate(cells):
            cell.width = col_widths[i]


def _build_mosa_implementation(
    doc: DocxDocument,
    answers: dict[str, str | None],
    mosa_scenarios: str,
) -> None:
    _h1(doc, "9. MOSA Implementation Plan")
    doc.add_paragraph(
        "This program will implement a Modular Open Systems Approach (MOSA) consistent with 10 U.S.C. \u00a7\u00a04401 "
        "and DoD Directive 5000.01. The following criteria constitute the minimum MOSA requirements for this program:"
    )

    for title, description in MOSA_MINIMUM_CRITERIA:
        p = doc.add_paragraph(style="List Bullet")
        run_title = p.add_run(f"{title}: ")
        run_title.bold = True
        p.add_run(description)

    _h2(doc, "9.1 MOSA Verification")
    doc.add_paragraph(
        "MOSA compliance will be verified at each major program milestone (PDR, CDR, IOT&E) "
        "through the following:"
    )
    verification_items = [
        (
            "Interface Control Document (ICD) Review",
            "All module interfaces shall be documented, Government-owned, and compliant with open standards.",
        ),
        (
            "Open Standards Compliance Matrix",
            "Contractor shall maintain and deliver an updated compliance matrix at each design review.",
        ),
        (
            "Competitive Demonstration",
            "By CDR, contractor shall demonstrate that at least one module can be replaced by an alternative "
            "vendor using only the published ICDs.",
        ),
        (
            "Technical Data Package (TDP) Delivery",
            "Deliverable TDPs shall be reviewed for completeness and Government usability.",
        ),
    ]
    for title, desc in verification_items:
        p = doc.add_paragraph(style="List Bullet")
        run_title = p.add_run(f"{title}: ")
        run_title.bold = True
        p.add_run(desc)

    if mosa_scenarios:
        _h2(doc, "9.2 Program-Specific MOSA Scenarios")
        for line in mosa_scenarios.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")


# ---------------------------------------------------------------------------
# Public document builders
# ---------------------------------------------------------------------------


def build_acq_strategy(
    program: Program,
    brief: ProgramBrief | None,
    answers: dict[str, str | None],
    files: list[ProgramFile],
    modules: list[Module],
    generated_date: str,
) -> DocxDocument:
    doc = DocxDocument()
    _setup_margins(doc)
    _add_title_page(doc, program, generated_date, "Acquisition Strategy")
    doc.add_page_break()

    brief_desc = (brief.program_description or "").strip() if brief else ""
    wizard_desc = (answers.get("a_program_description") or "").strip()
    program_desc = brief_desc or wizard_desc or "No program description provided."
    tech = (answers.get("f_tech_challenges_and_risk_areas") or "").strip()
    mosa_scenarios = (answers.get("g_mosa_scenarios") or "").strip()
    i_answer = (answers.get("i_known_standards_architectures_mapping") or "").strip()
    n_answer = (answers.get("n_software_standards_architectures") or "").strip()

    _build_executive_summary(doc, program, brief)
    doc.add_page_break()

    _build_program_overview(doc, program_desc, answers)
    doc.add_page_break()

    _build_constraints(doc, brief, answers, files, tech)
    doc.add_page_break()

    _build_modular_contracting_strategy(doc, modules, mosa_scenarios)
    doc.add_page_break()

    _build_data_rights(doc, i_answer)
    doc.add_page_break()

    _build_standards_architectures(doc, modules, answers, n_answer)
    doc.add_page_break()

    _build_industry_engagement(doc)
    doc.add_page_break()

    _build_risk_register(doc, modules, tech)
    doc.add_page_break()

    _build_mosa_implementation(doc, answers, mosa_scenarios)

    _configure_footer(doc, program.name)
    return doc


def build_rfi(  # noqa: PLR0912
    program: Program,
    brief: ProgramBrief | None,
    answers: dict[str, str | None],
    files: list[ProgramFile],
    modules: list[Module],
    generated_date: str,
) -> DocxDocument:
    doc = DocxDocument()
    _setup_margins(doc)
    _add_title_page(doc, program, generated_date, "Request for Information (RFI)")
    doc.add_page_break()

    # ---- Section 1: Program Overview ----
    _h1(doc, "1. Program Overview")
    brief_desc = (brief.program_description or "").strip() if brief else ""
    wizard_desc = (answers.get("a_program_description") or "").strip()
    overview_text = brief_desc or wizard_desc or "No program description provided."
    doc.add_paragraph(overview_text)

    if brief:
        bullets: list[str] = []
        if brief.dev_cost_estimate is not None:
            bullets.append(f"Estimated development cost: {format_currency(brief.dev_cost_estimate)}")
        if brief.production_unit_cost is not None:
            bullets.append(f"Estimated production unit cost: {format_currency(brief.production_unit_cost)}")
        if brief.attritable:
            bullets.append("Attritable system")
        if brief.sustainment_tail:
            bullets.append("Has sustainment tail")
        if brief.software_large_part:
            bullets.append("Software is a large part of the program")
        if brief.mission_critical:
            bullets.append("Mission-critical")
        if brief.safety_critical:
            bullets.append("Safety-critical")
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")

    if wizard_desc and brief_desc and wizard_desc != brief_desc:
        doc.add_paragraph()
        doc.add_paragraph(wizard_desc)

    doc.add_page_break()

    # ---- Section 2: Constraints ----
    _h1(doc, "2. Constraints")
    constraints_answer = (answers.get("f_tech_challenges_and_risk_areas") or "").strip()
    if constraints_answer:
        doc.add_paragraph("Technical challenges and risk areas identified during program planning:")
        for line in constraints_answer.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("No specific constraints have been identified at this time.")

    if files:
        doc.add_paragraph()
        doc.add_paragraph(
            "The following reference documents have been uploaded as program constraints and background material:"
        )
        for f in files:
            doc.add_paragraph(f.filename, style="List Bullet")
    else:
        doc.add_paragraph("No reference documents have been attached.")

    doc.add_page_break()

    # ---- Section 3: Desired Modular Decomposition ----
    _h1(doc, "3. Desired Modular Decomposition")
    h_answer = (answers.get("h_candidate_modules") or "").strip()
    if h_answer:
        doc.add_paragraph(
            "The Government has identified the following candidate modules as a starting point for the "
            "modular architecture. Respondents are encouraged to propose alternative decompositions."
        )

    if modules:
        col_widths = [Inches(0.9), Inches(1.2), Inches(1.0), Inches(0.9), Inches(0.6), Inches(0.8), Inches(0.6)]
        tbl = doc.add_table(rows=len(modules) + 1, cols=7)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, text in enumerate([
            "Module", "Rationale", "Key Interfaces", "Standards",
            "Tech Risk", "Obsolescence Risk", "COTS Candidate",
        ]):
            _set_cell_text(hdr[i], text, bold=True)
            _shade_cell(hdr[i])
            hdr[i].width = col_widths[i]
        for r, mod in enumerate(modules, start=1):
            cells = tbl.rows[r].cells
            _set_cell_text(cells[0], safe_text(mod.name))
            _set_cell_text(cells[1], safe_text(mod.rationale))
            _set_cell_text(cells[2], safe_text(mod.key_interfaces))
            _set_cell_text(cells[3], safe_text(mod.standards))
            _set_cell_text(cells[4], "Yes" if mod.tech_risk else "No", center=True)
            _set_cell_text(cells[5], "Yes" if mod.obsolescence_risk else "No", center=True)
            _set_cell_text(cells[6], "Yes" if mod.cots_candidate else "No", center=True)
            for i, cell in enumerate(cells):
                cell.width = col_widths[i]
    else:
        doc.add_paragraph("No modules have been defined for this program yet.")

    doc.add_page_break()

    # ---- Section 4: MOSA Approach ----
    _h1(doc, "4. MOSA Approach")
    doc.add_paragraph(
        "This acquisition will be structured in accordance with a Modular Open Systems Approach (MOSA) "
        "as required by 10 U.S.C. \u00a7 4401. At minimum, the delivered system must satisfy the following criteria:"
    )
    for title, description in MOSA_MINIMUM_CRITERIA:
        p = doc.add_paragraph(style="List Bullet")
        run_title = p.add_run(f"{title}: ")
        run_title.bold = True
        p.add_run(description)

    mosa_scenarios = (answers.get("g_mosa_scenarios") or "").strip()
    if mosa_scenarios:
        doc.add_paragraph()
        doc.add_paragraph("Program-specific MOSA scenarios and considerations:")
        for line in mosa_scenarios.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")

    standards_answer = (answers.get("i_known_standards_architectures_mapping") or "").strip()
    if standards_answer:
        doc.add_paragraph()
        doc.add_paragraph("Known applicable standards and architectures by module:")
        doc.add_paragraph(standards_answer)

    doc.add_page_break()

    # ---- Section 5: Requested Responses ----
    _h1(doc, "5. Requested Responses")
    doc.add_paragraph(
        "Industry respondents are requested to provide the following information. "
        "Responses are for market research purposes only and do not constitute a commitment to procure."
    )
    requested_items = [
        "Company name, point of contact, and relevant experience with similar programs or systems.",
        "Proposed modular architecture concept, including identification of modules, interfaces, and applicable open standards.",
        "Assessment of the feasibility of the desired modular decomposition described in Section 3, with recommended modifications.",
        "List of relevant commercial-off-the-shelf (COTS) or modified-off-the-shelf (MOTS) solutions applicable to each candidate module.",
        "Identification of open interface standards and reference architectures (e.g., FACE, SOSA, VICTORY, CMOSS) applicable to this program.",
        "Technical risks and obsolescence concerns associated with the proposed approach and suggested mitigation strategies.",
        "Estimated development cost range and schedule for the proposed architecture.",
        "Any additional information relevant to the Government's modular open systems strategy for this program.",
    ]
    for item in requested_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(
        "Responses should be submitted in accordance with the instructions provided in the accompanying synopsis. "
        "Questions regarding this RFI may be directed to the contracting office."
    )

    _configure_footer(doc, program.name)
    return doc


# ---------------------------------------------------------------------------
# SEP section builders  (OSD SEP Guide v4.1)
# ---------------------------------------------------------------------------


def _sep_program_context(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
    modules: list[Module],
    *,
    include_cost: bool = True,
    include_flags: bool = True,
    include_modules: bool = True,
) -> None:
    """Append a short program-specific paragraph after a major section heading.

    Draws from brief.program_description, cost/schedule/sustainment/criticality
    wizard fields, and the module list (name + key interfaces).
    """
    parts: list[str] = []

    if brief and brief.program_description:
        desc = brief.program_description.strip()
        if len(desc) > 220:
            desc = desc[:220].rsplit(" ", 1)[0] + "\u2026"
        parts.append(desc)

    if include_cost and brief:
        cost_parts: list[str] = []
        if brief.dev_cost_estimate is not None:
            cost_parts.append(f"estimated development cost of {format_currency(brief.dev_cost_estimate)}")
        if brief.production_unit_cost is not None:
            cost_parts.append(f"production unit cost of {format_currency(brief.production_unit_cost)}")
        if cost_parts:
            parts.append("The program carries an " + " and ".join(cost_parts) + ".")

    if include_flags and brief:
        flags: list[str] = []
        if brief.sustainment_tail:
            flags.append("long sustainment tail")
        if brief.software_large_part:
            flags.append("software-intensive architecture")
        if brief.mission_critical:
            flags.append("mission-critical designation")
        if brief.safety_critical:
            flags.append("safety-critical designation")
        if brief.attritable:
            flags.append("attritable platform")
        if flags:
            parts.append("Key program characteristics: " + "; ".join(flags) + ".")

    if include_modules and modules:
        entries = [
            f"{m.name} (interfaces: {safe_text(m.key_interfaces)})" for m in modules[:4]
        ]
        trail = f" and {len(modules) - 4} additional module(s)" if len(modules) > 4 else ""
        parts.append("Defined modules: " + ", ".join(entries) + trail + ".")

    if parts:
        doc.add_paragraph(" ".join(parts))


def _sep_introduction(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
    answers: dict[str, str | None],
    files: list[ProgramFile],
    modules: list[Module],
) -> None:
    _h1(doc, "1. Introduction")
    doc.add_paragraph(
        f"This Systems Engineering Plan (SEP) describes the systems engineering (SE) approach, "
        f"processes, and activities that will be employed to develop and deliver the {safe_text(program.name)} program. "
        "It is prepared in accordance with the DoD Systems Engineering Plan Preparation Guide v4.1 "
        "and supports program execution per DoDI 5000.02."
    )
    _sep_program_context(doc, program, brief, modules, include_modules=True)

    _h2(doc, "1.1 Purpose and Scope")
    doc.add_paragraph(
        "This SEP establishes the framework for technical planning, risk management, interface control, "
        "requirements management, verification and validation, and implementation of a Modular Open Systems "
        "Approach (MOSA) throughout the program lifecycle. It is a living document and will be updated at "
        "each major program milestone."
    )

    _h2(doc, "1.2 Program Overview")
    brief_desc = (brief.program_description or "").strip() if brief else ""
    wizard_desc = (answers.get("a_program_description") or "").strip()
    prog_desc = brief_desc or wizard_desc
    if prog_desc:
        doc.add_paragraph(prog_desc)
    else:
        doc.add_paragraph(
            f"The {safe_text(program.name)} program description has not yet been provided."
        )

    _h2(doc, "1.3 Reference Documents")
    if files:
        doc.add_paragraph(
            "The following reference documents have been identified as applicable to this program:"
        )
        for f in files:
            doc.add_paragraph(f.filename, style="List Bullet")
    else:
        doc.add_paragraph(
            "No reference documents have been attached. Reference documents will be identified and "
            "incorporated as the program matures."
        )


def _sep_technical_baseline(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
    answers: dict[str, str | None],
    modules: list[Module],
) -> None:
    _h1(doc, "2. Technical Baseline and Program Description")
    _sep_program_context(doc, program, brief, modules, include_cost=True, include_flags=True, include_modules=False)

    _h2(doc, "2.1 System Description")
    brief_desc = (brief.program_description or "").strip() if brief else ""
    wizard_desc = (answers.get("a_program_description") or "").strip()
    prog_desc = (
        brief_desc or wizard_desc
        or f"The {safe_text(program.name)} system description is to be determined."
    )
    doc.add_paragraph(prog_desc)

    similar = (answers.get("e_similar_previous_programs") or "").strip()
    if similar:
        doc.add_paragraph("Similar or heritage programs identified for reference:")
        doc.add_paragraph(similar)

    _h2(doc, "2.2 Operational Concept (CONOPS)")
    doc.add_paragraph(
        f"The {safe_text(program.name)} system will be employed in accordance with operational requirements "
        "established by the warfighter. The Concept of Operations (CONOPS) will be developed in coordination "
        "with the user community and will define operational scenarios, mission threads, and system usage patterns."
    )
    if brief:
        notes: list[str] = []
        if brief.attritable:
            notes.append("The system is designed as an attritable (low-cost, expendable) platform.")
        if brief.sustainment_tail:
            notes.append(
                "A long sustainment tail is anticipated; lifecycle supportability is a key design driver."
            )
        for note in notes:
            doc.add_paragraph(note)

    _h2(doc, "2.3 Technical Baseline Summary")
    doc.add_paragraph(
        "The following program parameters represent the current technical baseline for planning purposes:"
    )
    bullets: list[str] = [f"Program Name: {safe_text(program.name)}"]
    if brief:
        if brief.dev_cost_estimate is not None:
            bullets.append(
                f"Estimated Development Cost: {format_currency(brief.dev_cost_estimate)}"
            )
        if brief.production_unit_cost is not None:
            bullets.append(
                f"Estimated Production Unit Cost: {format_currency(brief.production_unit_cost)}"
            )
        if brief.mission_critical:
            bullets.append("Classification: Mission-Critical System")
        if brief.safety_critical:
            bullets.append("Classification: Safety-Critical System")
        if brief.software_large_part:
            bullets.append("Software Intensity: Software constitutes a major portion of the system")
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")


def _sep_se_approach(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
    modules: list[Module],
    answers: dict[str, str | None],
) -> None:
    _h1(doc, "3. Systems Engineering Approach")
    doc.add_paragraph(
        "The program will employ a disciplined Systems Engineering (SE) process consistent with "
        "MIL-STD-499C and ISO/IEC/IEEE 15288. SE activities will be integrated with program management "
        "through the Integrated Product Team (IPT) structure and governed by the Integrated Master Plan (IMP) "
        "and Integrated Master Schedule (IMS)."
    )
    _sep_program_context(doc, program, brief, modules, include_cost=False, include_flags=True, include_modules=True)

    _h2(doc, "3.1 SE Process Overview")
    doc.add_paragraph(
        "The SE process will follow a recursive and iterative technical management approach encompassing: "
        "stakeholder requirements definition, requirements analysis, architectural design, implementation, "
        "integration, verification, validation, and transition. SE processes will be tailored to program "
        "risk, complexity, and acquisition phase."
    )

    _h2(doc, "3.2 Technical Planning and Control")
    tech = (answers.get("f_tech_challenges_and_risk_areas") or "").strip()
    doc.add_paragraph(
        "Technical planning will be captured in the Systems Engineering Management Plan (SEMP), "
        "which will be developed and maintained by the prime contractor and reviewed by the Government SE team. "
        "Key technical planning activities include:"
    )
    for item in [
        "Development and maintenance of the IMS with SE milestones",
        "Technical Performance Measure (TPM) identification and tracking",
        "Engineering Change Proposal (ECP) review and disposition",
        "Technical review scheduling and entry/exit criteria enforcement",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    if tech:
        doc.add_paragraph("Program-specific technical challenges identified:")
        for line in tech.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")

    _h2(doc, "3.3 Requirements Management")
    doc.add_paragraph(
        "Requirements will be managed using a Government-approved requirements management tool. "
        "All system requirements will be traced from source documents through the functional and allocated "
        "baselines to test procedures. Requirements volatility will be tracked as a TPM."
    )

    _h2(doc, "3.4 Interface Management")
    i_answer = (answers.get("i_known_standards_architectures_mapping") or "").strip()
    doc.add_paragraph(
        "Interface management will be governed through Interface Control Documents (ICDs) and Interface "
        "Requirements Specifications (IRSs). The Interface Control Working Group (ICWG) will maintain all "
        "external and internal interface definitions."
    )
    if modules:
        doc.add_paragraph("Key interfaces identified at the module level:")
        col_widths = [Inches(1.8), Inches(4.2)]
        tbl = doc.add_table(rows=len(modules) + 1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, text in enumerate(["Module", "Key Interfaces"]):
            _set_cell_text(hdr[i], text, bold=True)
            _shade_cell(hdr[i])
            hdr[i].width = col_widths[i]
        for r, mod in enumerate(modules, start=1):
            cells = tbl.rows[r].cells
            _set_cell_text(cells[0], safe_text(mod.name))
            _set_cell_text(cells[1], safe_text(mod.key_interfaces))
            cells[0].width = col_widths[0]
            cells[1].width = col_widths[1]
    if i_answer:
        doc.add_paragraph()
        doc.add_paragraph("Known standards and architectures applicable to interfaces:")
        doc.add_paragraph(i_answer)

    _h2(doc, "3.5 Technical Reviews and Audits")
    doc.add_paragraph(
        "Technical reviews will be conducted at key program milestones to assess technical maturity "
        "and readiness to proceed. Standard reviews include:"
    )
    reviews = [
        ("SRR \u2013 System Requirements Review",
         "Confirm completeness and adequacy of system requirements"),
        ("SDR \u2013 System Design Review",
         "Assess initial architectural design and requirements allocation"),
        ("PDR \u2013 Preliminary Design Review",
         "Assess preliminary design against requirements and identify risks"),
        ("CDR \u2013 Critical Design Review",
         "Confirm detailed design is ready for build/code-to"),
        ("SVR \u2013 System Verification Review",
         "Confirm system meets requirements prior to operational test"),
        ("FCA/PCA \u2013 Functional/Physical Configuration Audit",
         "Confirm as-built system matches design documentation"),
    ]
    col_widths = [Inches(2.0), Inches(4.0)]
    tbl = doc.add_table(rows=len(reviews) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Review", "Purpose"]):
        _set_cell_text(hdr[i], text, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (review, purpose) in enumerate(reviews, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], review)
        _set_cell_text(cells[1], purpose)
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]


def _sep_risk_management(
    doc: DocxDocument,
    program: Program,
    modules: list[Module],
    answers: dict[str, str | None],
    brief: ProgramBrief | None,
) -> None:
    _h1(doc, "4. Technical Risk Management")
    doc.add_paragraph(
        "The program will implement a continuous risk management process consistent with the DoD Risk, "
        "Issue, and Opportunity (RIO) Management Guide. Risks will be identified, assessed, mitigated, "
        "and tracked through a formal Risk Register updated at each program review."
    )
    _sep_program_context(doc, program, brief, modules, include_cost=False, include_flags=True, include_modules=False)

    _h2(doc, "4.1 Risk Identification and Assessment")
    tech = (answers.get("f_tech_challenges_and_risk_areas") or "").strip()
    risks: list[tuple[str, str, str, str]] = []
    if tech:
        for line in tech.splitlines():
            line = line.strip()
            if line:
                risks.append((line[:80], "TBD", "TBD", "Modular design limits blast radius"))
    for mod in modules:
        if mod.tech_risk:
            risks.append((
                f"Technical: {safe_text(mod.name)[:50]}",
                "High", "High",
                f"Early prototype; modular boundary isolates risk to {safe_text(mod.name)[:30]}",
            ))
        if mod.obsolescence_risk:
            risks.append((
                f"Obsolescence: {safe_text(mod.name)[:45]}",
                "Medium", "Medium",
                "COTS insertion plan; modular replacement path in LCMP",
            ))
    if risks:
        col_widths = [Inches(2.0), Inches(0.8), Inches(0.8), Inches(2.4)]
        tbl = doc.add_table(rows=len(risks) + 1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, text in enumerate(
            ["Risk Description", "Likelihood", "Impact", "Initial Mitigation"]
        ):
            _set_cell_text(hdr[i], text, bold=True)
            _shade_cell(hdr[i])
            hdr[i].width = col_widths[i]
        for r, (desc, likelihood, impact, mitigation) in enumerate(risks, start=1):
            cells = tbl.rows[r].cells
            _set_cell_text(cells[0], desc)
            _set_cell_text(cells[1], likelihood, center=True)
            _set_cell_text(cells[2], impact, center=True)
            _set_cell_text(cells[3], mitigation)
            for i, cell in enumerate(cells):
                cell.width = col_widths[i]
    else:
        doc.add_paragraph(
            "Program-specific technical risks will be identified during system requirements analysis."
        )

    _h2(doc, "4.2 Mitigation Strategies")
    doc.add_paragraph(
        "Risk mitigation will leverage the modular architecture to limit the blast radius of individual "
        "technical failures. Standard mitigation approaches include:"
    )
    for item in [
        "Modular design enabling component-level risk containment",
        "Early prototyping and incremental demonstration milestones for high-risk modules",
        "Technology Readiness Level (TRL) tracking with exit criteria at each milestone",
        "Independent Technical Risk Assessments (ITRAs) at PDR and CDR",
        "Use of COTS/MOTS solutions to reduce development risk where possible",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _h2(doc, "4.3 Technology Readiness Assessment")
    doc.add_paragraph(
        "A Technology Readiness Assessment (TRA) will be conducted at Milestone B to evaluate the "
        "maturity of critical technologies. Critical Technology Elements (CTEs) will be identified "
        "and tracked against TRL exit criteria throughout the Engineering and Manufacturing Development (EMD) phase."
    )
    if brief and brief.attritable:
        doc.add_paragraph(
            "Given the attritable nature of this system, TRL thresholds may be tailored to balance "
            "risk with the reduced lifecycle cost and expendability of the platform."
        )


def _sep_config_management(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
    modules: list[Module],
    answers: dict[str, str | None],
    files: list[ProgramFile],
) -> None:
    _h1(doc, "5. Configuration Management and Technical Data")
    _sep_program_context(doc, program, brief, modules, include_cost=False, include_flags=False, include_modules=True)

    _h2(doc, "5.1 Configuration Management Approach")
    doc.add_paragraph(
        "Configuration management will be performed in accordance with MIL-HDBK-61B and EIA-649. "
        "The Configuration Control Board (CCB) will review and disposition all Engineering Change Proposals "
        "(ECPs) and Requests for Deviation (RFDs). Three technical baselines will be established and maintained:"
    )
    for name, desc in [
        ("Functional Baseline (FBL)",
         "Established at SRR/SDR; defines system-level functional requirements"),
        ("Allocated Baseline (ABL)",
         "Established at PDR; allocates requirements to system elements"),
        ("Product Baseline (PBL)",
         "Established at FCA/PCA; defines the as-built system configuration"),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    _h2(doc, "5.2 Technical Data Management")
    doc.add_paragraph(
        "Technical data will be managed in a Government-accessible data management system. "
        "Contractor deliverable data items (CDRLs) will be specified in the contract and reviewed "
        "for completeness and technical adequacy at each delivery."
    )
    if files:
        doc.add_paragraph("Program reference documents identified:")
        for f in files:
            doc.add_paragraph(f.filename, style="List Bullet")

    _h2(doc, "5.3 Data Rights")
    i_answer = (answers.get("i_known_standards_architectures_mapping") or "").strip()
    doc.add_paragraph(
        "The Government will negotiate for sufficient technical data rights to enable independent "
        "support, competitive re-sourcing, and future modernization of all program elements. "
        "Minimum data rights posture:"
    )
    rights = [
        ("Interface Control Documents",
         "Unlimited rights \u2014 required for MOSA compliance and module re-competition"),
        ("Software Source Code (Government-funded)", "Unlimited rights"),
        ("Software Source Code (contractor-funded)", "Government Purpose Rights (GPR) minimum"),
        ("Design Data",
         "Government Purpose Rights (GPR) minimum; unlimited rights sought where feasible"),
        ("Firmware and Embedded Software",
         "Rights sufficient to allow security patching and sustainment"),
    ]
    col_widths = [Inches(2.2), Inches(3.8)]
    tbl = doc.add_table(rows=len(rights) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Data Category", "Minimum Rights"]):
        _set_cell_text(hdr[i], text, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (category, rights_level) in enumerate(rights, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], category)
        _set_cell_text(cells[1], rights_level)
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]
    if i_answer:
        doc.add_paragraph()
        doc.add_paragraph("Known ICDs and interface standards by module:")
        doc.add_paragraph(i_answer)


def _sep_mosa(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
    modules: list[Module],
    answers: dict[str, str | None],
) -> None:
    _h1(doc, "6. Modular Open Systems Approach (MOSA)")
    doc.add_paragraph(
        "This program implements a Modular Open Systems Approach (MOSA) as required by 10 U.S.C. \u00a7\u00a04401 "
        "and consistent with DoD Directive 5000.01. MOSA is integrated throughout the SE process to ensure "
        "the system is designed for competitive upgrade, independent sustainment, and lifecycle affordability."
    )
    _sep_program_context(doc, program, brief, modules, include_cost=False, include_flags=False, include_modules=True)

    _h2(doc, "6.1 MOSA Strategy Overview")
    mosa_scenarios = (answers.get("g_mosa_scenarios") or "").strip()
    o_answer = (answers.get("o_mosa_repo_searched") or "").strip()
    doc.add_paragraph(
        "The MOSA strategy is built on four minimum criteria that govern all design and acquisition decisions:"
    )
    for title, description in MOSA_MINIMUM_CRITERIA:
        p = doc.add_paragraph(style="List Bullet")
        run_title = p.add_run(f"{title}: ")
        run_title.bold = True
        p.add_run(description)
    if mosa_scenarios:
        doc.add_paragraph("Program-specific MOSA scenarios and considerations:")
        for line in mosa_scenarios.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")
    if o_answer:
        status = "has been completed" if o_answer.lower() == "yes" else "has not yet been completed"
        doc.add_paragraph(
            f"A search of the DoD MOSA repository {status} for applicable reference architectures "
            "and interface standards."
        )

    _h2(doc, "6.2 Module Identification and Boundaries")
    doc.add_paragraph(
        "Modules are identified through functional decomposition of the operational mission threads and "
        "top-level system requirements. Each module boundary is drawn where a well-defined, "
        "standards-based interface can be established and where independent competition or replacement "
        "provides lifecycle value. Module boundaries are validated against the following criteria: "
        "functional cohesion within the module, minimal coupling across the interface, availability of "
        "applicable open standards, and commercial or Government source availability."
    )
    h_answer = (answers.get("h_candidate_modules") or "").strip()
    if h_answer:
        doc.add_paragraph("Candidate modules identified during program planning:")
        doc.add_paragraph(h_answer)
    if modules:
        doc.add_paragraph("Formal module definitions:")
        col_widths = [Inches(1.2), Inches(1.6), Inches(1.2), Inches(0.7), Inches(0.7), Inches(0.6)]
        tbl = doc.add_table(rows=len(modules) + 1, cols=6)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, text in enumerate(
            ["Module", "Rationale / Boundary", "Key Interfaces", "Tech Risk", "Obs. Risk", "COTS"]
        ):
            _set_cell_text(hdr[i], text, bold=True)
            _shade_cell(hdr[i])
            hdr[i].width = col_widths[i]
        for r, mod in enumerate(modules, start=1):
            cells = tbl.rows[r].cells
            _set_cell_text(cells[0], safe_text(mod.name))
            _set_cell_text(cells[1], safe_text(mod.rationale))
            _set_cell_text(cells[2], safe_text(mod.key_interfaces))
            _set_cell_text(cells[3], "Yes" if mod.tech_risk else "No", center=True)
            _set_cell_text(cells[4], "Yes" if mod.obsolescence_risk else "No", center=True)
            _set_cell_text(cells[5], "Yes" if mod.cots_candidate else "No", center=True)
            for i, cell in enumerate(cells):
                cell.width = col_widths[i]
    else:
        doc.add_paragraph(
            "Module definitions will be developed during system requirements analysis and documented "
            "in the Interface Control Working Group (ICWG) charter."
        )

    _h2(doc, "6.3 Interface Standards and Control Documents")
    i_answer = (answers.get("i_known_standards_architectures_mapping") or "").strip()
    n_answer = (answers.get("n_software_standards_architectures") or "").strip()
    doc.add_paragraph(
        "Open standards and reference architectures are applied at two levels: (1) system-level "
        "reference architectures (e.g., CMOSS, SOSA, VICTORY, GVA) shape the overall decomposition "
        "and slot definitions; and (2) module-level standards govern specific interface protocols, "
        "data formats, and transport mechanisms. Standards selection follows a preference order of "
        "open consensus-based standards, Government-developed standards, and — only when no open "
        "standard exists — contractor-proposed standards subject to PM approval and ICD documentation."
    )
    doc.add_paragraph(
        "Interface Control Documents (ICDs) will be Government-owned and maintained under "
        "configuration control. All module interfaces will reference open, consensus-based standards "
        "where available. Contractor-proprietary interfaces are prohibited unless specifically "
        "approved by the PM."
    )
    doc.add_paragraph(
        "Technical data and data rights: The Government will acquire unlimited rights to all ICDs and "
        "interface specifications to enable competitive re-sourcing. Design data for Government-funded "
        "development will carry unlimited rights; contractor-funded components will carry at minimum "
        "Government Purpose Rights (GPR). These rights enable the Government to share data with "
        "competing vendors, conduct independent integration testing, and sustain the system without "
        "sole-source dependency."
    )
    if i_answer:
        doc.add_paragraph("Known applicable standards and architectures by module:")
        doc.add_paragraph(i_answer)
    if n_answer:
        doc.add_paragraph("Software-specific standards and architectures:")
        doc.add_paragraph(n_answer)
    with_standards = [m for m in modules if m.standards and m.standards.strip()]
    if with_standards:
        doc.add_paragraph("Module-level standards summary:")
        col_widths = [Inches(1.8), Inches(4.2)]
        tbl = doc.add_table(rows=len(with_standards) + 1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, text in enumerate(["Module", "Applicable Standards"]):
            _set_cell_text(hdr[i], text, bold=True)
            _shade_cell(hdr[i])
            hdr[i].width = col_widths[i]
        for r, mod in enumerate(with_standards, start=1):
            cells = tbl.rows[r].cells
            _set_cell_text(cells[0], safe_text(mod.name))
            _set_cell_text(cells[1], safe_text(mod.standards))
            cells[0].width = col_widths[0]
            cells[1].width = col_widths[1]

    _h2(doc, "6.4 Competitive Upgrade Path")
    doc.add_paragraph(
        "The modular architecture enables competitive re-sourcing of individual modules without system-level "
        "redesign. To preserve this capability:"
    )
    for item in [
        "All ICDs will be delivered with unlimited data rights to enable alternative-source competition",
        "Module acceptance testing will rely only on published ICDs \u2014 not proprietary contractor tools",
        "The Government will maintain a validated integration test environment to support re-competition",
        "Each module contract will include options for follow-on competitive procurement",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    k_answer = (answers.get("k_commercial_solutions_by_module") or "").strip()
    if k_answer:
        doc.add_paragraph("COTS/MOTS solutions identified by module:")
        doc.add_paragraph(k_answer)
    doc.add_paragraph(
        "Industry collaboration: The program office will engage industry throughout the acquisition "
        "lifecycle to validate module boundaries, refine interface standards, and assess commercial "
        "availability. Planned activities include pre-solicitation industry engagement, release of "
        "draft ICDs for industry comment, and participation in relevant standards body working groups "
        "(e.g., FACE Consortium, SOSA Consortium, VICTORY). Vendor-proposed open standards will be "
        "evaluated for adoption through the ICWG process. This collaboration ensures module definitions "
        "align with market realities and maximizes the competitive supplier base."
    )


def _sep_vv(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
    modules: list[Module],
    answers: dict[str, str | None],
) -> None:
    _h1(doc, "7. Verification and Validation")
    doc.add_paragraph(
        "Verification and Validation (V&V) activities confirm that the system is built correctly (meets "
        "requirements) and that the right system is being built (meets stakeholder needs). V&V will be "
        "planned and documented in the Test and Evaluation Master Plan (TEMP)."
    )
    _sep_program_context(doc, program, brief, modules, include_cost=False, include_flags=True, include_modules=False)

    _h2(doc, "7.1 Verification Strategy")
    doc.add_paragraph(
        "System requirements will be verified through a combination of the following methods:"
    )
    for method, desc in [
        ("Inspection",
         "Review of design documentation, drawings, and code for compliance with requirements"),
        ("Analysis",
         "Mathematical, simulation, or model-based demonstration of compliance"),
        ("Demonstration",
         "Observable functional demonstration without formal measurement"),
        ("Test",
         "Controlled measurement of system performance against quantitative acceptance criteria"),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{method}: ")
        run.bold = True
        p.add_run(desc)
    doc.add_paragraph(
        "A Verification Cross-Reference Matrix (VCRM) will trace each requirement to its verification "
        "method, responsible organization, and completion milestone."
    )

    _h2(doc, "7.2 Validation Approach")
    doc.add_paragraph(
        "System validation will confirm that the delivered system satisfies warfighter operational needs "
        "as expressed in the Operational Requirements Document (ORD) or Capability Development Document (CDD). "
        "Validation activities include:"
    )
    for item in [
        "Concept of Operations (CONOPS) review with operational users at each major milestone",
        "Operational utility evaluation during Initial Operational Test & Evaluation (IOT&E)",
        "Mission thread analysis to confirm end-to-end mission accomplishment",
        "User acceptance testing with representative operational users prior to fielding",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _h2(doc, "7.3 Test and Evaluation Summary")
    doc.add_paragraph(
        "Testing will be conducted in phases, progressing from component-level to system-level:"
    )
    test_phases = [
        ("Unit / Component Testing",
         "Contractor-conducted at module level; results reviewed by Government SE"),
        ("Integration Testing",
         "Verify module interfaces function correctly per published ICDs"),
        ("System Testing",
         "Government-witnessed contractor testing against all system-level requirements"),
        ("Developmental Testing (DT)",
         "Government-controlled testing to characterize system performance"),
        ("Operational Testing (OT)",
         "User-controlled testing in representative operational environment"),
    ]
    col_widths = [Inches(2.0), Inches(4.0)]
    tbl = doc.add_table(rows=len(test_phases) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Test Phase", "Description"]):
        _set_cell_text(hdr[i], text, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (phase, desc) in enumerate(test_phases, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], phase)
        _set_cell_text(cells[1], desc)
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]
    if brief:
        if brief.mission_critical:
            doc.add_paragraph(
                "Mission-critical designation requires formal DT/OT execution with "
                "documented acceptance criteria."
            )
        if brief.safety_critical:
            doc.add_paragraph(
                "Safety-critical designation requires independent safety review board "
                "approval prior to IOT&E."
            )


def _sep_specialty_engineering(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
    answers: dict[str, str | None],
    modules: list[Module],
) -> None:
    _h1(doc, "8. Specialty Engineering")
    doc.add_paragraph(
        "Specialty engineering disciplines are integrated throughout the SE process to ensure a balanced, "
        "lifecycle-supportable design. Specialty engineering requirements will be captured in the SEMP "
        "and tracked through the IPT structure."
    )
    _sep_program_context(doc, program, brief, modules, include_cost=False, include_flags=True, include_modules=False)

    _h2(doc, "8.1 Reliability, Availability, and Maintainability")
    doc.add_paragraph(
        "RAM requirements will be derived from operational availability (Ao) and Mean Time Between "
        "Failure (MTBF) thresholds established in the CDD/CPD. RAM analysis activities include:"
    )
    for item in [
        "Failure Mode, Effects, and Criticality Analysis (FMECA)",
        "Reliability growth testing with Duane model tracking",
        "Maintainability demonstration at CDR and SVR",
        "Availability modeling updated at each design review",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    if brief and brief.sustainment_tail:
        doc.add_paragraph(
            "Given the anticipated long sustainment tail, particular emphasis will be placed on "
            "field-level maintainability and modular component replaceability."
        )

    _h2(doc, "8.2 Software Systems Engineering")
    n_answer = (answers.get("n_software_standards_architectures") or "").strip()
    l_answer = (answers.get("l_software_large_part") or "").strip()
    doc.add_paragraph(
        "Software engineering will be conducted in accordance with MIL-STD-498 (tailored as appropriate) "
        "and applicable service-specific software standards. Software development will follow an Agile or "
        "spiral methodology with formal Government checkpoint reviews."
    )
    if l_answer and l_answer.lower() in ("yes", "true", "1"):
        doc.add_paragraph(
            "Software is identified as a major portion of this system. Elevated SE rigor will be applied "
            "including formal Software Requirements Reviews (SRRs), Software Design Reviews (SDRs), "
            "and Software Test Reviews (STRs)."
        )
    if n_answer:
        doc.add_paragraph("Applicable software standards and architectures:")
        doc.add_paragraph(n_answer)
    sw_modules = [m for m in modules if m.name and "software" in m.name.lower()]
    if sw_modules:
        doc.add_paragraph("Software-intensive modules:")
        for mod in sw_modules:
            doc.add_paragraph(safe_text(mod.name), style="List Bullet")

    _h2(doc, "8.3 Human Systems Integration")
    doc.add_paragraph(
        "Human Systems Integration (HSI) will be conducted in accordance with MIL-STD-1472H "
        "and applicable HSI guidance. HSI domains addressed include:"
    )
    for item in [
        "Human factors engineering (HFE) \u2014 operator interface design and workload analysis",
        "Manpower \u2014 personnel and staffing requirements for operation and maintenance",
        "Training \u2014 operator and maintainer training system development",
        "Personnel \u2014 minimum qualification requirements for operators and maintainers",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _h2(doc, "8.4 Cybersecurity / Risk Management Framework")
    m_answer = (answers.get("m_mission_or_safety_critical") or "").strip()
    doc.add_paragraph(
        "Cybersecurity will be addressed through the DoD Risk Management Framework (RMF) in accordance "
        "with DoDI 8510.01. A Program Protection Plan (PPP) will be developed and maintained throughout "
        "the acquisition lifecycle. Key cybersecurity activities include:"
    )
    for item in [
        "System categorization and security control selection per CNSSI 1253",
        "Threat model development and adversarial risk assessment",
        "Cybersecurity requirements derivation and allocation to system elements",
        "Penetration testing at CDR and prior to IOT&E",
        "Software assurance including static analysis and fuzz testing",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    if m_answer and m_answer.lower() in ("mission critical", "yes", "true", "mission"):
        doc.add_paragraph(
            "Mission-critical designation requires elevated cybersecurity scrutiny and an approved "
            "Authority to Operate (ATO) prior to operational use."
        )

    _h2(doc, "8.5 Environment, Safety, and Occupational Health")
    doc.add_paragraph(
        "ESOH will be addressed in accordance with DoDI 4715.02. Key ESOH activities include:"
    )
    for item in [
        "ESOH risk assessment integrated into the program risk register",
        "Hazard Analysis (HA) and System Safety Analysis performed at each design phase",
        "Compliance with applicable environmental regulations (e.g., RoHS, REACH) for all components",
        "Occupational health and safety assessments for operator and maintainer populations",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    if brief and brief.safety_critical:
        doc.add_paragraph(
            "Safety-critical designation requires a System Safety Working Group (SSWG) and an approved "
            "System Safety Program Plan (SSPP) prior to PDR."
        )


def _sep_se_organization(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
    modules: list[Module],
) -> None:
    _h1(doc, "9. SE Organization and Resources")
    _sep_program_context(doc, program, brief, modules, include_cost=True, include_flags=False, include_modules=False)

    _h2(doc, "9.1 Government SE Organization")
    doc.add_paragraph(
        f"The Government SE team for {safe_text(program.name)} will be led by the Lead Systems Engineer (LSE) "
        "reporting to the Program Manager (PM). The LSE is responsible for SE process oversight, "
        "technical review execution, and MOSA compliance verification. Functional IPTs will support "
        "each major technical area."
    )

    _h2(doc, "9.2 Contractor SE Organization")
    doc.add_paragraph(
        "The prime contractor will maintain a Systems Engineering organization with a Chief Engineer "
        "responsible for technical direction and SE process execution. The contractor SE organization will "
        "mirror the Government IPT structure and participate in all Government-led technical reviews. "
        "Sub-contractor SE activities will be governed by the prime contractor."
    )

    _h2(doc, "9.3 SE Tools and Infrastructure")
    doc.add_paragraph(
        "The following SE tools and infrastructure will be used to support program execution:"
    )
    tools = [
        ("Requirements Management",
         "Government-approved RM tool (e.g., DOORS, Jama, Polarion)"),
        ("Model-Based Systems Engineering (MBSE)",
         "SysML-compliant modeling tool (e.g., Cameo, Rhapsody, Capella)"),
        ("Interface Control",
         "ICD repository accessible to all contractors and Government SE team"),
        ("Configuration Management",
         "Government-controlled CM database with automated change tracking"),
        ("Risk Management",
         "Program Risk Register maintained in a Government-accessible tool"),
        ("Test Management",
         "Linked to requirements for automated VCRM generation"),
    ]
    col_widths = [Inches(1.8), Inches(4.2)]
    tbl = doc.add_table(rows=len(tools) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["SE Function", "Tool / Approach"]):
        _set_cell_text(hdr[i], text, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (func, tool) in enumerate(tools, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], func)
        _set_cell_text(cells[1], tool)
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]


def _sep_summary(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
    modules: list[Module],
    answers: dict[str, str | None],
) -> None:
    _h1(doc, "10. Summary and Appendices")
    _sep_program_context(doc, program, brief, modules, include_cost=True, include_flags=True, include_modules=False)

    _h2(doc, "10.1 Key Assumptions and Constraints")
    doc.add_paragraph(
        "The following assumptions and constraints underpin this SEP. Changes to these assumptions "
        "will trigger a SEP update and may require re-baselining of associated program plans."
    )
    tech = (answers.get("f_tech_challenges_and_risk_areas") or "").strip()
    assumptions: list[str] = [
        "MOSA compliance will be verified at each major milestone as a condition of proceeding",
        "Government will maintain a validated integration test environment for module re-competition",
        "All ICDs will be delivered with unlimited Government data rights",
        "Contractor SEMP will be consistent with and subordinate to this SEP",
    ]
    if brief:
        if brief.dev_cost_estimate is not None:
            assumptions.append(
                f"Program development cost baseline: {format_currency(brief.dev_cost_estimate)}"
            )
        if brief.production_unit_cost is not None:
            assumptions.append(
                f"Production unit cost target: {format_currency(brief.production_unit_cost)}"
            )
    if tech:
        for line in tech.splitlines():
            line = line.strip()
            if line:
                assumptions.append(f"Known technical challenge: {line[:120]}")
    for a in assumptions:
        doc.add_paragraph(a, style="List Bullet")

    _h2(doc, "10.2 Acronyms and Abbreviations")
    acronyms = [
        ("ABL", "Allocated Baseline"),
        ("ATO", "Authority to Operate"),
        ("CCB", "Configuration Control Board"),
        ("CDD", "Capability Development Document"),
        ("CDR", "Critical Design Review"),
        ("CONOPS", "Concept of Operations"),
        ("COTS", "Commercial Off-the-Shelf"),
        ("CPD", "Capability Production Document"),
        ("CTE", "Critical Technology Element"),
        ("DT", "Developmental Testing"),
        ("ECP", "Engineering Change Proposal"),
        ("ESOH", "Environment, Safety, and Occupational Health"),
        ("FBL", "Functional Baseline"),
        ("FCA", "Functional Configuration Audit"),
        ("FMECA", "Failure Mode, Effects, and Criticality Analysis"),
        ("GPR", "Government Purpose Rights"),
        ("HSI", "Human Systems Integration"),
        ("ICD", "Interface Control Document"),
        ("ICWG", "Interface Control Working Group"),
        ("IMP", "Integrated Master Plan"),
        ("IMS", "Integrated Master Schedule"),
        ("IOT&E", "Initial Operational Test and Evaluation"),
        ("IPT", "Integrated Product Team"),
        ("LSE", "Lead Systems Engineer"),
        ("MBSE", "Model-Based Systems Engineering"),
        ("MOSA", "Modular Open Systems Approach"),
        ("MOTS", "Modified Off-the-Shelf"),
        ("OT", "Operational Testing"),
        ("PBL", "Product Baseline"),
        ("PCA", "Physical Configuration Audit"),
        ("PDR", "Preliminary Design Review"),
        ("PM", "Program Manager"),
        ("PPP", "Program Protection Plan"),
        ("RAM", "Reliability, Availability, and Maintainability"),
        ("RFD", "Request for Deviation"),
        ("RMF", "Risk Management Framework"),
        ("SDR", "System Design Review"),
        ("SE", "Systems Engineering"),
        ("SEMP", "Systems Engineering Management Plan"),
        ("SEP", "Systems Engineering Plan"),
        ("SRR", "System Requirements Review"),
        ("SSPP", "System Safety Program Plan"),
        ("SSWG", "System Safety Working Group"),
        ("SVR", "System Verification Review"),
        ("TEMP", "Test and Evaluation Master Plan"),
        ("TPM", "Technical Performance Measure"),
        ("TRA", "Technology Readiness Assessment"),
        ("TRL", "Technology Readiness Level"),
        ("VCRM", "Verification Cross-Reference Matrix"),
    ]
    col_widths = [Inches(1.2), Inches(4.8)]
    tbl = doc.add_table(rows=len(acronyms) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(["Acronym", "Definition"]):
        _set_cell_text(hdr[i], text, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (acr, defn) in enumerate(acronyms, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], acr, bold=True)
        _set_cell_text(cells[1], defn)
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]


# ---------------------------------------------------------------------------
# MOSA CONFORMANCE PLAN section builders
# ---------------------------------------------------------------------------


def _mosa_plan_program_overview(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
    answers: dict[str, str | None],
    files: list[ProgramFile],
) -> None:
    _h1(doc, "3. Program Overview & Reference Documents")
    doc.add_paragraph(
        f"This MOSA Conformance Plan documents the {safe_text(program.name)} program's approach to "
        "achieving and sustaining Modular Open Systems Approach (MOSA) compliance in accordance with "
        "10 U.S.C. \u00a7\u00a04401 and DoD Directive 5000.01. It establishes the modular architecture baseline, "
        "interface control strategy, technical data rights posture, and verification plan against which "
        "MOSA compliance will be assessed at each program milestone."
    )

    brief_desc = (brief.program_description or "").strip() if brief else ""
    wizard_desc = (answers.get("a_program_description") or "").strip()
    prog_desc = brief_desc or wizard_desc
    if prog_desc:
        doc.add_paragraph(prog_desc)

    bullets: list[str] = [f"Program: {safe_text(program.name)}"]
    if brief:
        if brief.dev_cost_estimate is not None:
            bullets.append(f"Estimated development cost: {format_currency(brief.dev_cost_estimate)}")
        if brief.production_unit_cost is not None:
            bullets.append(f"Estimated production unit cost: {format_currency(brief.production_unit_cost)}")
        if brief.mission_critical:
            bullets.append("Mission-critical system")
        if brief.safety_critical:
            bullets.append("Safety-critical system")
        if brief.attritable:
            bullets.append("Attritable system (low-cost, expendable)")
        if brief.sustainment_tail:
            bullets.append("Long sustainment tail anticipated")
        if brief.software_large_part:
            bullets.append("Software constitutes a major portion of the system")
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    if files:
        doc.add_paragraph()
        doc.add_paragraph("Reference documents applicable to this program:")
        for f in files:
            doc.add_paragraph(f.filename, style="List Bullet")


def _mosa_plan_modular_architecture(
    doc: DocxDocument,
    modules: list[Module],
    answers: dict[str, str | None],
) -> None:
    _h1(doc, "3.1 Module Inventory / Decomposition")
    doc.add_paragraph(
        "The following table summarizes the modular decomposition of the system. Each module "
        "represents a cohesive functional element with well-defined boundaries and standardized "
        "interfaces, enabling independent development, integration, and replacement by competing suppliers."
    )

    h_answer = (answers.get("h_candidate_modules") or "").strip()
    if h_answer:
        doc.add_paragraph("Candidate modules identified through program planning:")
        doc.add_paragraph(h_answer)

    headers = [
        "Module", "Functional Boundary", "Key Interfaces",
        "Standards", "Tech Risk", "Obsolescence Risk",
    ]
    col_widths = [Inches(1.0), Inches(1.2), Inches(1.2), Inches(1.2), Inches(0.7), Inches(0.7)]
    tbl = doc.add_table(rows=len(modules) + 1, cols=6)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(headers):
        _set_cell_text(hdr[i], text, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, mod in enumerate(modules, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], safe_text(mod.name))
        _set_cell_text(cells[1], safe_text(mod.rationale))
        _set_cell_text(cells[2], safe_text(mod.key_interfaces))
        _set_cell_text(cells[3], safe_text(mod.standards))
        _set_cell_text(cells[4], "High" if mod.tech_risk else "Low", center=True)
        _set_cell_text(cells[5], "High" if mod.obsolescence_risk else "Low", center=True)
        for i, cell in enumerate(cells):
            cell.width = col_widths[i]

    if not modules:
        doc.add_paragraph(
            "No modules have been defined for this program. Module definitions will be "
            "developed during system requirements analysis and documented in the ICWG charter."
        )


def _mosa_plan_interface_control(
    doc: DocxDocument,
    modules: list[Module],
    answers: dict[str, str | None],
) -> None:
    _h1(doc, "4. Interface Registry")
    doc.add_paragraph(
        "Interface Control Documents (ICDs) are Government-owned artifacts that define the technical "
        "standards and protocols governing communication between modules. The following registry identifies "
        "each interface by type (Hardware, Software, or Data), its owning module, and the applicable open "
        "standards and data rights posture. All ICDs will be maintained under configuration control by the "
        "Interface Control Working Group (ICWG). Per PEO AVN MIG §9.1, the Government will require delivery "
        "of Modular System Interface (MSI) documentation and will pursue Government Purpose Rights (GPR) "
        "or better on all interface definitions. Contractor-proprietary interfaces are prohibited unless "
        "specifically approved by the Program Manager."
    )

    i_answer = (answers.get("i_known_standards_architectures_mapping") or "").strip()
    if i_answer:
        doc.add_paragraph("Known interface standards and architecture mappings identified during program planning:")
        doc.add_paragraph(i_answer)

    # Assign interface types based on context clues in the interface text.
    def _infer_type(iface_name: str, mod_name: str) -> str:
        name_lower = iface_name.lower() + " " + mod_name.lower()
        if any(k in name_lower for k in ("power", "electrical", "connector", "mechanical", "emi", "thermal")):
            return "Hardware"
        if any(k in name_lower for k in ("data", "message", "bus", "network", "link", "comms", "communication",
                                          "serial", "ethernet", "can ", "1553", "arinc")):
            return "Data"
        return "Software"

    # Build one row per interface extracted from each module's key_interfaces field.
    # Fall back to a single row per module if no interface text is present.
    # Tuple: (iface_name, interface_type, owning_module, open_std, icd_req, rights)
    interface_rows: list[tuple[str, str, str, str, str, str]] = []
    for mod in modules:
        iface_text = (mod.key_interfaces or "").strip()
        has_standards = bool((mod.standards or "").strip())
        open_std = "Yes" if has_standards else "TBD"
        if iface_text:
            raw_ifaces = [
                x.strip()
                for x in iface_text.replace(";", "\n").replace(",", "\n").splitlines()
                if x.strip()
            ]
            for iface in raw_ifaces[:3]:  # cap at 3 per module
                iface_type = _infer_type(iface, mod.name or "")
                interface_rows.append((
                    iface[:80],
                    iface_type,
                    safe_text(mod.name),
                    open_std,
                    "Yes",
                    "Unlimited Rights",
                ))
        else:
            iface_type = _infer_type("", mod.name or "")
            interface_rows.append((
                f"{safe_text(mod.name)} External Interface",
                iface_type,
                safe_text(mod.name),
                open_std,
                "Yes",
                "Unlimited Rights",
            ))

    if not interface_rows:
        interface_rows = [
            ("External Command & Control", "Software", "TBD", "TBD", "Yes", "Unlimited Rights"),
            ("Data Link / Communications", "Data", "TBD", "TBD", "Yes", "Unlimited Rights"),
            ("Power Distribution", "Hardware", "TBD", "TBD", "Yes", "Unlimited Rights"),
        ]

    headers = [
        "Interface Name", "Type\n(HW/SW/Data)", "Owning Module",
        "Open Standard?", "ICD Required?", "Data Rights Required",
    ]
    col_widths = [Inches(1.3), Inches(0.7), Inches(1.1), Inches(0.8), Inches(0.8), Inches(1.3)]
    tbl = doc.add_table(rows=len(interface_rows) + 1, cols=6)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(headers):
        _set_cell_text(hdr[i], text, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (iface, itype, owner, open_std, icd_req, rights) in enumerate(interface_rows, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], iface)
        _set_cell_text(cells[1], itype, center=True)
        _set_cell_text(cells[2], owner)
        _set_cell_text(cells[3], open_std, center=True)
        _set_cell_text(cells[4], icd_req, center=True)
        _set_cell_text(cells[5], rights)
        for i, cell in enumerate(cells):
            cell.width = col_widths[i]


def _mosa_plan_data_rights(
    doc: DocxDocument,
    brief: ProgramBrief | None,
    modules: list[Module],
    answers: dict[str, str | None],
) -> None:
    _h1(doc, "6. Technical Data & Data Rights Strategy")
    doc.add_paragraph(
        "The Government will acquire sufficient technical data rights to enable independent support, "
        "competitive re-sourcing, and future modernization across all program elements. The governing "
        "principle is that the Government must retain the ability to re-compete any module using only "
        "Government-owned technical data, without reliance on the original developer. All contractor "
        "data rights assertions will be reviewed and negotiated prior to contract award."
    )

    if brief and brief.software_large_part:
        doc.add_paragraph(
            "Given the software-intensive nature of this program, particular emphasis is placed on "
            "acquiring source code rights sufficient to enable independent maintenance, security patching, "
            "and competitive re-development of software-intensive modules."
        )

    i_answer = (answers.get("i_known_standards_architectures_mapping") or "").strip()
    if i_answer:
        doc.add_paragraph("Interface standards and architecture mapping from program planning:")
        doc.add_paragraph(i_answer)

    artifacts: list[tuple[str, str, str]] = [
        (
            "Interface Control Documents (ICDs)",
            "Unlimited Rights",
            "Required for competitive module re-sourcing; must be available to any potential offeror",
        ),
        (
            "Software Source Code (Government-funded)",
            "Unlimited Rights",
            "Government-funded development yields unlimited rights per DFARS 252.227-7014",
        ),
        (
            "Software Source Code (contractor-funded)",
            "Government Purpose Rights (GPR) minimum",
            "Allows Government to use, modify, and share with other contractors for Government purposes",
        ),
        (
            "Design Data & Drawings",
            "Government Purpose Rights (GPR) minimum",
            "Enables independent sustainment and future competitive procurement",
        ),
        (
            "Firmware & Embedded Software",
            "Rights sufficient for sustainment",
            "Minimum rights to allow security patching, upgrades, and maintainer access",
        ),
        (
            "Test Procedures & Acceptance Criteria",
            "Unlimited Rights",
            "Required to conduct independent module acceptance testing during re-competition",
        ),
        (
            "Open Standards Documentation",
            "Not applicable (public domain)",
            "Open, consensus-based standards are publicly available; copies retained by Government",
        ),
    ]

    for mod in modules:
        if mod.cots_candidate:
            artifacts.append((
                f"COTS License: {safe_text(mod.name)[:40]}",
                "Perpetual, transferable license",
                f"{safe_text(mod.name)[:40]} is a COTS candidate; license terms must support competitive refresh",
            ))

    headers = ["Artifact", "Rights Required", "Rationale"]
    col_widths = [Inches(1.8), Inches(1.5), Inches(2.7)]
    tbl = doc.add_table(rows=len(artifacts) + 1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(headers):
        _set_cell_text(hdr[i], text, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (artifact, rights, rationale) in enumerate(artifacts, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], artifact)
        _set_cell_text(cells[1], rights)
        _set_cell_text(cells[2], rationale)
        for i, cell in enumerate(cells):
            cell.width = col_widths[i]


def _mosa_plan_verification(
    doc: DocxDocument,
    modules: list[Module],
    answers: dict[str, str | None],
) -> None:
    _h1(doc, "7. Conformance Verification Plan")
    doc.add_paragraph(
        "MOSA compliance will be verified at each major program milestone through review of compliance "
        "artifacts and demonstration of modular interoperability. Failure to satisfy MOSA compliance "
        "criteria at any milestone is a condition of the program not proceeding to the next phase. "
        "The program Lead Systems Engineer (LSE) is responsible for compliance verification."
    )

    milestones: list[tuple[str, str, str]] = [
        (
            "System Requirements Review (SRR)",
            "MOSA Requirements Compliance Matrix; initial module decomposition; candidate ICD list",
            "Government review of compliance matrix; MOSA checklist scoring; module boundary assessment",
        ),
        (
            "System Design Review (SDR)",
            "Draft ICDs for all external interfaces; open standards selection rationale; data rights plan",
            "ICWG review of draft ICDs; Government SE assessment of open standards compliance",
        ),
        (
            "Preliminary Design Review (PDR)",
            "Updated ICDs; Open Standards Compliance Matrix; TDP outline; data rights assertions",
            "Government review of ICDs and compliance matrix; legal review of data rights assertions",
        ),
        (
            "Critical Design Review (CDR)",
            "Final ICDs; completed Open Standards Compliance Matrix; Technical Data Package (TDP)",
            "Competitive demonstration: alternate-source module integrated using only published ICDs",
        ),
        (
            "System Verification Review (SVR)",
            "Verified compliance matrix; as-built ICD updates; final TDP delivery",
            "Government witness of MOSA integration testing; ICD accuracy verification vs. as-built system",
        ),
        (
            "Initial Operational Test & Evaluation (IOT&E)",
            "MOSA compliance certification; operational ICD library; data rights delivery complete",
            "Independent assessment of module replaceability; operational demonstration of open interface compliance",
        ),
    ]

    headers = ["Milestone", "Compliance Artifact", "Verification Method"]
    col_widths = [Inches(1.4), Inches(2.3), Inches(2.3)]
    tbl = doc.add_table(rows=len(milestones) + 1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(headers):
        _set_cell_text(hdr[i], text, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (milestone, artifact, method) in enumerate(milestones, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], milestone)
        _set_cell_text(cells[1], artifact)
        _set_cell_text(cells[2], method)
        for i, cell in enumerate(cells):
            cell.width = col_widths[i]


def _mosa_plan_scenario_mapping(
    doc: DocxDocument,
    modules: list[Module],
    answers: dict[str, str | None],
) -> None:
    _h1(doc, "10. MOSA Scenario Mapping")
    doc.add_paragraph(
        "The following table maps MOSA scenarios to the affected modules and the implementation strategy "
        "for each. MOSA scenarios represent anticipated lifecycle events — such as competitive re-sourcing, "
        "technology refresh, or capability upgrade — that the modular architecture must be designed to "
        "accommodate. Program-specific scenarios are listed first, followed by standard MOSA lifecycle scenarios."
    )

    mosa_scenarios_raw = (answers.get("g_mosa_scenarios") or "").strip()
    module_names_all = ", ".join(m.name for m in modules[:4]) if modules else "TBD"
    if len(modules) > 4:
        module_names_all += f" + {len(modules) - 4} more"

    tech_risk_mods = ", ".join(m.name for m in modules if m.tech_risk) or module_names_all
    obs_risk_mods = ", ".join(m.name for m in modules if m.obsolescence_risk) or module_names_all
    cots_mods = ", ".join(m.name for m in modules if m.cots_candidate) or module_names_all

    scenario_rows: list[tuple[str, str, str]] = []

    # Program-specific scenarios from wizard answer
    if mosa_scenarios_raw:
        for line in mosa_scenarios_raw.splitlines():
            line = line.strip()
            if not line:
                continue
            scenario_rows.append((
                line[:100],
                module_names_all,
                "Enforce open standards at module interfaces; ensure ICD data rights; "
                "validate replaceability at CDR",
            ))

    # Standard MOSA lifecycle scenarios
    standard_scenarios: list[tuple[str, str, str]] = [
        (
            "Competitive re-sourcing of a module at contract re-competition",
            module_names_all,
            "Maintain Government-owned ICDs with unlimited rights; conduct open competition using "
            "published interface specifications; validate alternate-source module at CDR",
        ),
        (
            "Technology refresh of an obsolete module without system redesign",
            obs_risk_mods,
            "Preserve interface standards across refresh boundary; update ICD and TDP; "
            "conduct module-level acceptance test using published ICDs",
        ),
        (
            "Integration of a COTS/MOTS replacement module",
            cots_mods,
            "Validate COTS module against published ICDs; acquire perpetual, transferable licenses; "
            "document deviations from open standards with PM approval",
        ),
        (
            "Independent Government sustainment without original contractor support",
            module_names_all,
            "Acquire TDP and source code with sufficient rights; maintain validated Government "
            "integration test environment; train Government engineers on interface standards",
        ),
        (
            "Rapid capability upgrade via plug-in module replacement",
            tech_risk_mods,
            "Enforce slot-based architecture with published slot definitions; require plug-in "
            "compliance demonstration at integration test; maintain ICD currency",
        ),
    ]
    scenario_rows.extend(standard_scenarios)

    headers = ["MOSA Scenario", "Affected Modules", "Implementation Strategy"]
    col_widths = [Inches(1.6), Inches(1.4), Inches(3.0)]
    tbl = doc.add_table(rows=len(scenario_rows) + 1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, text in enumerate(headers):
        _set_cell_text(hdr[i], text, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (scenario, affected, strategy) in enumerate(scenario_rows, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], scenario)
        _set_cell_text(cells[1], affected)
        _set_cell_text(cells[2], strategy)
        for i, cell in enumerate(cells):
            cell.width = col_widths[i]


def _mosa_plan_objectives(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
    answers: dict[str, str | None],
) -> None:
    """Section 1 – MOSA Objectives (statutory requirements + PEO AVN goals)."""
    _h1(doc, "1. MOSA Objectives")
    doc.add_paragraph(
        "Title 10 U.S. Code \u00a7 4401 defines a Modular Open Systems Approach (MOSA) as an integrated "
        "business and technical strategy that: employs a modular design using modular system interfaces; "
        "subjects those interfaces to verification; uses an architecture that allows severability; and "
        "complies with technical data rights guidance per 10 U.S.C. \u00a7 2320. The five MOSA Principles "
        "from the OSD Program Manager's Guide to Open Systems (v2.0) are:"
    )
    for principle in [
        "Establish an Enabling Environment \u2013 Governance, working groups, and shared architecture frameworks "
        "(e.g., PEO AVN Enterprise Architecture Framework) that sustain MOSA across the enterprise.",
        "Employ Modular Design \u2013 Cohesive, encapsulated, self-contained, and highly-bound modules with "
        "clearly delineated functional boundaries (MSCs/MSPs/MSIs per EAF).",
        "Designate Key Interfaces \u2013 Identify, document, and control Modular System Interfaces (MSIs) that "
        "are candidates for open-standards compliance and Government ownership.",
        "Select Open Standards \u2013 Choose consensus-based, widely supported standards (FACE, CMOSS/AMCE, "
        "SOSA, VICTORY, etc.) to prevent lock-in and enable competitive re-sourcing.",
        "Certify Conformance \u2013 Verify and certify that delivered components comply with the designated "
        "interface standards and the program\u2019s tailored MOSA at each lifecycle milestone.",
    ]:
        doc.add_paragraph(principle, style="List Bullet")

    _h2(doc, "1.1 PEO AVN MOSA Objectives for This Program")
    doc.add_paragraph(
        f"The {safe_text(program.name)} program applies MOSA in pursuit of the following PEO Aviation "
        "enterprise objectives:"
    )
    for obj in [
        "Improved Lifecycle Affordability \u2013 Reduce total ownership cost through competitive module "
        "re-sourcing and elimination of duplicative development across the PEO portfolio.",
        "Increased Readiness \u2013 Enable faster, lower-risk upgrades and component replacements that "
        "minimize system downtime and sustainment complexity.",
        "Enhanced Capabilities \u2013 Support rapid technology insertion by preserving open interface "
        "standards that allow new capabilities to be integrated with minimal redesign.",
        "Reduced Schedule Pressure \u2013 Enable parallel module development and independent module "
        "acceptance testing, compressing integration timelines.",
        "Reduced Supply Chain Risk \u2013 Eliminate single-source dependencies through open-standard "
        "interfaces and Government-owned technical data packages.",
    ]:
        doc.add_paragraph(obj, style="List Bullet")

    mosa_objectives_text = (answers.get("g_mosa_scenarios") or "").strip()
    if mosa_objectives_text:
        _h2(doc, "1.2 Program-Specific MOSA Objectives")
        for line in mosa_objectives_text.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")

    # PEO AVN Title 10 checklist highlights
    _h2(doc, "1.3 Title 10 Compliance Intent")
    doc.add_paragraph(
        "The following requirements from 10 U.S.C. \u00a7 4401 and the FY17/FY21 NDAAs are addressed by "
        "this Conformance Plan:"
    )
    t10_items = [
        ("Modular design with modular system interfaces",
         "Modules defined in Section 3; MSIs documented in Section 4 Interface Registry"),
        ("Verification of modular system interface compliance",
         "Verification events and evidence mapped in Section 7 Conformance Verification Plan"),
        ("Architecture enabling severability",
         "Module functional boundaries and severability rationale documented in Section 3"),
        ("Technical data rights per 10 U.S.C. \u00a7 2320",
         "Data rights strategy and deliverables matrix documented in Section 6"),
        ("Identification of modular systems in acquisition strategy",
         "MSCs and MSIs identified in Sections 3 and 4; reflected in program Acquisition Strategy"),
        ("Delivery and Government access to Modular System Interfaces",
         "ICD delivery requirements and Government rights documented in Section 4 and Section 6"),
    ]
    col_widths = [Inches(2.4), Inches(3.6)]
    tbl = doc.add_table(rows=len(t10_items) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(["10 U.S.C. \u00a7 4401 Requirement", "How This Plan Addresses It"]):
        _set_cell_text(hdr[i], htext, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (req, how) in enumerate(t10_items, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], req)
        _set_cell_text(cells[1], how)
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]


def _mosa_plan_tailoring(
    doc: DocxDocument,
    program: Program,
    brief: ProgramBrief | None,
    answers: dict[str, str | None],
) -> None:
    """Section 2 – Tailoring Rationale (per MIG §7.1)."""
    _h1(doc, "2. Tailoring Rationale")
    doc.add_paragraph(
        "Per PEO AVN MIG \u00a77.1, each program must tailor its MOSA to address enterprise concerns "
        "defined in the Enterprise Architecture Framework (EAF) while meeting program-specific objectives. "
        "This section documents the tailoring decisions made for this program, the rationale behind each "
        "decision, and the resulting conformance targets. Programs will be measured against these tailored "
        "degrees of openness and modularity to achieve their program and applicable enterprise objectives."
    )

    _h2(doc, "2.1 Lifecycle Phase and Program Context")
    lifecycle_answer = (answers.get("b_acquisition_phase") or "").strip()
    phase = lifecycle_answer or "To be determined"
    brief_desc = (brief.program_description or "").strip() if brief else ""
    wizard_desc = (answers.get("a_program_description") or "").strip()
    prog_desc = brief_desc or wizard_desc or "Not provided"

    doc.add_paragraph(
        f"Current acquisition phase: {phase}. "
        f"Program description: {prog_desc}"
    )

    _h2(doc, "2.2 MOSA Tailoring Decisions")
    doc.add_paragraph(
        "The following table documents each MOSA tailoring decision, its rationale, and the resulting "
        "conformance target. Tailoring must maintain alignment with PEO AVN MOSA Objectives and the EAF."
    )

    sw_intensive = bool(brief and brief.software_large_part) if brief else False
    attritable = bool(brief and brief.attritable) if brief else False
    mission_critical = bool(brief and brief.mission_critical) if brief else False
    safety_critical = bool(brief and brief.safety_critical) if brief else False
    sustainment = bool(brief and brief.sustainment_tail) if brief else False

    tailoring_rows: list[tuple[str, str, str]] = [
        (
            "FACE Technical Standard Applicability",
            "Software-intensive: full FACE applicability" if sw_intensive
            else "Applicability to be determined at SRR in coordination with APEO E&A",
            "All FACE UoCs at key interfaces shall complete at minimum internal FACE Verification "
            "with CTS test results delivered as CDRL; high-value UoCs shall complete independent FACE "
            "Conformance Program" if sw_intensive else "TBD",
        ),
        (
            "CMOSS/AMCE CSM Applicability",
            "Hardware architecture shall conform to AMCE Compute Sub-Module (CSM) requirements "
            "unless deviation is approved" if not attritable
            else "Attritable platform: AMCE CSM conformance is not required; modular hardware "
            "boundaries and ICD delivery remain required",
            "Contractor shall submit MOSA Standard Deviations List (DI-MISC-80508) if non-conformant; "
            "PM approval required prior to implementation",
        ),
        (
            "Severability Scope",
            "Mission-critical / safety-critical platform: all MSCs must be severable; "
            "replacement-supplier demonstration required at CDR" if (mission_critical or safety_critical)
            else "Severability defined per MSC boundaries in Section 3; demonstrated at CDR",
            "Each MSC shall be independently replaceable using only Government-owned ICDs without "
            "original contractor participation",
        ),
        (
            "Long-Term Sustainment Data Rights",
            "Long sustainment tail anticipated: Government will acquire unlimited or GPR for all "
            "major design artifacts to enable organic or third-party sustainment" if sustainment
            else "Data rights strategy defined in Section 6; tailored to program lifecycle",
            "Technical Data Package (TDP) delivery with unlimited rights on ICDs; minimum GPR on "
            "all other design data",
        ),
        (
            "Open System Architecture Framework Alignment",
            "Program architecture shall align to PEO AVN EAF and Enterprise Product Architecture (EPA). "
            "Any deviations require APEO E&A concurrence",
            "Solution Architecture model delivered at PDR; deviation register maintained and reviewed "
            "at each milestone",
        ),
    ]

    tech_challenges = (answers.get("f_tech_challenges_and_risk_areas") or "").strip()
    if tech_challenges:
        first_challenge = next((l.strip() for l in tech_challenges.splitlines() if l.strip()), "")
        if first_challenge:
            tailoring_rows.append((
                "Program-Specific Tailoring",
                f"Technical constraint identified: {first_challenge[:120]}",
                "Tailoring approved by APEO E&A; deviation documented and risk-tracked",
            ))

    col_widths = [Inches(1.6), Inches(2.2), Inches(2.2)]
    tbl = doc.add_table(rows=len(tailoring_rows) + 1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(["Tailoring Area", "Rationale", "Conformance Target"]):
        _set_cell_text(hdr[i], htext, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (area, rationale, target) in enumerate(tailoring_rows, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], area)
        _set_cell_text(cells[1], rationale)
        _set_cell_text(cells[2], target)
        for i, cell in enumerate(cells):
            cell.width = col_widths[i]


def _mosa_plan_standards_architecture(
    doc: DocxDocument,
    modules: list[Module],
    answers: dict[str, str | None],
    brief: ProgramBrief | None,
) -> None:
    """Section 5 – Standards & Architecture Mapping (per MIG §7.4 and §6.0)."""
    _h1(doc, "5. Standards & Architecture Mapping")
    doc.add_paragraph(
        "The following table maps each module to its applicable open standards and reference architectures. "
        "Selection of consensus-based standards (MOSA Principle #4) and their proper application are "
        "measured at each program review per PEO AVN MIG \u00a77.1. Compliance with the FACE Technical "
        "Standard, CMOSS/AMCE, SOSA, or other applicable standards shall be verified per Section 7. "
        "Standards marked 'Required' are contract requirements; those marked 'Preferred' are strongly "
        "encouraged and deviations require PM approval."
    )

    n_answer = (answers.get("n_software_standards_architectures") or "").strip()
    i_answer = (answers.get("i_known_standards_architectures_mapping") or "").strip()
    sw_intensive = bool(brief and brief.software_large_part) if brief else False

    if n_answer or i_answer:
        doc.add_paragraph("Program-identified standards and architecture notes:")
        if n_answer:
            doc.add_paragraph(n_answer)
        if i_answer and i_answer != n_answer:
            doc.add_paragraph(i_answer)

    _h2(doc, "5.1 Module-Level Standards Mapping")

    # Build standards rows from module data; supplement with well-known defaults where missing
    std_rows: list[tuple[str, str, str, str, str]] = []
    known_sw_stds = (
        "FACE Technical Standard Ed. 3.1 or later"
        if sw_intensive else "TBD at SRR"
    )
    known_hw_stds = "CMOSS/AMCE CSM Specification"

    for mod in modules:
        mod_stds = (mod.standards or "").strip() or known_sw_stds
        ifaces = (mod.key_interfaces or "").strip() or "TBD"
        arch = "FACE UoC" if sw_intensive else "MSC per EAF"
        applicability = "Required" if sw_intensive else "Preferred"
        std_rows.append((
            safe_text(mod.name),
            mod_stds[:120],
            ifaces[:80],
            arch,
            applicability,
        ))

    if not std_rows:
        std_rows = [
            ("Mission System Software", known_sw_stds, "TBD", "FACE UoC", "Required" if sw_intensive else "TBD"),
            ("Hardware Platform", known_hw_stds, "TBD", "AMCE CSM", "Required"),
            ("Communications", "VICTORY / CMOSS Comms", "MIL-STD-1553, Ethernet", "MSC per EAF", "Required"),
            ("Navigation & Control", "DO-178C / DO-254 as applicable", "TBD", "MSC per EAF", "Preferred"),
        ]

    col_widths = [Inches(1.2), Inches(1.8), Inches(1.3), Inches(1.0), Inches(0.7)]
    tbl = doc.add_table(rows=len(std_rows) + 1, cols=5)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(["Module", "Applicable Standards", "Key Interfaces", "Architecture Pattern", "Applicability"]):
        _set_cell_text(hdr[i], htext, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (mod_name, stds, ifaces, arch, appl) in enumerate(std_rows, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], mod_name)
        _set_cell_text(cells[1], stds)
        _set_cell_text(cells[2], ifaces)
        _set_cell_text(cells[3], arch)
        _set_cell_text(cells[4], appl, center=True)
        for i, cell in enumerate(cells):
            cell.width = col_widths[i]

    _h2(doc, "5.2 Enterprise Architecture Alignment")
    doc.add_paragraph(
        "The program architecture shall be derived from the PEO AVN Enterprise Architecture Framework "
        "(EAF) and shall align to the Enterprise Product Architecture (EPA). The Government-Furnished "
        "Information (GFI) Component System Model (CSM) will be provided to the prime contractor and "
        "shall form the baseline for the Solution Architecture. Architecture deviations require APEO "
        "Engineering & Architecture concurrence prior to implementation."
    )
    o_answer = (answers.get("o_mosa_repo_searched") or "").strip()
    if o_answer:
        status = "has been completed" if o_answer.lower() == "yes" else "has not yet been completed"
        doc.add_paragraph(
            f"A search of the PEO AVN MOSA reuse repository {status} for this program. "
            "Results will be incorporated into standards selection and component reuse decisions."
        )


def _mosa_plan_risk_register(
    doc: DocxDocument,
    modules: list[Module],
    answers: dict[str, str | None],
    brief: ProgramBrief | None,
) -> None:
    """Section 8 – MOSA Risk Register (per MIG §7.5 and §7.5.5)."""
    _h1(doc, "8. MOSA Risk Register")
    doc.add_paragraph(
        "Per PEO AVN MIG \u00a77.5.5, deviations from the MOSA plan at PDR and CDR are tracked as "
        "official program risks with clear mitigations. The following risk register documents MOSA-specific "
        "risks to the program\u2019s modularity and openness objectives. Risks are assessed for Likelihood "
        "(H/M/L) and Impact on MOSA objectives (H/M/L). Risk owners are responsible for mitigation "
        "execution and status reporting at each milestone review."
    )

    sw_intensive = bool(brief and brief.software_large_part) if brief else False
    tech_challenges = (answers.get("f_tech_challenges_and_risk_areas") or "").strip()
    obs_answer = (answers.get("j_obsolescence_candidates") or "").strip()

    # Standard MOSA-specific risks per MIG
    risks: list[tuple[str, str, str, str, str]] = [
        (
            "Proprietary Interface Lock-in",
            "Contractor delivers proprietary interface implementations that prevent competitive "
            "re-sourcing of modules at contract re-competition",
            "M", "H",
            "Require open-standard interfaces as contract requirement; Government review of ICD "
            "content at each milestone; enforce ICD delivery with unlimited rights",
        ),
        (
            "Inadequate Data Rights Acquisition",
            "Government fails to acquire sufficient technical data rights to enable independent "
            "module sustainment or re-competition",
            "M", "H",
            "Data rights assertions reviewed by legal counsel at PDR; unlimited rights required "
            "for all ICDs; GPR minimum for all design data per DFARS 252.227-7013/7014",
        ),
        (
            "Deviation from EAF/EPA Without Concurrence",
            "Contractor solution architecture diverges from PEO AVN EAF without APEO E&A approval, "
            "undermining enterprise reuse and interoperability objectives",
            "M", "M",
            "Solution Architecture CSM delivered at PDR for Government review; deviations logged in "
            "MOSA Deviations List (DI-MISC-80508); no deviation implemented without PM approval",
        ),
        (
            "FACE/CMOSS Conformance Schedule Risk",
            "Software or hardware components fail to achieve required FACE/CMOSS conformance by "
            "CDR, delaying competitive integration demonstration",
            "M" if sw_intensive else "L", "H" if sw_intensive else "M",
            "Early engagement with FACE Verification Agent (VA); FACE Conformance Verification "
            "Matrix (CVM) maintained starting at SRR; CTS testing milestones embedded in IMS",
        ),
        (
            "Insufficient Module Severability",
            "As-built system does not support independent module replacement without original "
            "contractor support, violating 10 U.S.C. \u00a7 4401 architecture requirements",
            "L", "H",
            "Competitive demonstration at CDR: alternate-source module integrated using only "
            "Government-owned ICDs; Government witness required for demonstration acceptance",
        ),
        (
            "Interface Creep / ICD Drift",
            "Undocumented interface changes accumulate, causing ICD to diverge from as-built "
            "system and reducing Government ability to independently specify replacement modules",
            "M", "M",
            "ICWG reviews all ECP / EO actions affecting interfaces; ICD change triggers "
            "Government notification; ICD configuration control enforced in contract CM plan",
        ),
    ]

    # Module-specific risks
    for mod in modules:
        if mod.tech_risk:
            risks.append((
                f"Tech Maturity Risk: {safe_text(mod.name)[:28]}",
                f"{safe_text(mod.name)[:50]} is identified as a high-risk technical area. "
                "Low TRL may delay interface stabilization and ICD finalization",
                "H", "M",
                f"Early prototype for {safe_text(mod.name)[:30]}; interface baseline locked by PDR; "
                "modular boundary isolates schedule risk to this module",
            ))
        if mod.obsolescence_risk:
            risks.append((
                f"Obsolescence: {safe_text(mod.name)[:32]}",
                f"{safe_text(mod.name)[:50]} identified as an obsolescence risk. "
                "Component end-of-life may force unplanned redesign",
                "M", "M",
                "Modular replacement path defined in ICD; COTS insertion plan in LCMP; "
                "technology refresh milestone in IMS; Government rights for independent refresh",
            ))

    if tech_challenges:
        first = next((l.strip() for l in tech_challenges.splitlines() if l.strip()), "")
        if first:
            risks.append((
                "Program Tech Risk Impact on MOSA",
                f"Technical challenge \u2013 {first[:100]} \u2013 may limit ability to achieve "
                "required modular boundaries or open-standard compliance",
                "M", "M",
                "Modular architecture limits blast radius; MOSA enables competitive re-sourcing "
                "even if one module requires proprietary solution with PM-approved deviation",
            ))

    col_widths = [Inches(1.2), Inches(1.8), Inches(0.5), Inches(0.5), Inches(2.0)]
    tbl = doc.add_table(rows=len(risks) + 1, cols=5)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(["Risk Area", "Description", "Like.", "Impact", "Mitigation"]):
        _set_cell_text(hdr[i], htext, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (area, desc, like, impact, mit) in enumerate(risks, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], area)
        _set_cell_text(cells[1], desc)
        _set_cell_text(cells[2], like, center=True)
        _set_cell_text(cells[3], impact, center=True)
        _set_cell_text(cells[4], mit)
        for i, cell in enumerate(cells):
            cell.width = col_widths[i]


def _mosa_plan_assessment_matrix(
    doc: DocxDocument,
    program: Program,
    modules: list[Module],
    answers: dict[str, str | None],
    brief: ProgramBrief | None,
) -> None:
    """Section 9 – PEO AVN MOSA Self-Assessment (per MIG §7.4 and Appendix E)."""
    _h1(doc, "9. PEO AVN MOSA Self-Assessment")
    doc.add_paragraph(
        "Per PEO AVN MIG \u00a77.4, the PMO is responsible for completing the PEO Aviation MOSA "
        "Assessment Matrix (Appendix E of the MIG). The matrix documents how the program intends to "
        "satisfy Title 10, ASA(ALT) Implementation Guidance, and PEO Aviation MOSA Implementation "
        "Guidance. This section captures the program\u2019s self-assessment responses against the four "
        "matrix tabs: (1) Title 10 Checklist, (2) ASA(ALT) Guidance mapping, (3) ASA(ALT) MOSA "
        "Checklist, and (4) PEO Aviation MOSA Checklist. The completed MOSA Assessment Matrix "
        "(Excel workbook) accompanies this plan as an attachment."
    )

    sw_intensive = bool(brief and brief.software_large_part) if brief else False
    attritable = bool(brief and brief.attritable) if brief else False
    mosa_scenarios = (answers.get("g_mosa_scenarios") or "").strip()
    i_answer = (answers.get("i_known_standards_architectures_mapping") or "").strip()
    o_answer = (answers.get("o_mosa_repo_searched") or "").strip()

    _h2(doc, "9.1 PEO Aviation MOSA Checklist Responses")
    doc.add_paragraph(
        "The following table provides self-assessment responses to PEO Aviation stakeholder concerns "
        "per MIG \u00a77.5. Responses with robust rationale enable PEO E&A architects to evaluate "
        "line-by-line conformance per MIG \u00a77.4. Rationale should remove as much subjectivity from "
        "the measurement as practical."
    )

    module_names = ", ".join(m.name for m in modules[:5]) if modules else "TBD"
    if len(modules) > 5:
        module_names += f" + {len(modules) - 5} more"

    checklist: list[tuple[str, str, str]] = [
        (
            "Has the program established MOSA objectives aligned to PEO AVN enterprise goals?",
            "Yes \u2013 documented in Section 1 of this MCP",
            "MOSA objectives are derived from PEO AVN MIG enterprise goals and tailored to "
            f"{safe_text(program.name)} program needs",
        ),
        (
            "Has the program defined MOSA scenarios and use cases driving modularity decisions?",
            "Yes" if mosa_scenarios else "Partially \u2013 initial scenarios defined; to be refined at SRR",
            mosa_scenarios[:200] if mosa_scenarios else
            "Standard MOSA lifecycle scenarios documented in Section 10; program-specific scenarios "
            "to be developed with APEO E&A at SRR",
        ),
        (
            "Is the Solution Architecture (SA) aligned with the EAF and mapped to MOSA objectives?",
            "In progress \u2013 SA to be derived from GFI CSM and delivered at PDR",
            "Program will work with APEO E&A to generate SA from EAF; model delivered at PDR "
            "with deviations documented in MOSA Deviations List",
        ),
        (
            "Are Major System Components (MSCs) and Modular System Interfaces (MSIs) defined?",
            "Partially \u2013 candidate modules identified; formal MSC/MSI boundaries at SRR",
            f"Candidate modules: {module_names}. Formal MSC boundaries will be established "
            "at SRR in coordination with APEO E&A",
        ),
        (
            "Has the program searched the MOSA reuse repository for applicable components?",
            "Yes" if o_answer.lower() == "yes" else "Not yet \u2013 planned prior to RFP release",
            "Repository search results will be incorporated into standards selection and "
            "statement of work; existing MOSA QMR library components will be considered",
        ),
        (
            "Does the program leverage FACE Technical Standard for software components?",
            "Yes \u2013 required for all UoCs at key interfaces" if sw_intensive else "TBD at SRR",
            "FACE applicability determined based on software reuse potential and interface "
            "criticality; FACE VA engagement planned" if sw_intensive else
            "FACE applicability will be determined at SRR with APEO E&A; not applicable to "
            "attritable-only software" if attritable else
            "FACE applicability to be determined at SRR",
        ),
        (
            "Are component models developed with open interfaces or GPR for used interfaces?",
            "Yes \u2013 required by contract; ICD delivery with unlimited rights",
            i_answer[:200] if i_answer else
            "All module interfaces shall be documented in Government-owned ICDs; contractor-proprietary "
            "interfaces prohibited without PM approval",
        ),
        (
            "Is the project CSM aligned with the EAF / FAF and GFI model?",
            "In progress \u2013 GFI CSM to be provided to contractor at contract award",
            "Government will provide GFI CSM derived from PEO AVN EAF; contractor shall deliver "
            "Solution Architecture model derived from GFI CSM at PDR",
        ),
        (
            "Are data rights being acquired consistent with requirements and MOSA expectations?",
            "Yes \u2013 data rights strategy documented in Section 6",
            "Unlimited rights on ICDs; minimum GPR on design data; DFARS 252.227-7013/7014 "
            "clause included in all contracts; assertions reviewed at PDR",
        ),
        (
            "Is the MOSA TPM defined and reported at each technical review?",
            "Planned \u2013 MOSA TPM to be defined at SRR and reported at each review thereafter",
            "MOSA TPM will track: % of MSIs with Government-owned ICDs; % of modules meeting "
            "open-standard conformance; severability demonstration status",
        ),
    ]

    col_widths = [Inches(2.0), Inches(0.8), Inches(3.2)]
    tbl = doc.add_table(rows=len(checklist) + 1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(["PEO AVN Stakeholder Concern / Question", "Response", "Rationale"]):
        _set_cell_text(hdr[i], htext, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r, (question, response, rationale) in enumerate(checklist, start=1):
        cells = tbl.rows[r].cells
        _set_cell_text(cells[0], question)
        _set_cell_text(cells[1], response)
        _set_cell_text(cells[2], rationale)
        for i, cell in enumerate(cells):
            cell.width = col_widths[i]

    _h2(doc, "9.2 Iterative Assessment Schedule")
    doc.add_paragraph(
        "Per MIG \u00a77.5.1, MOSA assessments are iterative across the program lifecycle. The following "
        "schedule establishes when formal self-assessments and independent assessments will be conducted:"
    )
    assessment_rows: list[tuple[str, str, str]] = [
        ("Pre-RFP", "PMO Self-Assessment (Step 1)", "Draft MCP; initial MOSA Assessment Matrix; GFI CSM issued"),
        ("Post-Contract Award", "PMO review of contractor OSMP", "Contractor OSMP reviewed against MCP; deviations documented"),
        ("Pre-SRR/SFR", "PMO Self-Assessment (Step 1)", "Updated Assessment Matrix; SRR entrance criteria include MOSA status"),
        ("PDR", "PMO + Independent Assessment (Steps 1\u20132)", "Solution Architecture CDM review; MOSA TPM status; ICD baseline review"),
        ("CDR", "Independent Assessment (Step 2) + AAE if MDAP", "Competitive module demonstration; final ICD review; data rights verification"),
        ("Pre-LRIP/FRP", "Third-party assessment via OpenITB or equivalent", "Independent integration demonstration; replaceability validation"),
    ]
    col_widths2 = [Inches(1.0), Inches(1.6), Inches(3.4)]
    tbl2 = doc.add_table(rows=len(assessment_rows) + 1, cols=3)
    tbl2.style = "Table Grid"
    hdr2 = tbl2.rows[0].cells
    for i, htext in enumerate(["Phase / Milestone", "Assessment Type", "Activities & Evidence"]):
        _set_cell_text(hdr2[i], htext, bold=True)
        _shade_cell(hdr2[i])
        hdr2[i].width = col_widths2[i]
    for r, (phase, atype, activities) in enumerate(assessment_rows, start=1):
        cells = tbl2.rows[r].cells
        _set_cell_text(cells[0], phase)
        _set_cell_text(cells[1], atype)
        _set_cell_text(cells[2], activities)
        for i, cell in enumerate(cells):
            cell.width = col_widths2[i]


def build_mosa_conformance_plan(
    program: Program,
    brief: ProgramBrief | None,
    answers: dict[str, str | None],
    files: list[ProgramFile],
    modules: list[Module],
    generated_date: str,
) -> DocxDocument:
    """Build a MOSA Conformance Plan per PEO AVN MIG v2.2, 10 U.S.C. § 4401, and DoD Directive 5000.01.

    Document structure (per MIG §7.2):
      1. MOSA Objectives (statutory + PEO AVN + program-specific)
      2. Tailoring Rationale (MIG §7.1)
      3. Module Inventory / Decomposition
      4. Interface Registry (HW / SW / Data)
      5. Standards & Architecture Mapping
      6. Technical Data & Data Rights
      7. Conformance Verification Plan (milestones → evidence)
      8. MOSA Risk Register
      9. PEO AVN Self-Assessment Matrix
     10. MOSA Scenario Mapping
    """
    doc = DocxDocument()
    _setup_margins(doc)
    _add_title_page(doc, program, generated_date, "MOSA Conformance Plan")
    doc.add_page_break()

    # 1. MOSA Objectives
    _mosa_plan_objectives(doc, program, brief, answers)
    doc.add_page_break()

    # 2. Tailoring Rationale
    _mosa_plan_tailoring(doc, program, brief, answers)
    doc.add_page_break()

    # 3. Program Overview & Reference Documents (renumbered to 3)
    _mosa_plan_program_overview(doc, program, brief, answers, files)
    doc.add_page_break()

    # 3.1 / Module Inventory – called from within program overview context
    _mosa_plan_modular_architecture(doc, modules, answers)
    doc.add_page_break()

    # 4. Interface Registry (HW/SW/Data) - renumbered to 4 inside function
    _mosa_plan_interface_control(doc, modules, answers)
    doc.add_page_break()

    # 5. Standards & Architecture Mapping
    _mosa_plan_standards_architecture(doc, modules, answers, brief)
    doc.add_page_break()

    # 6. Technical Data & Data Rights (renumber heading)
    _mosa_plan_data_rights(doc, brief, modules, answers)
    doc.add_page_break()

    # 7. Conformance Verification Plan (renumber heading)
    _mosa_plan_verification(doc, modules, answers)
    doc.add_page_break()

    # 8. MOSA Risk Register
    _mosa_plan_risk_register(doc, modules, answers, brief)
    doc.add_page_break()

    # 9. PEO AVN Self-Assessment Matrix
    _mosa_plan_assessment_matrix(doc, program, modules, answers, brief)
    doc.add_page_break()

    # 10. MOSA Scenario Mapping (renumber heading)
    _mosa_plan_scenario_mapping(doc, modules, answers)

    _configure_footer(doc, program.name)
    return doc


# ---------------------------------------------------------------------------
# SEP public builder
# ---------------------------------------------------------------------------


def build_sep(
    program: Program,
    brief: ProgramBrief | None,
    answers: dict[str, str | None],
    files: list[ProgramFile],
    modules: list[Module],
    generated_date: str,
) -> DocxDocument:
    """Build a Systems Engineering Plan (SEP) per OSD SEP Guide v4.1."""
    doc = DocxDocument()
    _setup_margins(doc)
    _add_title_page(doc, program, generated_date, "Systems Engineering Plan (SEP)")
    doc.add_page_break()

    _sep_introduction(doc, program, brief, answers, files, modules)
    doc.add_page_break()

    _sep_technical_baseline(doc, program, brief, answers, modules)
    doc.add_page_break()

    _sep_se_approach(doc, program, brief, modules, answers)
    doc.add_page_break()

    _sep_risk_management(doc, program, modules, answers, brief)
    doc.add_page_break()

    _sep_config_management(doc, program, brief, modules, answers, files)
    doc.add_page_break()

    _sep_mosa(doc, program, brief, modules, answers)
    doc.add_page_break()

    _sep_vv(doc, program, brief, modules, answers)
    doc.add_page_break()

    _sep_specialty_engineering(doc, program, brief, answers, modules)
    doc.add_page_break()

    _sep_se_organization(doc, program, brief, modules)
    doc.add_page_break()

    _sep_summary(doc, program, brief, modules, answers)

    _configure_footer(doc, program.name)
    return doc


# ===========================================================================
# Smart doc render system  (AI-assisted, template-driven)
#
# Architecture
# ------------
# • Section headings come from TEMPLATE_REGISTRY[doc_type]["section_order"].
#   Builders never hard-code heading strings; _h1() is always called with the
#   template string so the document matches the registry exactly.
# • _SMART_SECTION_RENDERERS maps (doc_type, heading) → renderer callable.
#   Each renderer fills content for one section; it never writes the H1 heading
#   itself (that is done by the shared _render_smart_doc_body loop).
# • section_format from the template gates formatting:
#   "narrative"  → _smart_render_narrative only
#   "bullet"     → _smart_render_bullets only
#   "mixed"      → narrative intro + table or bullets
#   "table"      → structured table only
# • Optional appendix sections (heading starts with "Appendix") are skipped
#   automatically when they have no content.
# ===========================================================================


# ---------------------------------------------------------------------------
# Shared render primitives
# ---------------------------------------------------------------------------

def _smart_render_narrative(doc: DocxDocument, text: str) -> None:
    """Render a string as multi-paragraph prose.  Splits on double-newlines."""
    text = (text or "").strip()
    if not text or text == "Not provided":
        doc.add_paragraph("Not provided.")
        return
    for block in text.split("\n\n"):
        stripped = block.strip()
        if stripped:
            doc.add_paragraph(stripped)


def _smart_render_bullets(doc: DocxDocument, items: list) -> None:
    """Render a list of strings as bullet points."""
    if not items:
        doc.add_paragraph("Not provided.")
        return
    for item in items:
        stripped = str(item).strip()
        if stripped:
            doc.add_paragraph(stripped, style="List Bullet")


def _smart_render_module_table(doc: DocxDocument, rows: list) -> None:
    """Render module_table_rows as a 7-column module summary table."""
    if not rows:
        doc.add_paragraph("No candidate modules have been defined for this program.")
        return
    col_widths = [
        Inches(0.9), Inches(1.2), Inches(1.0), Inches(0.9),
        Inches(0.6), Inches(0.8), Inches(0.6),
    ]
    tbl = doc.add_table(rows=len(rows) + 1, cols=7)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, label in enumerate([
        "Module", "Rationale", "Key Interfaces", "Standards",
        "Tech Risk", "Obs. Risk", "COTS?",
    ]):
        _set_cell_text(hdr[i], label, bold=True, font_size_pt=9)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r_idx, row in enumerate(rows, start=1):
        cells = tbl.rows[r_idx].cells
        _set_cell_text(cells[0], row.module_name or "", font_size_pt=9)
        _set_cell_text(cells[1], row.rationale or "", font_size_pt=9)
        _set_cell_text(cells[2], row.key_interfaces or "", font_size_pt=9)
        _set_cell_text(cells[3], row.standards or "", font_size_pt=9)
        _set_cell_text(cells[4], row.tech_risk or "", center=True, font_size_pt=9)
        _set_cell_text(cells[5], row.obsolescence_risk or "", center=True, font_size_pt=9)
        _set_cell_text(cells[6], row.cots_candidate or "", center=True, font_size_pt=9)
        for i, cell in enumerate(cells):
            cell.width = col_widths[i]


def _smart_render_sources(doc: DocxDocument, sources: list) -> None:
    """Render a sources_used list.

    Handles both RFI/ACQ/SEP (filename, excerpt) and MOSA (source_filename,
    excerpt, chunk_id) attribute shapes via getattr with fallback.
    """
    for src in sources:
        p = doc.add_paragraph(style="List Bullet")
        fname = getattr(src, "filename", None) or getattr(src, "source_filename", "")
        name_run = p.add_run(fname)
        name_run.bold = True
        excerpt = (getattr(src, "excerpt", "") or "").strip()
        chunk_id = getattr(src, "chunk_id", None)
        if excerpt:
            suffix = (
                f" (chunk {chunk_id}) \u2014 {excerpt}" if chunk_id else f" \u2014 {excerpt}"
            )
            p.add_run(suffix)


# ---------------------------------------------------------------------------
# RFI section renderers  (keyed to template section headings)
# ---------------------------------------------------------------------------

def _rfi_document_overview(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.overview)


def _rfi_rfi_purpose(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.rfi_purpose)


def _rfi_program_context(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.program_context)


def _rfi_mosa_requirements(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "This acquisition will be structured in accordance with a Modular Open Systems "
        "Approach (MOSA) as required by 10 U.S.C. \u00a7 4401. "
        "The following requirements apply to this specific program:"
    )
    _smart_render_bullets(doc, plan.mosa_requirements)


def _rfi_candidate_module_summary(doc: DocxDocument, plan) -> None:
    if plan.module_table_rows:
        doc.add_paragraph(
            "The Government has identified the following candidate modules as a starting "
            "point for the modular architecture. Respondents are encouraged to propose "
            "alternative decompositions."
        )
    _smart_render_module_table(doc, plan.module_table_rows)


def _rfi_questions_to_industry(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "Industry respondents are requested to address the following questions. "
        "Responses are for market research purposes only and do not constitute a "
        "commitment to procure."
    )
    _smart_render_bullets(doc, plan.questions_to_industry)


def _rfi_requested_deliverables(doc: DocxDocument, plan) -> None:
    doc.add_paragraph("Please provide the following items in your response:")
    _smart_render_bullets(doc, plan.requested_deliverables)


def _rfi_submission_instructions(doc: DocxDocument, plan) -> None:
    _smart_render_bullets(doc, plan.submission_instructions)


def _rfi_appendix_sources(doc: DocxDocument, plan) -> None:
    if plan.sources_used:
        doc.add_paragraph(
            "The following uploaded program documents were consulted in preparing "
            "this AI-assisted RFI:"
        )
        _smart_render_sources(doc, plan.sources_used)


# ---------------------------------------------------------------------------
# Acquisition Strategy section renderers
# ---------------------------------------------------------------------------

def _acq_executive_summary(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.executive_summary)


def _acq_acquisition_approach(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.acquisition_approach)


def _acq_schedule_milestones(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "The following milestones represent the planned acquisition schedule. "
        "Dates are subject to revision based on program execution."
    )
    if not plan.schedule_milestones:
        doc.add_paragraph("Schedule milestones to be defined at program initiation.")
        return
    ms_headers = ["Milestone", "Target Date", "Description"]
    ms_widths = [Inches(1.3), Inches(1.1), Inches(3.6)]
    tbl = doc.add_table(rows=len(plan.schedule_milestones) + 1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(ms_headers):
        _set_cell_text(hdr[i], h, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = ms_widths[i]
    for r_idx, ms in enumerate(plan.schedule_milestones, start=1):
        cells = tbl.rows[r_idx].cells
        _set_cell_text(cells[0], ms.name or "")
        _set_cell_text(cells[1], ms.date or "", center=True)
        _set_cell_text(cells[2], ms.description or "")
        for i, cell in enumerate(cells):
            cell.width = ms_widths[i]


def _acq_cost_estimates(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "The following cost estimates are provided as program planning figures. "
        "Final cost estimates will be developed through the Cost Analysis Requirements "
        "Description (CARD) process."
    )
    ce = plan.cost_estimates
    cost_rows = [
        ("Development Cost Estimate", ce.development),
        ("Production Unit Cost", ce.production_unit),
        ("Annual Sustainment Cost", ce.sustainment_annual),
    ]
    tbl = doc.add_table(rows=len(cost_rows) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Cost Element", "Estimate"]):
        _set_cell_text(hdr[i], h, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = Inches(3.0)
    for r_idx, (label, val) in enumerate(cost_rows, start=1):
        cells = tbl.rows[r_idx].cells
        _set_cell_text(cells[0], label)
        _set_cell_text(cells[1], val or "Not provided", center=True)
        cells[0].width = Inches(3.0)
        cells[1].width = Inches(3.0)


def _acq_candidate_module_summary(doc: DocxDocument, plan) -> None:
    if plan.module_table_rows:
        doc.add_paragraph(
            "The program has identified the following candidate modules consistent with a "
            "Modular Open Systems Approach (MOSA). Each module will be designed with "
            "well-defined interfaces to enable competitive upgrade and re-compete opportunities."
        )
    _smart_render_module_table(doc, plan.module_table_rows)


def _acq_mosa(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "This acquisition will implement a Modular Open Systems Approach (MOSA) as required "
        "by 10 U.S.C. \u00a7 4401 and DoD Directive 5000.01."
    )
    _smart_render_narrative(doc, plan.mosa_approach)
    if plan.mosa_bullets:
        _h2(doc, "6.1 MOSA Commitments")
        _smart_render_bullets(doc, plan.mosa_bullets)


def _acq_data_rights(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.data_rights_approach)


def _acq_standards_constraints(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "The following standards, regulations, and constraints apply to this program:"
    )
    if not plan.standards_references:
        doc.add_paragraph("Standards to be identified during system design.")
        return
    for ref in plan.standards_references:
        p = doc.add_paragraph(style="List Bullet")
        name_run = p.add_run(ref.name)
        name_run.bold = True
        if ref.description and ref.description.strip():
            p.add_run(f" \u2014 {ref.description.strip()}")


def _acq_test_verification(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.test_verification_approach)


def _acq_contracting_strategy(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.contracting_strategy)


def _acq_risk_register(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "The following risks have been identified for this acquisition. "
        "Risks will be tracked and managed through the program's Risk Management process."
    )
    if not plan.risk_register:
        doc.add_paragraph("No acquisition risks identified at this time.")
        return
    risk_headers = ["ID", "Description", "Probability", "Impact", "Mitigation", "Owner"]
    risk_widths = [
        Inches(0.4), Inches(1.5), Inches(0.75), Inches(0.65), Inches(2.0), Inches(0.7),
    ]
    tbl = doc.add_table(rows=len(plan.risk_register) + 1, cols=6)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(risk_headers):
        _set_cell_text(hdr[i], h, bold=True, font_size_pt=9)
        _shade_cell(hdr[i])
        hdr[i].width = risk_widths[i]
    for r_idx, risk in enumerate(plan.risk_register, start=1):
        cells = tbl.rows[r_idx].cells
        _set_cell_text(cells[0], risk.risk_id or "", center=True, font_size_pt=9)
        _set_cell_text(cells[1], risk.description or "", font_size_pt=9)
        _set_cell_text(cells[2], risk.probability or "", center=True, font_size_pt=9)
        _set_cell_text(cells[3], risk.impact or "", center=True, font_size_pt=9)
        _set_cell_text(cells[4], risk.mitigation or "", font_size_pt=9)
        _set_cell_text(cells[5], risk.owner or "", font_size_pt=9)
        for i, cell in enumerate(cells):
            cell.width = risk_widths[i]


def _acq_appendix_sources(doc: DocxDocument, plan) -> None:
    if plan.sources_used:
        doc.add_paragraph(
            "The following uploaded program documents were consulted in preparing "
            "this AI-assisted Acquisition Strategy:"
        )
        _smart_render_sources(doc, plan.sources_used)


# ---------------------------------------------------------------------------
# MOSA Conformance Plan section renderers
# ---------------------------------------------------------------------------

def _mcp_executive_summary(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.executive_summary)


def _mcp_mosa_objectives(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "This MOSA Conformance Plan documents how the program will satisfy the requirements "
        "of 10 U.S.C. \u00a7 4401 and DoD Directive 5000.01. "
        "The following objectives apply:"
    )
    _smart_render_bullets(doc, plan.mosa_objectives)


def _mcp_tailoring_rationale(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.tailoring_rationale)


def _mcp_program_context(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.program_context)


def _mcp_module_inventory(doc: DocxDocument, plan) -> None:
    if not plan.module_inventory:
        doc.add_paragraph("No modules have been defined for this program.")
        return
    doc.add_paragraph(
        "The following table summarizes the program's modular decomposition and the "
        "MOSA conformance status of each module."
    )
    headers = [
        "Module", "Open Standards", "Interface Type", "Data Rights",
        "Conformance", "Verification Method", "Notes",
    ]
    col_widths = [
        Inches(0.9), Inches(1.1), Inches(0.7), Inches(1.0),
        Inches(0.8), Inches(1.0), Inches(0.5),
    ]
    tbl = doc.add_table(rows=len(plan.module_inventory) + 1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True, font_size_pt=9)
        _shade_cell(hdr[i])
        hdr[i].width = col_widths[i]
    for r_idx, row in enumerate(plan.module_inventory, start=1):
        cells = tbl.rows[r_idx].cells
        _set_cell_text(cells[0], row.module_name or "", font_size_pt=9)
        _set_cell_text(cells[1], row.open_standards or "", font_size_pt=9)
        _set_cell_text(cells[2], row.interface_type or "", center=True, font_size_pt=9)
        _set_cell_text(cells[3], row.data_rights_category or "", font_size_pt=9)
        _set_cell_text(cells[4], row.conformance_status or "", center=True, font_size_pt=9)
        _set_cell_text(cells[5], row.verification_method or "", font_size_pt=9)
        _set_cell_text(cells[6], row.notes or "", font_size_pt=9)
        for i, cell in enumerate(cells):
            cell.width = col_widths[i]


def _mcp_interface_registry(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "The following interfaces have been identified. All Government-owned interfaces "
        "will be documented in Interface Control Documents (ICDs) using open, "
        "consensus-based standards."
    )
    if not plan.interface_registry:
        doc.add_paragraph("Interface registry to be populated during system design.")
        return
    headers = ["Interface", "Type", "Standard", "Owner", "Status"]
    widths = [Inches(1.4), Inches(0.7), Inches(1.5), Inches(1.0), Inches(0.8)]
    tbl = doc.add_table(rows=len(plan.interface_registry) + 1, cols=5)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = widths[i]
    for r_idx, iface in enumerate(plan.interface_registry, start=1):
        cells = tbl.rows[r_idx].cells
        _set_cell_text(cells[0], iface.interface_name or "")
        _set_cell_text(cells[1], iface.interface_type or "", center=True)
        _set_cell_text(cells[2], iface.standard or "")
        _set_cell_text(cells[3], iface.owner or "", center=True)
        _set_cell_text(cells[4], iface.status or "", center=True)
        for i, cell in enumerate(cells):
            cell.width = widths[i]


def _mcp_standards_architecture_mapping(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "The program will adopt the following open, consensus-based standards and "
        "reference architectures:"
    )
    _smart_render_bullets(doc, plan.standards_mapping)


def _mcp_technical_data_rights(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.data_rights_posture)


def _mcp_conformance_verification(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "MOSA conformance will be verified at each major program milestone through "
        "the following activities:"
    )
    if not plan.verification_milestones:
        doc.add_paragraph("Verification milestones to be defined at program initiation.")
        return
    headers = ["Milestone", "Evidence Required", "Responsible Party", "Completion Criteria"]
    widths = [Inches(0.9), Inches(1.8), Inches(1.1), Inches(1.7)]
    tbl = doc.add_table(rows=len(plan.verification_milestones) + 1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = widths[i]
    for r_idx, vm in enumerate(plan.verification_milestones, start=1):
        cells = tbl.rows[r_idx].cells
        _set_cell_text(cells[0], vm.milestone or "", center=True)
        _set_cell_text(cells[1], vm.evidence_required or "")
        _set_cell_text(cells[2], vm.responsible_party or "")
        _set_cell_text(cells[3], vm.completion_criteria or "")
        for i, cell in enumerate(cells):
            cell.width = widths[i]


def _mcp_risk_register(doc: DocxDocument, plan) -> None:
    if not plan.risk_register:
        doc.add_paragraph("No MOSA-specific risks identified at this time.")
        return
    headers = ["Risk Area", "Description", "Likelihood", "Impact", "Mitigation"]
    widths = [Inches(1.0), Inches(1.7), Inches(0.8), Inches(0.7), Inches(1.8)]
    tbl = doc.add_table(rows=len(plan.risk_register) + 1, cols=5)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = widths[i]
    for r_idx, risk in enumerate(plan.risk_register, start=1):
        cells = tbl.rows[r_idx].cells
        _set_cell_text(cells[0], risk.risk_area or "")
        _set_cell_text(cells[1], risk.description or "")
        _set_cell_text(cells[2], risk.likelihood or "", center=True)
        _set_cell_text(cells[3], risk.impact or "", center=True)
        _set_cell_text(cells[4], risk.mitigation or "")
        for i, cell in enumerate(cells):
            cell.width = widths[i]


def _mcp_self_assessment_matrix(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "The following assessment matrix tracks conformance against PEO AVN MIG "
        "MOSA criteria:"
    )
    if not plan.assessment_matrix:
        doc.add_paragraph("Assessment matrix to be populated during system design.")
        return
    headers = ["Criterion", "Status", "Evidence / Artifact", "Gap Description"]
    widths = [Inches(1.5), Inches(0.9), Inches(1.6), Inches(1.5)]
    tbl = doc.add_table(rows=len(plan.assessment_matrix) + 1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = widths[i]
    for r_idx, am in enumerate(plan.assessment_matrix, start=1):
        cells = tbl.rows[r_idx].cells
        _set_cell_text(cells[0], am.criterion or "")
        _set_cell_text(cells[1], am.status or "", center=True)
        _set_cell_text(cells[2], am.evidence or "")
        _set_cell_text(cells[3], am.gap_description or "")
        for i, cell in enumerate(cells):
            cell.width = widths[i]


def _mcp_scenario_mapping(doc: DocxDocument, plan) -> None:
    if not plan.mosa_scenarios:
        doc.add_paragraph(
            "No program-specific MOSA scenarios have been identified at this time. "
            "Scenarios will be developed during system design."
        )
        return
    doc.add_paragraph(
        "The following scenarios illustrate how MOSA principles apply to specific "
        "program situations:"
    )
    headers = ["Scenario", "Affected Modules", "Approach", "Applicable Standards"]
    widths = [Inches(1.2), Inches(1.2), Inches(1.8), Inches(1.3)]
    tbl = doc.add_table(rows=len(plan.mosa_scenarios) + 1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = widths[i]
    for r_idx, sc in enumerate(plan.mosa_scenarios, start=1):
        cells = tbl.rows[r_idx].cells
        _set_cell_text(cells[0], sc.scenario_title or "")
        _set_cell_text(cells[1], sc.affected_modules or "")
        _set_cell_text(cells[2], sc.approach or "")
        _set_cell_text(cells[3], sc.applicable_standards or "")
        for i, cell in enumerate(cells):
            cell.width = widths[i]


def _mcp_appendix_sources(doc: DocxDocument, plan) -> None:
    if plan.sources_used:
        doc.add_paragraph(
            "The following document chunks were retrieved and used to inform this "
            "AI-assisted plan:"
        )
        _smart_render_sources(doc, plan.sources_used)


# ---------------------------------------------------------------------------
# SEP section renderers
# ---------------------------------------------------------------------------

def _sep_executive_summary(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.executive_summary)


def _sep_program_overview(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.program_overview)


def _sep_se_strategy(doc: DocxDocument, plan) -> None:
    _smart_render_narrative(doc, plan.se_strategy)


def _sep_technical_reviews(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "The program will conduct the following technical reviews in accordance with "
        "the OSD SEP Guide v4.1. Reviews serve as decision gates; entry and exit "
        "criteria must be satisfied before proceeding."
    )
    if not plan.tech_reviews:
        doc.add_paragraph("Technical reviews to be defined at program initiation.")
        return
    headers = ["Review", "Purpose", "Entry Criteria", "Exit Criteria", "Artifacts"]
    widths = [Inches(1.05), Inches(1.35), Inches(1.1), Inches(1.1), Inches(1.4)]
    tbl = doc.add_table(rows=len(plan.tech_reviews) + 1, cols=5)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True, font_size_pt=9)
        _shade_cell(hdr[i])
        hdr[i].width = widths[i]
    for r_idx, rev in enumerate(plan.tech_reviews, start=1):
        cells = tbl.rows[r_idx].cells
        _set_cell_text(cells[0], rev.name or "", bold=True, font_size_pt=9)
        _set_cell_text(cells[1], rev.purpose or "", font_size_pt=9)
        _set_cell_text(cells[2], rev.entry_criteria or "", font_size_pt=9)
        _set_cell_text(cells[3], rev.exit_criteria or "", font_size_pt=9)
        artifacts = "; ".join(a for a in rev.artifacts if a.strip()) if rev.artifacts else ""
        _set_cell_text(cells[4], artifacts, font_size_pt=9)
        for i, cell in enumerate(cells):
            cell.width = widths[i]


def _sep_requirements_traceability(doc: DocxDocument, plan) -> None:
    rt = plan.requirements_traceability
    _h2(doc, "5.1 Approach")
    _smart_render_narrative(doc, rt.approach)
    _h2(doc, "5.2 Tools")
    doc.add_paragraph(rt.tools or "Tools to be determined.")
    _h2(doc, "5.3 Digital Thread")
    _smart_render_narrative(doc, rt.digital_thread_notes)


def _sep_architecture_mosa(doc: DocxDocument, plan) -> None:
    am = plan.architecture_mosa
    _h2(doc, "6.1 MOSA Approach")
    _smart_render_narrative(doc, am.mosa_approach)
    _h2(doc, "6.2 Modules")
    _smart_render_narrative(doc, am.modules_summary)
    _h2(doc, "6.3 Interfaces")
    _smart_render_narrative(doc, am.interfaces_summary)
    _h2(doc, "6.4 Standards")
    _smart_render_narrative(doc, am.standards_summary)


def _sep_risk_register(doc: DocxDocument, plan) -> None:
    doc.add_paragraph(
        "The following technical risks have been identified. Risks will be managed "
        "and tracked through the program's Risk Management process in accordance "
        "with DoDI 5000.02."
    )
    if not plan.risk_register:
        doc.add_paragraph("No technical risks identified at this time.")
        return
    headers = ["Risk", "Cause", "Likelihood", "Impact", "Mitigation", "Owner"]
    widths = [
        Inches(1.0), Inches(1.1), Inches(0.75), Inches(0.65), Inches(1.75), Inches(0.75),
    ]
    tbl = doc.add_table(rows=len(plan.risk_register) + 1, cols=6)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True, font_size_pt=9)
        _shade_cell(hdr[i])
        hdr[i].width = widths[i]
    for r_idx, risk in enumerate(plan.risk_register, start=1):
        cells = tbl.rows[r_idx].cells
        _set_cell_text(cells[0], risk.risk or "", font_size_pt=9)
        _set_cell_text(cells[1], risk.cause or "", font_size_pt=9)
        _set_cell_text(cells[2], risk.likelihood or "", center=True, font_size_pt=9)
        _set_cell_text(cells[3], risk.impact or "", center=True, font_size_pt=9)
        _set_cell_text(cells[4], risk.mitigation or "", font_size_pt=9)
        _set_cell_text(cells[5], risk.owner or "", font_size_pt=9)
        for i, cell in enumerate(cells):
            cell.width = widths[i]


def _sep_config_mgmt(doc: DocxDocument, plan) -> None:
    cm = plan.config_mgmt
    _h2(doc, "8.1 Approach")
    _smart_render_narrative(doc, cm.approach)
    if cm.baselines:
        _h2(doc, "8.2 Baselines")
        _smart_render_bullets(doc, cm.baselines)
    _h2(doc, "8.3 Change Control")
    _smart_render_narrative(doc, cm.change_control)


def _sep_vnv(doc: DocxDocument, plan) -> None:
    vnv = plan.vnv
    _h2(doc, "9.1 Strategy")
    _smart_render_narrative(doc, vnv.strategy)
    if vnv.test_levels:
        _h2(doc, "9.2 Test Levels")
        _smart_render_bullets(doc, vnv.test_levels)
    _h2(doc, "9.3 Acceptance Criteria")
    _smart_render_narrative(doc, vnv.acceptance_criteria)


def _sep_data_mgmt(doc: DocxDocument, plan) -> None:
    dm = plan.data_mgmt
    if dm.data_items:
        _h2(doc, "10.1 Contract Data Requirements (CDRLs)")
        _smart_render_bullets(doc, dm.data_items)
    _h2(doc, "10.2 Data Rights Strategy")
    _smart_render_narrative(doc, dm.data_rights_strategy)
    _h2(doc, "10.3 Repository and Tools")
    _smart_render_narrative(doc, dm.repo_notes)


def _sep_specialty_eng(doc: DocxDocument, plan) -> None:
    se = plan.specialty_eng
    for sub_heading, content in [
        ("11.1 Cybersecurity / RMF",          se.cyber),
        ("11.2 System Safety",                se.safety),
        ("11.3 Airworthiness",                se.airworthiness),
        ("11.4 Human Systems Integration",    se.human_factors),
        ("11.5 Reliability",                  se.reliability),
        ("11.6 Maintainability",              se.maintainability),
    ]:
        _h2(doc, sub_heading)
        _smart_render_narrative(doc, content or "Not addressed.")


def _sep_appendix_glossary(doc: DocxDocument, plan) -> None:
    glossary = plan.appendices.glossary
    if not glossary:
        doc.add_paragraph("Glossary to be populated during program execution.")
        return
    headers = ["Term", "Definition"]
    widths = [Inches(1.5), Inches(4.5)]
    tbl = doc.add_table(rows=len(glossary) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True)
        _shade_cell(hdr[i])
        hdr[i].width = widths[i]
    for r_idx, entry in enumerate(glossary, start=1):
        cells = tbl.rows[r_idx].cells
        _set_cell_text(cells[0], entry.term or "", bold=True)
        _set_cell_text(cells[1], entry.definition or "")
        for i, cell in enumerate(cells):
            cell.width = widths[i]


def _sep_appendix_references(doc: DocxDocument, plan) -> None:
    refs = plan.appendices.references
    if refs:
        _smart_render_bullets(doc, refs)
    else:
        doc.add_paragraph("References to be determined.")


def _sep_appendix_sources(doc: DocxDocument, plan) -> None:
    if plan.sources_used:
        doc.add_paragraph(
            "The following uploaded program documents were consulted in preparing "
            "this AI-assisted Systems Engineering Plan:"
        )
        _smart_render_sources(doc, plan.sources_used)


# ---------------------------------------------------------------------------
# Section renderer registry
# Keys are the exact heading strings from TEMPLATE_REGISTRY section_order.
# ---------------------------------------------------------------------------

_SMART_SECTION_RENDERERS: dict[str, dict[str, object]] = {
    "rfi": {
        "1. Document Overview":        _rfi_document_overview,
        "2. RFI Purpose":              _rfi_rfi_purpose,
        "3. Program Context":          _rfi_program_context,
        "4. MOSA Requirements":        _rfi_mosa_requirements,
        "5. Candidate Module Summary": _rfi_candidate_module_summary,
        "6. Questions to Industry":    _rfi_questions_to_industry,
        "7. Requested Deliverables":   _rfi_requested_deliverables,
        "8. Submission Instructions":  _rfi_submission_instructions,
        "Appendix A: Sources Used":    _rfi_appendix_sources,
    },
    "acq_strategy": {
        "1. Executive Summary":                       _acq_executive_summary,
        "2. Acquisition Approach":                    _acq_acquisition_approach,
        "3. Schedule and Milestones":                 _acq_schedule_milestones,
        "4. Cost Estimates":                          _acq_cost_estimates,
        "5. Candidate Module Summary":                _acq_candidate_module_summary,
        "6. Modular Open Systems Approach (MOSA)":    _acq_mosa,
        "7. Technical Data and Data Rights Strategy": _acq_data_rights,
        "8. Standards and Constraints":               _acq_standards_constraints,
        "9. Test and Verification Strategy":          _acq_test_verification,
        "10. Contracting Strategy":                   _acq_contracting_strategy,
        "11. Risk Register":                          _acq_risk_register,
        "Appendix A: Sources Used":                   _acq_appendix_sources,
    },
    "mcp": {
        "1. Executive Summary":                       _mcp_executive_summary,
        "2. MOSA Objectives":                         _mcp_mosa_objectives,
        "3. Tailoring Rationale":                     _mcp_tailoring_rationale,
        "4. Program Context":                         _mcp_program_context,
        "5. Module Inventory and Conformance Status": _mcp_module_inventory,
        "6. Interface Registry":                      _mcp_interface_registry,
        "7. Standards and Architecture Mapping":      _mcp_standards_architecture_mapping,
        "8. Technical Data and Data Rights":          _mcp_technical_data_rights,
        "9. Conformance Verification Plan":           _mcp_conformance_verification,
        "10. MOSA Risk Register":                     _mcp_risk_register,
        "11. MOSA Self-Assessment Matrix":            _mcp_self_assessment_matrix,
        "12. MOSA Scenario Mapping":                  _mcp_scenario_mapping,
        "Appendix A: Sources Used":                   _mcp_appendix_sources,
    },
    "sep": {
        "1. Executive Summary":            _sep_executive_summary,
        "2. Program Overview":             _sep_program_overview,
        "3. Systems Engineering Strategy": _sep_se_strategy,
        "4. Technical Reviews":            _sep_technical_reviews,
        "5. Requirements Traceability":    _sep_requirements_traceability,
        "6. System Architecture and MOSA": _sep_architecture_mosa,
        "7. Technical Risk Register":      _sep_risk_register,
        "8. Configuration Management":     _sep_config_mgmt,
        "9. Verification and Validation":  _sep_vnv,
        "10. Data Management":             _sep_data_mgmt,
        "11. Specialty Engineering":       _sep_specialty_eng,
        "Appendix A: Glossary":            _sep_appendix_glossary,
        "Appendix B: References":          _sep_appendix_references,
        "Appendix C: Sources Used":        _sep_appendix_sources,
    },
}


def _should_skip_section(plan, doc_type: str, heading: str) -> bool:
    """Return True for optional appendix sections whose content is empty."""
    from document_templates import TEMPLATE_REGISTRY
    tmpl = TEMPLATE_REGISTRY[doc_type]
    if heading in set(tmpl.get("required_sections", [])):
        return False
    # Skip sources appendices when no sources were retrieved
    if heading.lower().startswith("appendix") and "source" in heading.lower():
        return not getattr(plan, "sources_used", None)
    return False


def _render_smart_doc_body(doc: DocxDocument, plan, doc_type: str) -> None:
    """Iterate section_order from the template, emit H1 headings and content.

    Heading strings come exclusively from TEMPLATE_REGISTRY so the document
    always matches the template contract.  The LLM cannot introduce new headings.
    """
    from document_templates import TEMPLATE_REGISTRY
    section_order = TEMPLATE_REGISTRY[doc_type]["section_order"]
    renderers = _SMART_SECTION_RENDERERS.get(doc_type, {})

    first_section = True
    for heading in section_order:
        renderer = renderers.get(heading)
        if renderer is None:
            continue
        if _should_skip_section(plan, doc_type, heading):
            continue
        if not first_section:
            doc.add_page_break()
        _h1(doc, heading)
        renderer(doc, plan)
        first_section = False


# ---------------------------------------------------------------------------
# Public smart builder functions
# ---------------------------------------------------------------------------

def build_smart_rfi(
    program: Program,
    generated_date: str,
    rfi_plan,  # schemas.RfiPlan — imported lazily to avoid circular imports
) -> DocxDocument:
    """Render a DOCX from a validated RfiPlan.

    Headings come from TEMPLATE_REGISTRY["rfi"]["section_order"]; the LLM
    cannot introduce or rename sections.
    """
    doc = DocxDocument()
    _setup_margins(doc)
    _add_title_page(doc, program, generated_date, "Request for Information (AI-Assisted)")
    _render_smart_doc_body(doc, rfi_plan, "rfi")
    _configure_footer(doc, program.name)
    return doc


# ---------------------------------------------------------------------------
# Smart MOSA Conformance Plan builder  (AI-assisted, structured-output)
# ---------------------------------------------------------------------------


def build_smart_mosa_conformance_plan(
    program: Program,
    generated_date: str,
    mosa_plan,  # llm.mosa_schema.MosaPlan – imported lazily to avoid circular imports
) -> DocxDocument:
    """Render a DOCX from a validated MosaPlan.

    Headings come from TEMPLATE_REGISTRY["mcp"]["section_order"]; the LLM
    cannot introduce or rename sections.
    """
    doc = DocxDocument()
    _setup_margins(doc)
    _add_title_page(doc, program, generated_date, "MOSA Conformance Plan (AI-Assisted)")
    _render_smart_doc_body(doc, mosa_plan, "mcp")
    _configure_footer(doc, program.name)
    return doc


# ---------------------------------------------------------------------------
# Smart Acquisition Strategy builder  (AI-assisted, structured-output)
# ---------------------------------------------------------------------------


def build_acq_strategy_smart(
    program: Program,
    brief: ProgramBrief | None,
    answers: dict[str, str | None],
    files: list[ProgramFile],
    modules: list[Module],
    generated_date: str,
    plan,  # schemas.AcqStrategyPlan — imported lazily to avoid circular imports
) -> DocxDocument:
    """Render a DOCX from a validated AcqStrategyPlan.
    Headings come from TEMPLATE_REGISTRY["acq_strategy"]["section_order"]; no freeform sections.
    """
    doc = DocxDocument()
    _setup_margins(doc)
    _add_title_page(doc, program, generated_date, "Acquisition Strategy (AI-Assisted)")
    _render_smart_doc_body(doc, plan, "acq_strategy")
    _configure_footer(doc, program.name)
    return doc


# ---------------------------------------------------------------------------
# Smart SEP builder  (AI-assisted, structured-output)
# ---------------------------------------------------------------------------


def build_sep_smart(
    plan,           # llm.sep_schema.SepPlan — imported lazily to avoid circular imports
    program: Program,
    generated_date: str,
) -> DocxDocument:
    """Render a DOCX from a validated SepPlan.
    Headings come from TEMPLATE_REGISTRY["sep"]["section_order"]; no freeform sections.
    Title page uses plan.title_block for program name, org, and version.
    """
    doc = DocxDocument()
    _setup_margins(doc)

    # ---- SEP title page: extended with org / version from plan.title_block ----
    tb = plan.title_block
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    title_run = title_para.add_run(tb.program_name or program.name)
    title_run.bold = True
    title_run.font.size = Pt(24)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_before = Pt(12)
    sub_run = sub_para.add_run("Systems Engineering Plan (AI-Assisted)")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    if tb.organization:
        org_para = doc.add_paragraph()
        org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        org_para.paragraph_format.space_before = Pt(10)
        org_run = org_para.add_run(tb.organization)
        org_run.font.size = Pt(12)
        org_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta_parts: list[str] = []
    if tb.version:
        meta_parts.append(f"Version {tb.version}")
    meta_parts.append(generated_date)
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.paragraph_format.space_before = Pt(8)
    meta_para.add_run("  |  ".join(meta_parts)).font.size = Pt(11)

    doc.add_page_break()
    _render_smart_doc_body(doc, plan, "sep")
    _configure_footer(doc, program.name)
    return doc
