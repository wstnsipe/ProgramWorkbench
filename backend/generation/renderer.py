"""
renderer.py — DOCX rendering layer for assembled section dicts.

Converts orchestrator output (assembled dict) to a .docx file.
Uses a generic section-aware renderer that handles all doc types.
"""
import logging
import re
from typing import Any

logger = logging.getLogger(__name__)

_ASSUMPTION_RE = re.compile(r'\[ASSUMPTION:[^\]]*\]', re.IGNORECASE)


def _clean(text: str) -> str:
    """Strip inline assumption markers before rendering."""
    return _ASSUMPTION_RE.sub('', str(text)).strip()


def render_document(*, assembled: dict[str, Any], doc_type: str, output_path: str) -> None:
    """
    Write assembled section data to a .docx file at output_path.
    Handles all doc types via the generic section renderer.
    """
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()

    # Standard margins
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.25)
        sec.right_margin = Inches(1.25)

    meta = assembled.get("_meta", {})
    doc_label = {
        "rfi":          "Request for Information (RFI)",
        "acq_strategy": "Acquisition Strategy",
        "sep":          "Systems Engineering Plan (SEP)",
        "mcp":          "MOSA Conformance Plan",
    }.get(doc_type, doc_type.upper().replace("_", " "))

    # Title
    title = doc.add_heading(doc_label, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # MIG reference line if available
    mig_id = meta.get("mig_id")
    if mig_id:
        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ref_para.add_run(f"Reference: {mig_id}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Classification line
    class_para = doc.add_paragraph()
    class_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    class_run = class_para.add_run("UNCLASSIFIED // FOR OFFICIAL USE ONLY")
    class_run.font.size = Pt(9)
    class_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()  # spacer

    # Render sections
    section_num = 0
    for section_name, section_data in assembled.items():
        if section_name.startswith("_") or not isinstance(section_data, dict):
            continue
        section_num += 1
        doc.add_heading(f"{section_num}. {section_name}", level=1)
        _render_section_fields(doc, section_data)

    # Footer on all pages
    for sec in doc.sections:
        footer = sec.footer
        if footer.paragraphs:
            fp = footer.paragraphs[0]
        else:
            fp = footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("UNCLASSIFIED // FOR OFFICIAL USE ONLY")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(output_path)
    logger.info("Rendered %s → %s (%d sections)", doc_type, output_path, section_num)


def _render_section_fields(doc, section_data: dict) -> None:
    """Dispatch each field in a section dict to the appropriate renderer."""
    for field_name, value in section_data.items():
        if value is None or value == "" or value == [] or value == {}:
            continue

        if isinstance(value, str):
            cleaned = _clean(value)
            if cleaned:
                doc.add_paragraph(cleaned)

        elif isinstance(value, list) and value:
            first = value[0]
            if isinstance(first, str):
                for item in value:
                    item_clean = _clean(str(item))
                    if item_clean:
                        doc.add_paragraph(item_clean, style="List Bullet")
            elif isinstance(first, dict):
                _render_dict_list(doc, field_name, value)


def _render_dict_list(doc, field_name: str, items: list) -> None:
    """Render a list of dicts with formatting based on field name."""
    if field_name == "risks":
        _render_risk_table(doc, items)
    elif field_name == "milestones":
        _render_milestone_list(doc, items)
    elif field_name in ("tech_review_schedule",):
        _render_review_schedule(doc, items)
    elif field_name in ("module_sustainability", "module_boundaries",
                        "module_assessments", "module_rights"):
        for entry in items:
            _render_module_entry(doc, entry)
    elif field_name in ("interface_boundaries",):
        for entry in items:
            _render_interface_entry(doc, entry)
    elif field_name in ("test_events", "milestones", "assessment_criteria",
                        "conformance_objectives"):
        for entry in items:
            _render_generic_dict_entry(doc, entry)
    else:
        for entry in items:
            _render_generic_dict_entry(doc, entry)


def _render_risk_table(doc, risks: list) -> None:
    from docx.shared import Inches
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    headers = ["Risk", "P", "I", "Mitigation", "Owner"]
    widths = [Inches(1.6), Inches(0.35), Inches(0.35), Inches(2.7), Inches(0.8)]
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].clear()
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        hdr[i].width = widths[i]
    for risk in risks:
        row = tbl.add_row().cells
        row[0].text = _clean(risk.get("title", ""))
        row[1].text = str(risk.get("probability", ""))
        row[2].text = str(risk.get("impact", ""))
        row[3].text = _clean(risk.get("mitigation", ""))
        row[4].text = str(risk.get("owner", ""))
    doc.add_paragraph()


def _render_milestone_list(doc, milestones: list) -> None:
    from docx.shared import Pt
    for m in milestones:
        name = _clean(m.get("name", ""))
        month = m.get("month_offset", "")
        desc = _clean(m.get("description", ""))
        rationale = _clean(m.get("compression_rationale") or "")
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"Month {month} — {name}: ")
        run.bold = True
        p.add_run(desc)
        if rationale:
            doc.add_paragraph(f"Schedule rationale: {rationale}", style="List Bullet 2")


def _render_review_schedule(doc, reviews: list) -> None:
    from docx.shared import Inches
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    headers = ["Review", "Month", "Entry Criteria", "Exit Criteria"]
    widths = [Inches(0.8), Inches(0.6), Inches(2.3), Inches(2.1)]
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].clear()
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        hdr[i].width = widths[i]
    for r in reviews:
        row = tbl.add_row().cells
        row[0].text = _clean(r.get("name", ""))
        row[1].text = str(r.get("month_offset", ""))
        row[2].text = _clean(r.get("entry_criteria", ""))
        row[3].text = _clean(r.get("exit_criteria", ""))
    doc.add_paragraph()


def _render_module_entry(doc, entry: dict) -> None:
    name = _clean(entry.get("module_name") or entry.get("name") or "")
    if name:
        doc.add_heading(name, level=3)
    for key, value in entry.items():
        if key in ("module_name", "name") or not value:
            continue
        label = key.replace("_", " ").title()
        if isinstance(value, str):
            cleaned = _clean(value)
            if cleaned:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(cleaned)
        elif isinstance(value, list) and value:
            doc.add_paragraph(f"{label}:").runs[0].bold = True if doc.paragraphs else None
            p = doc.add_paragraph(f"{label}:")
            if p.runs:
                p.runs[0].bold = True
            if isinstance(value[0], str):
                for item in value:
                    item_clean = _clean(str(item))
                    if item_clean:
                        doc.add_paragraph(item_clean, style="List Bullet")
            elif isinstance(value[0], dict):
                for sub in value:
                    _render_interface_entry(doc, sub)


def _render_interface_entry(doc, entry: dict) -> None:
    standard = _clean(entry.get("standard_name") or entry.get("interface_standard") or "")
    btype = entry.get("boundary_type", "")
    crosses = _clean(entry.get("what_crosses") or entry.get("boundary_description") or "")
    rights = entry.get("government_rights", "")

    p = doc.add_paragraph(style="List Bullet")
    if standard:
        p.add_run(standard).bold = True
    if btype:
        p.add_run(f" [{btype}]")
    if crosses:
        doc.add_paragraph(f"    Signals/data: {crosses}", style="List Bullet 2")
    if rights:
        doc.add_paragraph(f"    Government rights: {rights}", style="List Bullet 2")


def _render_generic_dict_entry(doc, entry: dict) -> None:
    for key, value in entry.items():
        if not value:
            continue
        label = key.replace("_", " ").title()
        if isinstance(value, str):
            cleaned = _clean(value)
            if cleaned:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{label}: ").bold = True
                p.add_run(cleaned)
        elif isinstance(value, list):
            if value and isinstance(value[0], str):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{label}: ").bold = True
                p.add_run(", ".join(_clean(str(i)) for i in value if i))
